# -*- coding: utf-8 -*-
"""Generate the Levi's Retail (SES x EBR) FSD as DOCX for approval request.

Usage:  python docs/levis/build_fsd_docx.py
Output: docs/levis/FSD - Levis Retail Odoo (SES x EBR) v1.0.docx
All body paragraphs are justified per the approval-document requirement.
"""

import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "FSD - Levis Retail Odoo (SES x EBR) v1.0.docx")

INK = RGBColor(0x1A, 0x23, 0x32)
ACCENT = RGBColor(0x71, 0x4B, 0x67)  # Odoo purple
HEAD_BG = "1A2332"

doc = Document()

# ---- base styles -----------------------------------------------------------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
normal.paragraph_format.space_after = Pt(6)

for level, size in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)):
    st = doc.styles[level]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.color.rgb = INK if level != "Heading 1" else ACCENT
    st.font.bold = True

for section in doc.sections:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def p(text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=None, color=None, space_after=6):
    par = doc.add_paragraph()
    par.alignment = align
    par.paragraph_format.space_after = Pt(space_after)
    run = par.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return par


def bullets(items):
    for it in items:
        par = doc.add_paragraph(style="List Bullet")
        par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if isinstance(it, tuple):  # (bold lead, rest)
            r = par.add_run(it[0])
            r.bold = True
            par.add_run(it[1])
        else:
            par.add_run(it)


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def table(headers, rows, widths=None, header_bg=HEAD_BG):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9.5)
        shade(hdr[i], header_bg)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            par = cells[i].paragraphs[0]
            par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = par.add_run(str(val))
            run.font.size = Pt(9.5)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ============================ COVER =========================================
for _ in range(5):
    doc.add_paragraph()
p("FUNCTIONAL SPECIFICATION DOCUMENT", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=22, color=ACCENT, space_after=2)
p("(FSD)", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, color=ACCENT)
doc.add_paragraph()
p("Implementasi Odoo untuk Bisnis Retail Levi's Indonesia", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
p(
    "PT Sinar Eka Selaras (SES — SAP)  ×  PT EBR (Retail Levi's — Odoo)  ×  POS Principal",
    align=WD_ALIGN_PARAGRAPH.CENTER,
    size=11,
)
for _ in range(3):
    doc.add_paragraph()
table(
    ["Atribut", "Keterangan"],
    [
        ["Judul Dokumen", "FSD — Proses Bisnis & Integrasi Retail Levi's (SES × EBR) pada Odoo"],
        ["Nomor Dokumen", "FSD/ERAJAYA/LEVIS-ODOO/2026/001"],
        ["Versi", "1.0"],
        ["Tanggal", "12 Juni 2026"],
        ["Klasifikasi", "Internal — Erajaya Group"],
        ["Status", "Draft — Diajukan untuk Persetujuan (Approval Request)"],
    ],
    widths=[5, 11],
)
doc.add_page_break()

# ===================== APPROVAL & REVISION ==================================
doc.add_heading("Lembar Persetujuan (Approval Sheet)", level=1)
p(
    "Dokumen ini diajukan untuk mendapatkan persetujuan atas rancangan fungsional "
    "implementasi Odoo bagi entitas retail Levi's (PT EBR) beserta integrasinya dengan "
    "PT Sinar Eka Selaras (SAP) dan POS bawaan principal. Dengan menandatangani lembar "
    "ini, para pihak menyatakan telah membaca, memahami, dan menyetujui ruang lingkup, "
    "kebutuhan fungsional, serta asumsi dan keterbatasan yang tercantum di dalam dokumen ini. "
    "Perubahan terhadap ruang lingkup setelah persetujuan akan dikelola melalui mekanisme "
    "change request."
)
table(
    ["Peran", "Nama", "Jabatan / Fungsi", "Tanda Tangan", "Tanggal"],
    [
        ["Disusun oleh (Prepared by)", "", "Functional Consultant — Erajaya", "", ""],
        ["Diperiksa oleh (Reviewed by)", "", "IT / ERP Lead — Erajaya", "", ""],
        ["Diperiksa oleh (Reviewed by)", "", "Finance & Accounting — PT EBR", "", ""],
        ["Disetujui oleh (Approved by)", "", "Business Owner — Levi's Retail", "", ""],
        ["Disetujui oleh (Approved by)", "", "Manajemen PT SES", "", ""],
    ],
    widths=[4.4, 3.2, 4.6, 2.5, 2.0],
)

doc.add_heading("Riwayat Revisi", level=2)
table(
    ["Versi", "Tanggal", "Penyusun", "Deskripsi Perubahan"],
    [
        ["1.0", "12 Juni 2026", "Tim Erajaya", "Penerbitan awal untuk pengajuan persetujuan."],
    ],
    widths=[2, 3, 4, 7.2],
)
doc.add_page_break()

# ============================ 1. PENDAHULUAN ================================
doc.add_heading("1. Pendahuluan", level=1)

doc.add_heading("1.1 Latar Belakang", level=2)
p(
    "Erajaya Group melalui PT Sinar Eka Selaras (SES) bertindak sebagai distributor resmi "
    "produk Levi's di Indonesia, dengan seluruh proses distribusi dan pengadaan dijalankan "
    "pada sistem SAP. Operasional penjualan retail dijalankan oleh PT EBR sebagai entitas "
    "retail Levi's, sementara seluruh toko beroperasi menggunakan sistem POS bawaan principal "
    "yang menjadi standar global Levi's. PT EBR membutuhkan sistem pembukuan dan pelaporan "
    "yang andal namun tidak menggantikan sistem transaksional toko; untuk kebutuhan tersebut "
    "dipilih Odoo sebagai sistem ERP pencatatan (system of record akuntansi) PT EBR. Data "
    "operasional toko — goods receipt, penjualan POS beserta tender pembayarannya, dan mutasi "
    "persediaan — akan dimirror ke Odoo melalui mekanisme file Excel/CSV pada fase awal dan "
    "melalui FTP/SFTP Erajaya pada fase berikutnya."
)

doc.add_heading("1.2 Tujuan Dokumen", level=2)
p(
    "Dokumen ini mendefinisikan kebutuhan fungsional (functional requirements) implementasi "
    "Odoo untuk PT EBR, mencakup proses bisnis end-to-end antara Levi's Principal, PT SES, "
    "toko retail, dan PT EBR; modul Odoo yang akan diinstal beserta kustomisasinya; "
    "spesifikasi file data yang dipertukarkan; serta konfigurasi data yang harus disiapkan "
    "sebelum go-live. Dokumen ini menjadi dasar persetujuan (approval) sebelum tahap "
    "konfigurasi, pengembangan, dan migrasi data dilaksanakan."
)

doc.add_heading("1.3 Ruang Lingkup", level=2)
p("Ruang lingkup pekerjaan yang tercakup dalam dokumen ini (in-scope):", bold=True)
bullets(
    [
        "Pencatatan Purchase Order PT EBR kepada PT SES di Odoo, termasuk approval berjenjang, "
        "pencatatan vendor bill, dan pembayaran kepada SES.",
        "Mirroring data operasional toko dari POS principal ke Odoo: goods receipt toko, "
        "penjualan POS dan tender pembayaran, mutasi persediaan, posisi stok on-hand, serta "
        "data diskon/promosi.",
        "Pembangunan adapter import Excel/CSV (modul custom_retail_import) dengan profil per "
        "jenis file, pratinjau dry-run, deduplikasi, log audit, dan eksekusi asinkron.",
        "Penyediaan feed otomatis melalui FTP/SFTP Erajaya pada Phase 2.",
        "Penyiapan master data dan konfigurasi awal: company, chart of accounts, rekening bank "
        "dan jurnal, pajak, vendor, master produk, warehouse per toko, dan saldo awal stok.",
        "Pelaporan keuangan dan rekonsiliasi (penjualan, HPP, persediaan, settlement tender terhadap penerimaan bank).",
    ]
)
p("Hal-hal berikut berada di luar ruang lingkup (out-of-scope):", bold=True)
bullets(
    [
        "Proses transaksional di toko. Toko tetap menggunakan POS bawaan principal; Odoo tidak "
        "menggantikan fungsi POS toko.",
        "Proses internal SAP PT SES, termasuk pengadaan SES kepada Levi's Principal dan goods "
        "receipt di DC SES; proses tersebut dijelaskan hanya sebagai konteks alur bisnis.",
        "Integrasi API real-time dua arah antara SAP, POS principal, dan Odoo (di luar mekanisme "
        "file/SFTP yang didefinisikan).",
        "Perubahan proses bisnis pada sisi Levi's Principal.",
    ]
)

doc.add_heading("1.4 Definisi dan Singkatan", level=2)
table(
    ["Istilah", "Definisi"],
    [
        ["SES", "PT Sinar Eka Selaras — distributor produk Levi's (Erajaya Group), bersistem SAP."],
        ["EBR", "PT EBR — entitas retail Levi's yang pembukuannya dijalankan di Odoo."],
        ["Principal", "Levi's selaku pemilik merek dan pemasok produk."],
        ["POS Principal", "Sistem point-of-sale bawaan principal yang digunakan seluruh toko Levi's."],
        ["GR", "Goods Receipt — penerimaan barang."],
        ["Tender", "Metode pembayaran pada transaksi POS (tunai, EDC, e-wallet) beserta settlement-nya."],
        ["PO / SO", "Purchase Order / Sales Order."],
        ["CoA", "Chart of Accounts — bagan akun."],
        ["Mirroring", "Replikasi data transaksi dari sistem sumber (POS principal) ke Odoo melalui file."],
        ["FSD", "Functional Specification Document."],
    ],
    widths=[3.5, 12.7],
)

doc.add_heading("1.5 Referensi", level=2)
bullets(
    [
        "Diagram proses bisnis: docs/levis/business-process-flow.html.",
        "Paket data dari principal/SES: X101 Material Master, X20 On-hand Inventory, X24DN "
        "Retail Sales Detail, X31 Discount Journal, X32P Stock Movement, X70D Tender Detail, "
        "X70T Tender Settlement, Store Master, CoA EBR.",
        "Runbook onboarding data: scripts/tenants/levis/README.md.",
    ]
)
doc.add_page_break()

# ===================== 2. GAMBARAN PROSES BISNIS ============================
doc.add_heading("2. Gambaran Umum Proses Bisnis", level=1)
p(
    "Rantai proses melibatkan empat aktor dengan tiga sistem yang berbeda. Levi's Principal "
    "bertindak sebagai pemasok produk dan penyedia sistem POS toko. PT SES menjalankan fungsi "
    "distribusi dengan SAP sebagai sistem pencatatannya. Toko retail beroperasi sepenuhnya "
    "pada POS bawaan principal. PT EBR menjalankan pembukuan dan pelaporan keuangan pada "
    "Odoo, yang menerima mirroring data operasional toko. Secara keseluruhan terdapat tiga "
    "alur utama sebagaimana diuraikan berikut."
)

doc.add_heading("2.1 Flow A — Pengadaan SES kepada Levi's Principal (SAP)", level=2)
p(
    "PT SES menerbitkan Purchase Order kepada Levi's Principal melalui SAP. Principal "
    "mengonfirmasi pesanan dan mengirimkan barang beserta dokumen pengiriman dan invoice. "
    "Atas barang yang diterima, SES melakukan Goods Receipt di SAP sehingga stok bertambah "
    "pada gudang/DC SES. Proses dilanjutkan dengan verifikasi invoice (three-way match antara "
    "PO, GR, dan invoice) serta pemrosesan pembayaran kepada principal di SAP. Seluruh "
    "tahapan Flow A berada pada sistem SAP milik SES dan dijelaskan dalam dokumen ini sebagai "
    "konteks; tidak ada pekerjaan konfigurasi Odoo pada alur ini."
)
table(
    ["Langkah", "Aktivitas", "Sistem"],
    [
        ["A1", "SES membuat PO produk Levi's dan mengirimkannya kepada principal.", "SAP — SES"],
        ["A2", "Principal mengonfirmasi pesanan dan mengirim barang beserta dokumen.", "Principal"],
        ["A3", "Goods Receipt atas kiriman principal; stok bertambah di DC SES.", "SAP — SES"],
        ["A4", "Verifikasi invoice (three-way match) dan pembayaran kepada principal.", "SAP — SES"],
    ],
    widths=[2, 10.7, 3.5],
)

doc.add_heading("2.2 Flow B — Pembelian EBR kepada SES dan GR di Toko", level=2)
p(
    "PT EBR menerbitkan Purchase Order kepada PT SES melalui Odoo. SES menerima PO tersebut, "
    "membentuk Sales Order dan Delivery di SAP, kemudian mengirimkan barang dari DC ke toko "
    "tujuan. Penerimaan barang fisik di toko dilakukan menggunakan POS bawaan principal "
    "(Goods Receipt toko), sehingga stok toko bertambah pada sistem POS. Data GR toko "
    "selanjutnya dimirror ke Odoo untuk memvalidasi penerimaan atas PO EBR, sehingga stok dan "
    "hutang usaha EBR tercatat dengan benar. Alur ditutup dengan pencatatan vendor bill dari "
    "SES di Odoo, pencocokan terhadap PO dan GR, serta pemrosesan pembayaran kepada SES."
)
table(
    ["Langkah", "Aktivitas", "Sistem"],
    [
        ["B1", "EBR membuat PO pembelian produk Levi's kepada SES (intercompany B2B).", "Odoo — EBR"],
        ["B2", "SES menerima PO, membuat SO dan Delivery, mengirim barang ke toko.", "SAP — SES"],
        ["B3", "Toko menerima barang; Goods Receipt dilakukan pada POS principal.", "POS Principal"],
        ["B4", "Data GR toko dimirror ke Odoo untuk validasi receipt atas PO EBR.", "Mirroring"],
        ["B5", "Vendor bill SES dicatat, dicocokkan dengan PO dan GR, lalu dibayar.", "Odoo — EBR"],
    ],
    widths=[2, 10.7, 3.5],
)

doc.add_heading("2.3 Flow C — Operasional Toko dan Mirroring ke Odoo", level=2)
p(
    "Seluruh transaksi operasional toko terjadi pada POS bawaan principal, meliputi goods "
    "receipt, penjualan retail beserta tender pembayaran dan settlement-nya, serta mutasi "
    "persediaan lainnya (transfer, retur, adjustment, posisi on-hand). Report standar "
    "principal diekspor secara periodik dalam format Excel/CSV, kemudian ditransfer ke Odoo: "
    "pada Phase 1 melalui unggah manual, dan pada Phase 2 melalui drop otomatis ke server "
    "FTP/SFTP Erajaya yang dipoll terjadwal oleh Odoo. Adapter import pada Odoo memetakan "
    "setiap report ke dokumen yang sesuai (receipt, POS/sales order, stock move, payment "
    "register), membentuk jurnal otomatis untuk penjualan, HPP, dan mutasi persediaan, serta "
    "menjalankan rekonsiliasi settlement tender terhadap penerimaan bank."
)
table(
    ["Langkah", "Aktivitas", "Sistem"],
    [
        ["C1", "Transaksi di POS: GR toko, penjualan & tender, mutasi stok, on-hand.", "POS Principal"],
        ["C2", "Ekspor report standar principal per periode (Excel/CSV).", "Export"],
        ["C3", "Transfer file: Phase 1 unggah manual; Phase 2 FTP/SFTP Erajaya.", "File / SFTP"],
        ["C4", "Adapter import memetakan report ke dokumen Odoo.", "Odoo — EBR"],
        ["C5", "Posting jurnal otomatis dan rekonsiliasi tender vs bank.", "Odoo — EBR"],
    ],
    widths=[2, 10.7, 3.5],
)
doc.add_page_break()

# ===================== 3. KEBUTUHAN FUNGSIONAL ==============================
doc.add_heading("3. Kebutuhan Fungsional", level=1)
p(
    "Kebutuhan fungsional berikut menjadi dasar konfigurasi dan pengembangan. Prioritas "
    "dinyatakan sebagai M (Mandatory — wajib pada go-live) atau P2 (Phase 2 / menunggu "
    "keputusan)."
)
table(
    ["ID", "Kebutuhan Fungsional", "Deskripsi", "Prioritas"],
    [
        [
            "FR-01",
            "Purchase Order EBR → SES",
            "Pengguna EBR dapat membuat, mengubah, dan mengonfirmasi PO kepada SES di Odoo dengan "
            "approval berjenjang sesuai matriks otorisasi (custom_approval_engine).",
            "M",
        ],
        [
            "FR-02",
            "Mirror GR toko ke Odoo",
            "GR yang dilakukan di POS principal direplikasi ke Odoo sebagai receipt atas PO EBR, "
            "menambah stok warehouse toko dan membentuk dasar pencocokan vendor bill.",
            "M",
        ],
        [
            "FR-03",
            "Vendor bill & pembayaran SES",
            "Vendor bill dari SES dicatat dan dicocokkan terhadap PO dan GR (three-way match) "
            "sebelum pembayaran diproses.",
            "M",
        ],
        [
            "FR-04",
            "Mirror penjualan POS (X24DN)",
            "Detail penjualan harian per toko diimpor dan membentuk jurnal penjualan beserta HPP. "
            "Representasi dokumen (pos.order), kedalaman histori, dan tax mapping menunggu "
            "keputusan Phase 5; executor menahan posting hingga diaktifkan.",
            "P2",
        ],
        [
            "FR-05",
            "Mirror tender & settlement (X70D/X70T)",
            "Rincian metode pembayaran dan settlement-nya diimpor sebagai payment register untuk "
            "rekonsiliasi terhadap penerimaan bank per acquirer.",
            "P2",
        ],
        [
            "FR-06",
            "Mirror mutasi persediaan (X32P) & on-hand (X20)",
            "Mutasi persediaan bernilai (GR, transfer, retur, adjustment) direplikasi sebagai stock "
            "move; X20 dipakai untuk saldo awal (one-shot) dan validasi stok periodik.",
            "M",
        ],
        [
            "FR-07",
            "Import master data",
            "Master produk X101 (artikel, varian size/inseam, barcode, kategori, harga), CoA, "
            "company detail, dan store master dapat diimpor melalui profil import.",
            "M",
        ],
        [
            "FR-08",
            "Wizard import dengan dry-run",
            "Setiap unggah file menampilkan pratinjau hasil parsing (jumlah baris, sampel mapping) "
            "sebelum commit; file besar diproses asinkron via queue_job.",
            "M",
        ],
        [
            "FR-09",
            "Deduplikasi & idempotensi",
            "File yang sama (SHA256) ditolak bila diimpor ulang; re-import record menghasilkan "
            "update melalui external ID, bukan duplikasi.",
            "M",
        ],
        [
            "FR-10",
            "Log audit import",
            "Setiap import tercatat (pengguna, waktu, jumlah baris, error) dan file sumber "
            "diarsipkan sebagai attachment untuk kebutuhan audit.",
            "M",
        ],
        [
            "FR-11",
            "SFTP feed otomatis",
            "Odoo melakukan polling terjadwal ke FTP/SFTP Erajaya per jenis file (pola nama), "
            "memproses file baru melalui executor yang sama dengan unggah manual, dan memberi "
            "notifikasi bila parsing gagal.",
            "P2",
        ],
        [
            "FR-12",
            "Klaim promosi (X31)",
            "Data diskon/promo diimpor sebagai dasar klaim promosi kepada principal/SES.",
            "P2",
        ],
        [
            "FR-13",
            "Pelaporan keuangan",
            "Laporan keuangan lengkap (Excel/PDF) tersedia melalui custom_accounting_full; dokumen "
            "transaksi (Invoice, Quotation, PO) menggunakan template PDF ber-branding per company.",
            "M",
        ],
        [
            "FR-14",
            "Pajak Indonesia",
            "PPN keluaran/masukan dikonfigurasi sesuai tarif berlaku; integrasi Coretax/e-Faktur "
            "dan bukti potong melalui custom_coretax dan custom_coretax_bupot.",
            "M",
        ],
    ],
    widths=[1.6, 4.0, 8.6, 2.0],
)
doc.add_page_break()

# ===================== 4. MODUL & KUSTOMISASI ===============================
doc.add_heading("4. Modul Odoo yang Diinstal dan Kustomisasi", level=1)
p(
    'Tenant EBR diprovision menggunakan Industry Pack "Retail / POS" pada platform Odoo '
    "Erajaya (deployment satu aksi). Pack tersebut memasang modul standar Odoo Community "
    "beserta modul custom platform; di atasnya dipasang modul yang dibangun khusus untuk "
    "kebutuhan Levi's."
)

doc.add_heading("4.1 Modul Standar Odoo (Community Edition)", level=2)
table(
    ["Modul", "Fungsi pada Implementasi Ini"],
    [
        ["purchase", "Purchase Order EBR → SES; pencocokan vendor bill (PO–GR–Invoice)."],
        ["stock (Inventory)", "Warehouse per toko, receipt (mirror GR), stock move, inventory valuation."],
        ["account (Accounting)", "CoA, jurnal, AP/AR, rekonsiliasi bank."],
        ["point_of_sale", "Representasi transaksi POS hasil mirroring (bukan POS transaksional toko)."],
        ["sale_management", "Sales order dan pricelist."],
        ["crm, website_sale", "Bawaan pack retail; diaktifkan sesuai kebutuhan."],
    ],
    widths=[4.5, 11.7],
)

doc.add_heading("4.2 Modul Custom Platform (Bawaan Pack Retail)", level=2)
table(
    ["Modul", "Fungsi"],
    [
        ["custom_accounting_full", "Pelengkap akuntansi CE→EE: laporan keuangan lengkap (Excel/PDF)."],
        ["custom_coretax", "Integrasi pajak Indonesia: e-Faktur/Coretax."],
        ["custom_coretax_bupot", "Bukti potong pajak."],
        ["custom_pos_id", "Lokalisasi POS Indonesia (format struk, pembayaran lokal)."],
        ["custom_retail_pos_offline", "Kapabilitas POS offline (cadangan; toko tetap menggunakan POS principal)."],
        ["custom_wms_putaway", "Putaway gudang/DC."],
        ["custom_wms_cycle_count", "Cycle count persediaan."],
        ["custom_approval_engine", "Approval berjenjang untuk PO dan dokumen lainnya."],
    ],
    widths=[5.2, 11],
)

doc.add_heading("4.3 Modul Custom Khusus Levi's", level=2)
table(
    ["Komponen", "Fungsi"],
    [
        ["custom_retail_import", "Adapter mirroring Excel/CSV → Odoo dan SFTP feed; komponen inti integrasi Flow C."],
        ["custom_report_templates", "Template PDF ber-branding per company: Invoice, Quotation, PO."],
        [
            "queue_job (OCA)",
            "Eksekusi import file besar secara asinkron (X101 ± 159 ribu SKU) tanpa terputus timeout proxy.",
        ],
        ["openpyxl, paramiko", "Library Python untuk parsing Excel dan koneksi SFTP (ditambahkan pada image Odoo)."],
    ],
    widths=[5.2, 11],
)

doc.add_heading("4.4 Rincian Fungsional custom_retail_import", level=2)
p(
    "Modul custom_retail_import dibangun untuk onboarding Levi's dan dirancang reusable untuk "
    "tenant retail lain. Kemampuan utamanya adalah sebagai berikut."
)
bullets(
    [
        (
            "Import profile per jenis file. ",
            "Setiap report (X101, X20, X24, X70D, CoA, Company, Store Master) memiliki profil "
            "tersendiri: format xlsx/csv, baris awal data untuk file bergaya report, mapping kolom "
            "dalam JSON, aturan parsing tanggal/desimal, serta perbaikan karakter rusak (U+FFFD → ®).",
        ),
        (
            "Upload wizard dengan dry-run preview. ",
            "Pengguna mengunggah file lalu melihat pratinjau hasil parsing sebelum commit, sehingga "
            "kesalahan mapping terdeteksi sebelum posting.",
        ),
        (
            "Deduplikasi dan idempotensi. ",
            "Hash SHA256 per file menolak file identik yang diimpor dua kali; external ID per record "
            "menjadikan re-import bersifat update, bukan duplikasi.",
        ),
        ("Log audit dan arsip file. ", "Setiap import tercatat lengkap dan file sumber disimpan sebagai attachment."),
        (
            "SFTP feed poller. ",
            "Cron terjadwal menarik file baru dari FTP/SFTP Erajaya per jenis file dan memprosesnya "
            "melalui executor yang sama dengan unggah manual.",
        ),
        ("Eksekusi asinkron. ", "File besar diproses di background melalui queue_job."),
        (
            "Phase-5 gating. ",
            "Executor X24 (penjualan) dan X70D (tender) telah mampu mem-parsing dan mengelompokkan "
            "data, namun menahan posting hingga keputusan bisnis Phase 5 (representasi POS, "
            "kedalaman histori, tax mapping) ditetapkan.",
        ),
    ]
)
doc.add_page_break()

# ===================== 5. SPESIFIKASI FILE FEED =============================
doc.add_heading("5. Spesifikasi File Feed Mirroring", level=1)
p(
    "Report standar dari sistem principal berikut menjadi sumber data mirroring ke Odoo. "
    "Seluruh file berformat Excel/CSV pada Phase 1 dan dikirim melalui FTP/SFTP Erajaya pada "
    "Phase 2 tanpa perubahan struktur."
)
table(
    ["Report / File", "Isi Data", "Tujuan di Odoo", "Frekuensi"],
    [
        [
            "X101_Material_Master",
            "Master produk Levi's (artikel, SKU, barcode, harga).",
            "Product master (create/update produk).",
            "Saat ada perubahan",
        ],
        [
            "X20_Current_Onhand_Inventory",
            "Posisi stok on-hand per toko.",
            "Saldo awal (one-shot) dan validasi stok/inventory adjustment.",
            "Harian / periodik",
        ],
        [
            "X24DN_Retail_Sales_Detail",
            "Detail transaksi penjualan POS per toko.",
            "Pencatatan penjualan (jurnal penjualan + HPP).",
            "Harian (EOD)",
        ],
        [
            "X70D_Tender_Detail / X70T_Tender_Settlement",
            "Rincian metode pembayaran (tunai, EDC, e-wallet) dan settlement-nya.",
            "Payment register dan rekonsiliasi bank.",
            "Harian (EOD)",
        ],
        [
            "X32P_Stock_Movement_With_Price",
            "Mutasi persediaan bernilai: GR toko, transfer, retur, adjustment.",
            "Stock move dan receipt (termasuk GR atas PO EBR → SES).",
            "Harian / periodik",
        ],
        [
            "X31_Discount_Journal",
            "Data diskon/promo.",
            "Dasar klaim promosi kepada principal/SES.",
            "Bulanan / per program",
        ],
        [
            "Store Master",
            "Daftar 24 lokasi (1 DC + 23 toko).",
            "Pembentukan warehouse per toko dan crosswalk kode toko.",
            "Saat ada perubahan",
        ],
        ["CoA EBR", "Bagan akun (code, name, account_type).", "Chart of Accounts PT EBR.", "Satu kali (go-live)"],
    ],
    widths=[4.4, 4.6, 4.7, 2.5],
)

# ===================== 6. KONFIGURASI DATA ==================================
doc.add_heading("6. Konfigurasi Data dan Urutan Setup", level=1)
p(
    "Penyiapan data tenant EBR mengikuti urutan berikut: konfigurasi dasar terlebih dahulu "
    "(company, CoA, bank, pajak), dilanjutkan master data, kemudian saldo awal, dan terakhir "
    "transaksi harian yang mengalir melalui mirroring. Setiap langkah pada tabel di bawah "
    "merupakan prasyarat bagi langkah setelahnya."
)
table(
    ["No", "Data", "File / Sumber", "Isi yang Dibutuhkan", "Cara Load"],
    [
        [
            "1",
            "Company Detail (EBR)",
            "Sheet legal entity (format key/value; acuan: PT Sinar Eka Selaras (SES).xlsx).",
            "Nama legal PT, NPWP, alamat lengkap, telepon, email, logo perusahaan, mata uang IDR.",
            "Profil import levis_company atau manual pada Settings ▸ Companies.",
        ],
        [
            "2",
            "Chart of Accounts",
            "CoA EBR.xlsx (lengkap dan clean).",
            "Kolom code, name, account_type (enum Odoo); mencakup akun persediaan, HPP, penjualan, kas/bank, AP/AR.",
            "Profil import levis_coa atau Accounting ▸ CoA ▸ Import.",
        ],
        [
            "3",
            "Bank Account & Journals",
            "Data rekening dari Finance EBR.",
            "Rekening bank EBR (bank, nomor rekening, atas nama) → jurnal Bank per rekening; jurnal "
            "Cash per toko untuk tender tunai; jurnal Sales/Purchase/Inventory; mapping akun "
            "settlement per jenis tender (tunai, EDC per acquirer, e-wallet).",
            "Manual pada Accounting ▸ Configuration (satu kali).",
        ],
        [
            "4",
            "Pajak (PPN)",
            "Konfigurasi tax Odoo dan data Coretax.",
            "PPN keluaran/masukan sesuai tarif berlaku; NPWP dan sertifikat Coretax untuk e-Faktur. "
            "Tax mapping penjualan POS (harga inklusif/eksklusif PPN) merupakan keputusan Phase 5.",
            "custom_coretax dan konfigurasi tax.",
        ],
        [
            "5",
            "Vendor: PT SES",
            "Data vendor dari Procurement.",
            "SES sebagai vendor pada Odoo EBR: NPWP, alamat, term of payment, rekening tujuan pembayaran.",
            "Manual atau import partner (satu kali).",
        ],
        [
            "6",
            "Master Produk",
            "X101_Material_Master.xlsx (full: 14.885 artikel / 159.658 SKU; atribut Size 89 dan Inseam 24).",
            "Artikel, varian (size/inseam), barcode, kategori (170), harga.",
            "Profil import levis_x101 (asinkron, ± 30 menit).",
        ],
        [
            "7",
            "Store Master → Warehouse",
            "LEVI'S - STORE MASTER DATA (for Odoo).xlsx.",
            "24 lokasi (1 DC + 23 toko) → satu warehouse per toko. Membutuhkan kode toko SAP untuk "
            "seluruh 24 toko sebagai kunci join semua file X (saat ini baru 1 kode diketahui).",
            "Script load store / profil store_master.",
        ],
        [
            "8",
            "Saldo Awal Stok",
            "X20_Current_Onhand_Inventory.",
            "On-hand per toko per SKU pada tanggal cut-over (one-shot; dikunci agar tidak terimpor dua kali).",
            "Profil import levis_x20 (prasyarat: No. 6 dan 7 selesai).",
        ],
        [
            "9",
            "Transaksi Harian",
            "X24DN (sales), X70D/X70T (tender), X32P (stock movement), X31 (discount).",
            "Mengalir rutin pasca go-live sesuai Flow C; X24/X70D menunggu keputusan Phase 5.",
            "Wizard upload (Phase 1) → SFTP feed (Phase 2).",
        ],
    ],
    widths=[1.0, 3.2, 3.6, 5.4, 3.0],
)
doc.add_page_break()

# ===================== 7. ROADMAP ===========================================
doc.add_heading("7. Roadmap Integrasi", level=1)
doc.add_heading("7.1 Phase 1 — Mirroring Manual Excel/CSV (Berjalan)", level=2)
p(
    "Pada fase ini tim toko/operasional mengekspor report dari sistem principal dalam format "
    "Excel/CSV, kemudian mengunggahnya ke Odoo melalui wizard import dengan mapping kolom "
    "yang telah ditetapkan pada profil. Validasi dilakukan sebelum posting, meliputi jumlah "
    "baris, total nilai, dan exception report. Fase ini sesuai untuk masa awal pembukaan toko "
    "dan stabilisasi master data."
)
doc.add_heading("7.2 Phase 2 — FTP/SFTP Erajaya (Berikutnya)", level=2)
p(
    "Pada fase ini file report di-drop secara otomatis ke server FTP/SFTP Erajaya sesuai "
    "jadwal (EOD atau intra-day). Odoo melakukan polling terjadwal — mengambil file, mem-parsing, "
    "mengimpor, lalu mengarsipkan ke folder incoming/processed/error — dan mengirimkan "
    "notifikasi otomatis apabila file gagal diparse atau terdapat selisih rekonsiliasi. Fase "
    "ini menghilangkan proses manual sehingga data toko pada Odoo selalu mutakhir pada H+0 "
    "hingga H+1."
)

# ===================== 8. ASUMSI, DEPENDENSI, OPEN ITEMS ====================
doc.add_heading("8. Asumsi, Dependensi, dan Isu Terbuka", level=1)
doc.add_heading("8.1 Asumsi", level=2)
bullets(
    [
        "Odoo PT EBR berperan sebagai sistem pembukuan (system of record akuntansi), bukan "
        "sistem transaksional toko; sumber transaksi tetap pada POS principal dan SAP SES.",
        "Struktur report principal (X101, X20, X24DN, X31, X32P, X70D, X70T) tidak berubah "
        "tanpa pemberitahuan; perubahan struktur ditangani melalui penyesuaian profil import.",
        "Kode toko SAP merupakan kunci penghubung seluruh file X terhadap warehouse Odoo.",
        "Akses server FTP/SFTP Erajaya (host, kredensial, struktur folder) disediakan sebelum Phase 2 dimulai.",
    ]
)
doc.add_heading("8.2 Isu Terbuka (Harus Diselesaikan Sebelum Go-Live Penuh)", level=2)
(
    table(
        ["No", "Isu", "Dampak", "Tindak Lanjut"],
        [
            [
                "1",
                "File X20/X24/X70D yang diterima masih sampel satu toko (store 14694 — Tunjungan Plaza 3).",
                "Saldo awal dan penjualan hanya dapat dimuat untuk satu toko.",
                "Minta ekstrak lengkap seluruh toko kepada principal/SES, atau aktifkan langsung feed SFTP.",
            ],
            [
                "2",
                "Store Master memuat nama toko tanpa kode toko untuk 23 dari 24 lokasi.",
                "Crosswalk kode toko → warehouse tidak lengkap; file X tidak dapat di-join untuk 23 toko.",
                "Minta store master lengkap dengan kode SAP; proses load store bersifat idempoten dan dapat diulang.",
            ],
            [
                "3",
                "Keputusan Phase 5 belum ditetapkan: representasi dokumen POS, kedalaman histori "
                "penjualan, dan tax mapping (harga inklusif/eksklusif PPN).",
                "Posting penjualan (X24) dan tender (X70D) tertahan oleh gating pada executor.",
                "Workshop keputusan bersama Finance EBR dan tim pajak; setelah ditetapkan, gating dibuka.",
            ],
        ],
        widths=[1.0, 5.4, 4.6, 5.2],
    ),
)

p(
    "Dengan disetujuinya dokumen ini, tim implementasi akan melanjutkan ke tahap konfigurasi, "
    "migrasi data, dan persiapan go-live sesuai ruang lingkup, urutan setup, dan roadmap yang "
    "telah diuraikan. Hal-hal yang berada di luar ruang lingkup atau yang timbul setelah "
    "persetujuan akan dikelola melalui mekanisme change request."
)

doc.save(OUT)
print("Saved:", OUT)
