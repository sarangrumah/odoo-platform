# -*- coding: utf-8 -*-
"""Generate three Levi's Retail (SES x EBR) project documents as DOCX:

  1. TSD  — Technical Specification Document (deep detail on custom_retail_import)
  2. Project Charter
  3. Business Workflow

Usage:  python docs/levis/build_tsd_charter_workflow_docx.py
All body paragraphs are justified. Companion of build_fsd_docx.py (FSD v1.0).
"""

import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
INK = RGBColor(0x1A, 0x23, 0x32)
ACCENT = RGBColor(0x71, 0x4B, 0x67)
HEAD_BG = "1A2332"


class Builder:
    def __init__(self):
        self.doc = Document()
        normal = self.doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)
        normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        normal.paragraph_format.space_after = Pt(6)
        for level, size in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)):
            st = self.doc.styles[level]
            st.font.name = "Calibri"
            st.font.size = Pt(size)
            st.font.color.rgb = ACCENT if level == "Heading 1" else INK
            st.font.bold = True
        for section in self.doc.sections:
            section.top_margin = Cm(2.2)
            section.bottom_margin = Cm(2.2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

    def p(self, text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=None, color=None, space_after=6):
        par = self.doc.add_paragraph()
        par.alignment = align
        par.paragraph_format.space_after = Pt(space_after)
        run = par.add_run(text)
        run.bold = bold
        run.italic = italic
        if size:
            run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
        return par

    def h(self, text, level):
        self.doc.add_heading(text, level=level)

    def bullets(self, items):
        for it in items:
            par = self.doc.add_paragraph(style="List Bullet")
            par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if isinstance(it, tuple):
                r = par.add_run(it[0])
                r.bold = True
                par.add_run(it[1])
            else:
                par.add_run(it)

    @staticmethod
    def _shade(cell, hexcolor):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), hexcolor)
        tcPr.append(shd)

    def table(self, headers, rows, widths=None):
        t = self.doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = t.rows[0].cells
        for i, htxt in enumerate(headers):
            hdr[i].text = ""
            run = hdr[i].paragraphs[0].add_run(htxt)
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9.5)
            self._shade(hdr[i], HEAD_BG)
        for row in rows:
            cells = t.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = ""
                par = cells[i].paragraphs[0]
                par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                run = par.add_run(str(val))
                run.font.size = Pt(9.5)
        if widths:
            for i, w in enumerate(widths):
                for row in t.rows:
                    row.cells[i].width = Cm(w)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)
        return t

    def cover(self, kind, title, subtitle, number):
        for _ in range(5):
            self.doc.add_paragraph()
        self.p(kind, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=22, color=ACCENT, space_after=2)
        self.doc.add_paragraph()
        self.p(title, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
        self.p(subtitle, align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
        for _ in range(3):
            self.doc.add_paragraph()
        self.table(
            ["Atribut", "Keterangan"],
            [
                ["Judul Dokumen", title],
                ["Nomor Dokumen", number],
                ["Versi", "1.0"],
                ["Tanggal", "12 Juni 2026"],
                ["Klasifikasi", "Internal — Erajaya Group"],
                ["Status", "Draft — Diajukan untuk Persetujuan (Approval Request)"],
                ["Dokumen Induk", "FSD/ERAJAYA/LEVIS-ODOO/2026/001 v1.0"],
            ],
            widths=[5, 11],
        )
        self.doc.add_page_break()

    def approval(self, intro, rows):
        self.h("Lembar Persetujuan (Approval Sheet)", 1)
        self.p(intro)
        self.table(
            ["Peran", "Nama", "Jabatan / Fungsi", "Tanda Tangan", "Tanggal"], rows, widths=[4.4, 3.2, 4.6, 2.5, 2.0]
        )
        self.h("Riwayat Revisi", 2)
        self.table(
            ["Versi", "Tanggal", "Penyusun", "Deskripsi Perubahan"],
            [
                ["1.0", "12 Juni 2026", "Tim Erajaya", "Penerbitan awal untuk pengajuan persetujuan."],
            ],
            widths=[2, 3, 4, 7.2],
        )
        self.doc.add_page_break()

    def save(self, filename):
        path = os.path.join(HERE, filename)
        self.doc.save(path)
        print("Saved:", path)


SUBTITLE = "PT Sinar Eka Selaras (SES — SAP)  ×  PT EBR (Retail Levi's — Odoo)  ×  POS Principal"

APPROVAL_INTRO = (
    "Dokumen ini diajukan untuk mendapatkan persetujuan. Dengan menandatangani lembar ini, "
    "para pihak menyatakan telah membaca, memahami, dan menyetujui isi dokumen. Perubahan "
    "setelah persetujuan dikelola melalui mekanisme change request."
)

APPROVAL_ROWS = [
    ["Disusun oleh (Prepared by)", "", "Functional/Technical Consultant — Erajaya", "", ""],
    ["Diperiksa oleh (Reviewed by)", "", "IT / ERP Lead — Erajaya", "", ""],
    ["Diperiksa oleh (Reviewed by)", "", "Finance & Accounting — PT EBR", "", ""],
    ["Disetujui oleh (Approved by)", "", "Business Owner — Levi's Retail", "", ""],
    ["Disetujui oleh (Approved by)", "", "Manajemen PT SES", "", ""],
]


# ============================================================================
# 1. TECHNICAL SPECIFICATION DOCUMENT
# ============================================================================
def build_tsd():
    b = Builder()
    b.cover(
        "TECHNICAL SPECIFICATION DOCUMENT",
        "TSD — Implementasi Odoo untuk Bisnis Retail Levi's Indonesia",
        SUBTITLE,
        "TSD/ERAJAYA/LEVIS-ODOO/2026/002",
    )
    b.approval(APPROVAL_INTRO, APPROVAL_ROWS)

    # ---- 1. Pendahuluan ----
    b.h("1. Pendahuluan", 1)
    b.h("1.1 Tujuan Dokumen", 2)
    b.p(
        "Dokumen ini menjabarkan spesifikasi teknis implementasi Odoo untuk PT EBR "
        "(retail Levi's) sebagaimana kebutuhan fungsionalnya telah didefinisikan pada FSD "
        "FSD/ERAJAYA/LEVIS-ODOO/2026/001. Fokus utama dokumen adalah rancangan teknis modul "
        "custom_retail_import — adapter mirroring data Excel/CSV dan SFTP — termasuk model "
        "data, alur pemrosesan, mekanisme deduplikasi dan idempotensi, eksekusi asinkron, "
        "penanganan error, keamanan, dan deployment. Dokumen ditujukan bagi tim teknis "
        "(developer, DevOps) dan reviewer arsitektur."
    )
    b.h("1.2 Referensi", 2)
    b.bullets(
        [
            "FSD — FSD/ERAJAYA/LEVIS-ODOO/2026/001 v1.0.",
            "Source code modul: addons/ee_gap/custom_retail_import (Odoo 19, LGPL-3).",
            "Runbook onboarding data: scripts/tenants/levis/README.md (Track A).",
            "Diagram proses bisnis: docs/levis/business-process-flow.html.",
        ]
    )

    # ---- 2. Arsitektur ----
    b.h("2. Arsitektur Sistem", 1)
    b.h("2.1 Platform", 2)
    b.p(
        "Odoo 19 Community Edition berjalan di atas platform multi-tenant Erajaya dengan "
        "pola database-per-tenant; tenant Levi's menggunakan database tersendiri yang "
        'diprovision melalui Industry Pack "Retail / POS". Aplikasi berjalan dalam '
        "kontainer Docker di belakang reverse proxy, dengan PostgreSQL sebagai basis data "
        "dan filestore bersama untuk ir.attachment. Pemrosesan latar belakang menggunakan "
        "modul OCA queue_job dengan channel khusus root.retail_import."
    )
    b.h("2.2 Dependensi Teknis", 2)
    b.table(
        ["Komponen", "Versi/Sumber", "Kegunaan"],
        [
            [
                "custom_retail_import",
                "19.0.0.1.0 (addons/ee_gap)",
                "Adapter import retail: profil, wizard, executor, log, SFTP feed.",
            ],
            ["queue_job", "OCA", "Eksekusi asinkron file besar; channel root.retail_import."],
            ["openpyxl", "Python (image Odoo)", "Parser .xlsx streaming (read_only, data_only)."],
            ["paramiko", "Python (image Odoo)", "Klien SFTP untuk feed otomatis (Phase 2)."],
            ["Modul dependensi", "custom_core, product, stock, account", "Dependensi manifest custom_retail_import."],
        ],
        widths=[4.2, 4.2, 7.8],
    )
    b.p(
        "Library openpyxl dan paramiko ditambahkan pada odoo/requirements.txt dan tersedia "
        "setelah rebuild image. Perubahan kode Python pada modul memerlukan restart "
        "kontainer Odoo (reload ORM saja tidak cukup untuk method baru)."
    )
    b.h("2.3 Topologi Aliran Data", 2)
    b.p(
        "Phase 1: pengguna mengunduh report dari sistem POS principal, lalu mengunggahnya "
        "melalui wizard di Odoo. Phase 2: file di-drop ke server FTP/SFTP Erajaya; cron Odoo "
        "menarik file melalui koneksi SFTP keluar dari kontainer Odoo (jaringan odoo-net), "
        "tanpa volume Docker tambahan — file yang diunduh disimpan sebagai ir.attachment "
        "pada filestore bersama. Kedua jalur bermuara pada executor yang sama sehingga "
        "perilaku parsing, deduplikasi, dan audit identik."
    )

    # ---- 3. Custom Retail Import — detail ----
    b.h("3. Spesifikasi Teknis Modul custom_retail_import", 1)
    b.p(
        "Modul terdiri atas lima komponen: profil import (konfigurasi parsing per jenis "
        "file), wizard unggah (interaksi pengguna), executor (logika bisnis per jenis "
        "file), log (audit dan deduplikasi), serta feed SFTP (otomasi Phase 2). Desain "
        "memisahkan konfigurasi (data) dari kode: tenant retail baru cukup menambah/ubah "
        "record profil tanpa menulis kode."
    )

    b.h("3.1 Model Data", 2)
    b.h("3.1.1 retail.import.profile — Profil Parsing per Jenis File", 3)
    b.p(
        "Mendeklarasikan cara mem-parsing satu jenis file sumber. Mendukung 10 jenis file: "
        "x101 (Material Master), x20 (On-hand), x24 (Retail Sales Detail), x70d (Tender "
        "Detail), x70t (Tender Settlement), x31 (Discount Journal), x32p (Stock Movement), "
        "coa (Chart of Accounts), store_master, dan company (legal entity)."
    )
    b.table(
        ["Field", "Tipe", "Keterangan"],
        [
            ["code", "Char (unique per company)", "Identitas stabil profil, mis. 'levis_x101'."],
            ["file_type", "Selection (10 jenis)", "Menentukan executor yang dipakai."],
            [
                "namespace",
                "Char (default retail_import)",
                "Namespace ir.model.data untuk external ID (idempotensi); per tenant, mis. 'levis'.",
            ],
            ["file_format", "Selection xlsx/csv", "Format sumber."],
            ["sheet_name", "Char", "Sheet xlsx yang dibaca; kosong = sheet aktif/pertama."],
            [
                "data_start_row",
                "Integer (1-based)",
                "Baris pertama DATA; baris judul/header di atasnya dilewati. X101=3, Store Master=3, X24/X70D=2.",
            ],
            [
                "encoding / delimiter",
                "Selection / Char",
                "CSV: utf-8, utf-8-sig (default, BOM SSRS), latin-1; delimiter default ','.",
            ],
            [
                "column_map_json",
                "Text (JSON)",
                'Pemetaan nama field logis → indeks kolom 1-based (A=1), mis. {"product_code": 2, "sku": 10}.',
            ],
            ["date_format", "Char", "Format strptime; fallback toleran ke 5 format umum."],
            ["decimal_separator / thousand_separator", "Char", "Konfigurasi parsing angka."],
            [
                "fix_encoding",
                "Boolean (default True)",
                "Pulihkan karakter rusak U+FFFD → '®' pada semua sel string (kekhasan X101).",
            ],
            ["sample_file", "Binary", "Contoh file untuk dokumentasi profil."],
        ],
        widths=[4.0, 3.6, 8.6],
    )

    b.h("3.1.2 retail.import.log — Audit Trail", 3)
    b.table(
        ["Field", "Tipe", "Keterangan"],
        [
            ["name", "Char (computed)", "Format RIL/<id>/<filename>."],
            ["profile_id / file_type", "M2O / related", "Profil yang dipakai."],
            ["file_hash", "Char (indexed)", "SHA256 byte mentah file — kunci deduplikasi."],
            ["attachment_id", "M2O ir.attachment", "File sumber diarsipkan untuk audit dan pemrosesan ulang."],
            ["job_uuid", "Char (indexed)", "UUID queue_job bila diproses asinkron."],
            [
                "line_count / records_created / records_matched / records_skipped / error_count",
                "Integer",
                "Penghitung per tahap.",
            ],
            ["state", "Selection", "queued → running → imported | partial | failed (tracked)."],
            [
                "error_message / raw_payload",
                "Text",
                "Pesan kegagalan; ringkasan error per baris (dibatasi 200 baris pertama).",
            ],
            ["imported_at / imported_by_id", "Datetime / M2O res.users", "Jejak siapa dan kapan."],
        ],
        widths=[4.8, 3.2, 8.2],
    )

    b.h("3.1.3 retail.import.feed — Feed SFTP", 3)
    b.table(
        ["Field", "Tipe", "Keterangan"],
        [
            ["host / port / username", "Char / Integer / Char", "Tujuan SFTP (default port 22)."],
            [
                "auth_type",
                "Selection password/key",
                "Password dari ir.config_parameter (kunci pada password_param) atau private key "
                "RSA (path file dalam kontainer).",
            ],
            [
                "remote_dir / file_glob",
                "Char",
                "Direktori remote dan pola nama file (fnmatch), mis. 'X20_*.csv', 'X24DN_*.xlsx'.",
            ],
            ["run_async", "Boolean (default True)", "Serahkan tiap file ke queue_job."],
            [
                "last_run / last_status / last_message / files_imported",
                "Status",
                "Telemetri polling terakhir (ok/error/idle) untuk monitoring.",
            ],
        ],
        widths=[4.8, 3.4, 8.0],
    )

    b.h("3.1.4 retail.import.wizard dan retail.import.executor", 3)
    b.p(
        "retail.import.wizard adalah model transient untuk unggah file (Binary), dengan "
        "field force (bypass dedup) dan preview_text (hasil dry-run). retail.import.executor "
        "adalah AbstractModel tanpa tabel; method run(log) men-dispatch ke handler "
        "_load_<file_type> sesuai profil pada log."
    )

    b.h("3.2 Pipeline Parsing", 2)
    b.p(
        "Method read_records(file_b64, limit) pada profil mengubah byte file menjadi daftar "
        "dict berkunci nama field logis, dengan tahapan sebagai berikut."
    )
    b.bullets(
        [
            (
                "Pembacaan streaming. ",
                "File xlsx dibuka dengan openpyxl mode read_only + data_only sehingga konsumsi "
                "memori terkendali untuk file besar; CSV didekode sesuai encoding profil dengan "
                "errors=replace. Iterasi dimulai dari data_start_row.",
            ),
            (
                "Pemetaan kolom. ",
                "Setiap baris dipetakan melalui column_map_json (indeks kolom 1-based agar cocok "
                "dengan huruf kolom spreadsheet); indeks di luar rentang menghasilkan None tanpa "
                "error (toleran terhadap baris pendek).",
            ),
            ("Pembersihan string. ", "Trim whitespace dan perbaikan encoding U+FFFD → '®' bila fix_encoding aktif."),
            (
                "Baris kosong dilewati ",
                "dan dihitung terpisah (blank_rows); setiap record membawa nomor baris asli (_row) "
                "untuk pelaporan error yang presisi.",
            ),
            (
                "Koersi angka dan tanggal per-field. ",
                "_parse_amount membuang pemisah ribuan, menormalkan pemisah desimal, dan "
                "menafsirkan format kurung '(123)' sebagai negatif; _parse_date mencoba "
                "date_format profil lalu fallback toleran ke %Y-%m-%d, %d/%m/%Y, %d-%m-%Y, "
                "%m/%d/%Y, %Y/%m/%d.",
            ),
        ]
    )

    b.h("3.3 Deduplikasi dan Idempotensi (Detail)", 2)
    b.p(
        "Perlindungan duplikasi diterapkan pada dua lapis: lapis file (mencegah file yang "
        "sama diproses dua kali) dan lapis record (menjadikan pemrosesan ulang bersifat "
        "update, bukan penggandaan)."
    )
    b.h("3.3.1 Deduplikasi Lapis File — SHA256", 3)
    b.bullets(
        [
            (
                "Penghitungan hash. ",
                "Saat import, byte mentah file di-hash dengan SHA256 (retail.import.log."
                "compute_hash) dan disimpan pada field file_hash yang ter-index.",
            ),
            (
                "Pemeriksaan duplikat. ",
                "find_duplicate(file_hash) mencari log lain dengan hash sama yang berstatus "
                "imported, partial, atau running. Status failed sengaja dikecualikan sehingga file "
                "yang gagal dapat langsung diunggah ulang tanpa override.",
            ),
            (
                "Perilaku pada wizard (manual). ",
                "Bila duplikat ditemukan, import ditolak dengan pesan yang menyebutkan nomor log "
                'sebelumnya: "This exact file was already imported/queued (log #N)". Pengguna '
                "dapat mencentang opsi Force re-import untuk menimpa penolakan — tindakan sadar "
                "yang tercatat pada log baru.",
            ),
            (
                "Perilaku pada feed (otomatis). ",
                "File remote yang hash-nya sudah dikenal dilewati secara diam-diam (skip), sehingga "
                "polling berulang terhadap direktori yang sama tidak menghasilkan import ganda; "
                "file kosong (0 byte) juga dilewati.",
            ),
        ]
    )
    b.h("3.3.2 Idempotensi Lapis Record — External ID", 3)
    b.bullets(
        [
            (
                "Mekanisme. ",
                "Setiap record hasil import diberi external ID (ir.model.data) di bawah namespace "
                "profil (untuk Levi's: 'levis'), mis. kategori, atribut, template, varian produk "
                "pada X101 dan akun pada CoA. Saat re-import, executor mencari external ID terlebih "
                "dahulu: bila ada, record di-update; bila tidak, dibuat baru lalu external ID "
                "didaftarkan (noupdate=True).",
            ),
            (
                "Normalisasi kunci. ",
                "Helper _safe_xid menormalkan nilai kunci menjadi alfanumerik kapital (karakter "
                "lain diganti '_') sehingga external ID selalu valid dan stabil.",
            ),
            (
                "Isolasi antar tenant. ",
                "Namespace per tenant mencegah tabrakan external ID bila beberapa tenant retail "
                "menggunakan modul yang sama.",
            ),
        ]
    )
    b.h("3.3.3 Guard Tambahan — Saldo Awal One-shot", 3)
    b.p(
        "Import saldo awal stok (X20) bersifat one-shot: setelah berhasil, penanda "
        "ir.config_parameter (levis.x20_opening_stock_applied) dipasang sehingga eksekusi "
        "ulang tertolak. Guard ini melindungi dari penggandaan saldo awal yang berdampak "
        "langsung pada nilai persediaan."
    )

    b.h("3.4 Eksekusi Asinkron (queue_job)", 2)
    b.bullets(
        [
            (
                "Jenis file asinkron. ",
                "x101, x20, x24, x70d, dan x32p diproses melalui queue_job (channel "
                "root.retail_import) karena volumenya besar — X101 berisi ±159 ribu SKU dengan "
                "durasi proses ±30 menit, jauh melampaui timeout reverse proxy (gejala "
                "ERR_EMPTY_RESPONSE bila sinkron).",
            ),
            ("Pelacakan. ", "UUID job disimpan pada log (job_uuid) sehingga status dapat ditelusuri dari UI."),
            (
                "Degradasi anggun. ",
                "Bila queue_job tidak tersedia, eksekusi otomatis jatuh ke mode sinkron dengan "
                "peringatan di log server — fungsionalitas tetap berjalan.",
            ),
            (
                "Commit berbatch. ",
                "Executor menulis dengan batch 200 record per commit untuk menjaga ukuran transaksi "
                "dan memungkinkan progres parsial yang dapat diaudit.",
            ),
        ]
    )

    b.h("3.5 Cakupan Executor per Jenis File", 2)
    b.table(
        ["File Type", "Status", "Perilaku Teknis"],
        [
            [
                "x101 — Material Master",
                "FULL",
                "Agregasi SKU→best price by effective date; membentuk kategori (170), atribut "
                "Size/Inseam, product.template, dan product.product (varian) dengan external ID.",
            ],
            [
                "coa — Chart of Accounts",
                "FULL",
                "account.account dari file code/name/account_type; pencarian per company "
                "(company_ids M2M, perilaku Odoo 19); update nama bila akun sudah ada.",
            ],
            [
                "company — Legal Entity",
                "FULL",
                "res.company dan partner-nya dari sheet key/value (label Indonesia/Inggris: nama, "
                "NPWP→vat, alamat→street, telepon, email); mendukung sel gabungan 'Label : Value'.",
            ],
            [
                "x20 — On-hand",
                "FULL (one-shot)",
                "stock.quant saldo awal per warehouse toko per SKU; dikunci ir.config_parameter.",
            ],
            [
                "x24 — Retail Sales",
                "Phase-5 gated",
                "Membentuk pos.order (tanpa stock move). Prasyarat pos.config, session, dan payment "
                "method; gagal dengan pesan panduan bila prasyarat belum ada.",
            ],
            [
                "x70d — Tender Detail",
                "Phase-5 gated",
                "pos.payment, di-join ke order X24 berdasarkan kunci (store, date, register, transnum).",
            ],
            [
                "x70t / x31 / x32p / store_master",
                "STAGED",
                "Diparse dan dihitung, file diarsipkan; belum menulis model (menunggu keputusan "
                "Phase 3/5). Pembentukan warehouse dilakukan loader Track A (odoo shell).",
            ],
        ],
        widths=[4.0, 2.8, 9.6],
    )

    b.h("3.6 Penanganan Error dan State Machine", 2)
    b.p(
        "Log mengikuti state machine queued → running → imported/partial/failed. Executor "
        "meng-commit penanda running sebelum memproses agar status terlihat saat job "
        "berjalan. Exception tak terduga menyebabkan rollback transaksi data, lalu log "
        "ditandai failed beserta pesan; UserError diteruskan ke pengguna setelah rollback. "
        "Error per baris dikumpulkan sebagai pasangan (baris, pesan) dan diringkas maksimal "
        "200 baris pertama pada raw_payload; bila terdapat error namun sebagian record "
        "berhasil, status akhir adalah partial. Status failed tidak memblokir unggah ulang "
        "file yang sama (lihat 3.3.1)."
    )

    b.h("3.7 Profil yang Di-seed untuk Levi's", 2)
    b.p(
        "Enam profil di-seed sebagai data (noupdate=1) dengan namespace 'levis'; penyesuaian "
        "per tenant dilakukan dengan mengubah record, bukan kode."
    )
    b.table(
        ["Profil (code)", "Format", "Data Mulai Baris", "Kolom Kunci (column_map_json)"],
        [
            [
                "levis_x101",
                "xlsx",
                "3",
                "product_code(2), description(3), brand(4), category(5), class(6), subclass(7), "
                "sku(10), size(11), inseam(12), gtin(13), retail_price(15), price_eff(16).",
            ],
            ["levis_coa", "xlsx", "3", "code(1), account_name(2), account_type(3)."],
            ["levis_company", "xlsx", "4", "field(2), value(3) — layout key/value."],
            [
                "levis_x20",
                "csv (utf-8-sig)",
                "2",
                "store_code(16), sap_store_code(17), store_name(18), item_id(19), ean(23), onhand_qty(29).",
            ],
            [
                "levis_x24",
                "xlsx",
                "2",
                "store(1-3), trans_date(4), register(5), transnum(6), item(14-22), harga/qty/"
                "discount/amount(23-27), tax_rate(28), tax_amount(29), total(30).",
            ],
            [
                "levis_x70d",
                "xlsx",
                "2",
                "store(1-3), trans_date(4), register(5), transnum(6), tender_type(9), "
                "tender_amount(10), auth(11), voucher(12).",
            ],
        ],
        widths=[3.0, 2.6, 2.4, 8.4],
    )

    b.h("3.8 Keamanan dan Hak Akses", 2)
    b.bullets(
        [
            (
                "Grup akses. ",
                "Dua grup: Retail Import / User (menjalankan wizard, melihat log) dan Retail Import "
                "/ Manager (mengelola profil dan feed; mewarisi User). Hak akses model diatur via "
                "ir.model.access.csv.",
            ),
            (
                "Kredensial SFTP. ",
                "Password disimpan pada ir.config_parameter dengan kunci yang dirujuk field "
                "password_param. Platform belum menyediakan enkripsi at-rest untuk config "
                "parameter, sehingga nilainya tersimpan mentah — akses baca parameter harus "
                "dibatasi dan autentikasi berbasis private key lebih disarankan.",
            ),
            (
                "Arsip file. ",
                "File sumber tersimpan sebagai ir.attachment terkait log; retensi mengikuti "
                "kebijakan filestore tenant.",
            ),
        ]
    )

    b.h("3.9 Penjadwalan (ir.cron)", 2)
    b.p(
        'Cron "Retail Import: poll SFTP feeds" memanggil _cron_poll_feeds() yang mem-poll '
        "seluruh feed aktif. Default nonaktif dan berkadens harian — sesuai steady state "
        "(penjualan X24 hari sebelumnya dan refresh on-hand X20); diaktifkan saat feed "
        "pertama dikonfigurasi pada Phase 2. Setiap feed juga dapat dipicu manual (Poll Now) "
        "dan diuji koneksinya (Test Connection menampilkan jumlah file yang cocok dengan "
        "glob)."
    )

    # ---- 4. Deployment ----
    b.h("4. Deployment dan Operasional", 1)
    b.h("4.1 Langkah Deployment", 2)
    b.bullets(
        [
            "Tambahkan openpyxl dan paramiko pada odoo/requirements.txt, lalu rebuild image Odoo.",
            "Install modul custom_retail_import pada database tenant levis; restart kontainer "
            "Odoo setelah perubahan kode Python.",
            "Verifikasi profil ter-seed (menu Retail Import ▸ Configuration ▸ Profiles) dan "
            "sesuaikan column map bila struktur file pelanggan berubah.",
            "Phase 2: buat record SFTP Feed per jenis file (host, kredensial, remote_dir, "
            "file_glob), uji dengan Test Connection, lalu aktifkan cron poller.",
        ]
    )
    b.h("4.2 Monitoring dan Verifikasi", 2)
    b.p(
        "Pemantauan operasional menggunakan: (1) daftar Retail Import Log — status, "
        "penghitung record, ringkasan error per baris; (2) status feed — last_run, "
        "last_status, last_message, files_imported; (3) antrian queue_job pada channel "
        "root.retail_import. Verifikasi pasca-import master data dilakukan dengan "
        "penghitungan record (product.template ±14.885; product.product ber-default_code "
        "±159.658; account.account = jumlah baris CoA; stock.warehouse = 24; total "
        "stock.quant = total X20)."
    )
    b.h("4.3 Batasan yang Diketahui", 2)
    b.bullets(
        [
            "Executor x70t, x31, x32p, dan store_master masih staged (parse dan hitung saja) — "
            "penulisan model menunggu keputusan Phase 3/5.",
            "Posting x24/x70d memerlukan keputusan Phase 5 (representasi POS, kedalaman "
            "histori, tax mapping) serta prasyarat pos.config/session/payment method.",
            "Kredensial SFTP tersimpan tanpa enkripsi at-rest pada ir.config_parameter "
            "(mitigasi: pembatasan akses dan preferensi key-based auth).",
            "File transaksi yang diterima saat ini sampel satu toko; crosswalk kode toko untuk "
            "23 toko menunggu data pelanggan.",
        ]
    )
    b.save("TSD - Levis Retail Odoo (SES x EBR) v1.0.docx")


# ============================================================================
# 2. PROJECT CHARTER
# ============================================================================
def build_charter():
    b = Builder()
    b.cover(
        "PROJECT CHARTER",
        "Project Charter — Implementasi Odoo untuk Bisnis Retail Levi's Indonesia",
        SUBTITLE,
        "PC/ERAJAYA/LEVIS-ODOO/2026/003",
    )
    b.approval(
        APPROVAL_INTRO,
        [
            ["Disusun oleh (Prepared by)", "", "Project Manager — Erajaya", "", ""],
            ["Diperiksa oleh (Reviewed by)", "", "IT / ERP Lead — Erajaya", "", ""],
            ["Diperiksa oleh (Reviewed by)", "", "Finance & Accounting — PT EBR", "", ""],
            ["Disetujui oleh (Project Sponsor)", "", "Business Owner — Levi's Retail", "", ""],
            ["Disetujui oleh (Approved by)", "", "Manajemen PT SES", "", ""],
        ],
    )

    b.h("1. Latar Belakang dan Justifikasi Bisnis", 1)
    b.p(
        "PT EBR sebagai entitas retail Levi's membutuhkan sistem pembukuan dan pelaporan "
        "keuangan yang terpisah dari sistem transaksional toko (POS bawaan principal) dan "
        "dari sistem distribusi PT SES (SAP). Tanpa sistem pencatatan tersendiri, "
        "pembukuan EBR bergantung pada rekap manual yang lambat, rawan salah, dan sulit "
        "diaudit. Implementasi Odoo memberikan: (1) pencatatan pembelian EBR kepada SES "
        "yang tertib dengan approval berjenjang; (2) replikasi otomatis data operasional "
        "toko — GR, penjualan dan tender, mutasi persediaan — ke dalam pembukuan; (3) "
        "laporan keuangan dan rekonsiliasi yang tepat waktu; serta (4) fondasi yang "
        "reusable bagi tenant retail Erajaya berikutnya melalui modul adapter yang "
        "dikonfigurasi sebagai data, bukan kode."
    )

    b.h("2. Tujuan dan Kriteria Keberhasilan", 1)
    b.table(
        ["No", "Tujuan", "Kriteria Keberhasilan (Terukur)"],
        [
            [
                "1",
                "Pembukuan EBR berjalan di Odoo",
                "PO EBR→SES, vendor bill, dan pembayaran tercatat di Odoo; tutup buku bulanan "
                "dapat dihasilkan dari Odoo.",
            ],
            [
                "2",
                "Master data lengkap dan valid",
                "±14.885 artikel / ±159.658 SKU termuat; CoA sesuai file EBR; 24 warehouse "
                "terbentuk; verifikasi count lolos.",
            ],
            [
                "3",
                "Mirroring operasional toko berjalan",
                "GR, penjualan, tender, dan mutasi stok harian termuat H+1 dengan log audit "
                "lengkap dan nol duplikasi (dedup SHA256 aktif).",
            ],
            [
                "4",
                "Otomasi Phase 2",
                "Feed SFTP aktif untuk seluruh jenis file rutin; tanpa unggah manual pada steady state.",
            ],
            [
                "5",
                "Rekonsiliasi tender",
                "Settlement tender (X70T) terekonsiliasi terhadap penerimaan bank per acquirer setiap periode.",
            ],
        ],
        widths=[1.0, 4.6, 10.6],
    )

    b.h("3. Ruang Lingkup", 1)
    b.p(
        "In-scope: implementasi Odoo PT EBR (procurement ke SES, akuntansi, persediaan, "
        "pelaporan); modul custom_retail_import beserta profil Levi's; penyiapan master "
        "data dan saldo awal; integrasi file Phase 1 dan SFTP Phase 2. Out-of-scope: "
        "penggantian POS toko, perubahan proses SAP SES, integrasi API real-time, dan "
        "perubahan proses pada Levi's Principal. Rincian lengkap mengacu pada FSD Bab 1.3."
    )

    b.h("4. Deliverables", 1)
    b.bullets(
        [
            "Tenant Odoo levis terprovision dengan Industry Pack Retail / POS.",
            "Modul custom_retail_import terpasang dengan 6 profil Levi's ter-seed.",
            "Master data termuat: company, CoA, produk X101, 24 warehouse, saldo awal X20.",
            "Konfigurasi keuangan: bank account dan journals, pajak/Coretax, vendor SES.",
            "Feed SFTP terkonfigurasi dan teruji (Phase 2).",
            "Dokumen: FSD, TSD, Project Charter, Business Workflow, runbook onboarding, materi pelatihan pengguna.",
        ]
    )

    b.h("5. Milestone dan Tahapan", 1)
    b.p("Tanggal target diisi saat penjadwalan bersama; urutan tahapan dan dependensinya adalah sebagai berikut.")
    b.table(
        ["Milestone", "Tahapan", "Keluaran Utama", "Dependensi"],
        [
            ["M0", "Provisioning tenant (Phase 0)", "Database levis + pack retail terpasang.", "—"],
            [
                "M1",
                "Konfigurasi dasar & master data (Phase 1)",
                "Company, CoA, bank/journals, pajak, vendor SES, produk X101, 24 warehouse.",
                "M0; file X101 dan CoA (sudah lengkap).",
            ],
            [
                "M2",
                "Saldo awal stok (Phase 2 data)",
                "X20 termuat one-shot pada tanggal cut-over.",
                "M1; ekstrak X20 seluruh toko dari pelanggan.",
            ],
            [
                "M3",
                "Go-live procurement",
                "PO EBR→SES, GR mirror, vendor bill, pembayaran berjalan di produksi.",
                "M1; kode toko lengkap untuk crosswalk.",
            ],
            [
                "M4",
                "Keputusan dan go-live penjualan (Phase 5)",
                "Posting X24/X70D diaktifkan setelah keputusan representasi POS, histori, dan tax mapping.",
                "M3; workshop Finance EBR dan pajak.",
            ],
            [
                "M5",
                "Otomasi SFTP (Phase 6)",
                "Feed SFTP aktif, cron poller berjalan, unggah manual dihentikan.",
                "M3; akses server FTP/SFTP Erajaya.",
            ],
            [
                "M6",
                "Stabilisasi dan serah terima",
                "Rekonsiliasi 1 siklus bulanan penuh; handover ke operasional.",
                "M4, M5.",
            ],
        ],
        widths=[1.6, 4.6, 6.0, 4.0],
    )

    b.h("6. Organisasi Proyek", 1)
    b.table(
        ["Peran", "Pihak", "Tanggung Jawab Utama"],
        [
            [
                "Project Sponsor",
                "Business Owner Levi's Retail (Erajaya)",
                "Keputusan bisnis, pendanaan, eskalasi akhir.",
            ],
            ["Project Manager", "Erajaya", "Perencanaan, koordinasi lintas pihak, pelaporan status, manajemen risiko."],
            ["Functional Consultant", "Erajaya", "Konfigurasi Odoo, mapping proses, UAT, pelatihan."],
            [
                "Technical Lead / Developer",
                "Erajaya",
                "Modul custom_retail_import, profil import, feed SFTP, deployment.",
            ],
            [
                "Finance & Accounting",
                "PT EBR",
                "CoA, kebijakan pembukuan, keputusan Phase 5, rekonsiliasi, UAT keuangan.",
            ],
            ["IT / Master Data", "PT SES", "Penyediaan ekstrak SAP/POS, kode toko, koordinasi dengan principal."],
            ["Principal Liaison", "Levi's Principal", "Akses report POS, jadwal drop file, perubahan struktur report."],
        ],
        widths=[3.6, 4.4, 8.2],
    )

    b.h("7. Risiko Utama dan Mitigasi", 1)
    b.table(
        ["No", "Risiko", "Dampak", "Mitigasi"],
        [
            [
                "1",
                "Data transaksi baru tersedia untuk 1 dari 24 toko (sampel).",
                "Go-live penjualan tertunda.",
                "Eskalasi permintaan ekstrak penuh / aktivasi feed SFTP lebih awal; go-live bertahap per toko.",
            ],
            [
                "2",
                "Kode toko SAP belum tersedia untuk 23 toko.",
                "Crosswalk file X ke warehouse tidak lengkap.",
                "Daftar kebutuhan data formal ke SES/principal; loader store bersifat idempoten "
                "sehingga dapat diulang saat kode tiba.",
            ],
            [
                "3",
                "Keputusan Phase 5 (representasi POS, histori, tax mapping) tertunda.",
                "Posting penjualan dan tender tertahan gating.",
                "Workshop keputusan dijadwalkan sebelum M4 dengan opsi default yang direkomendasikan.",
            ],
            [
                "4",
                "Struktur report principal berubah tanpa pemberitahuan.",
                "Parsing gagal atau salah kolom.",
                "Profil import berbasis data (ubah mapping tanpa kode); dry-run preview; notifikasi kegagalan parse.",
            ],
            [
                "5",
                "Kredensial SFTP tersimpan tanpa enkripsi at-rest.",
                "Risiko kebocoran kredensial.",
                "Preferensi autentikasi private key; pembatasan akses ir.config_parameter; roadmap enkripsi platform.",
            ],
            [
                "6",
                "Volume X101 besar (±159 ribu SKU) membebani import.",
                "Timeout / beban server.",
                "Eksekusi asinkron queue_job, parsing streaming, commit berbatch (teruji ±30 menit).",
            ],
        ],
        widths=[1.0, 4.6, 3.6, 7.0],
    )

    b.h("8. Asumsi dan Batasan", 1)
    b.bullets(
        [
            "Odoo EBR adalah sistem pembukuan; POS principal dan SAP SES tetap menjadi sumber transaksi.",
            "Format report principal stabil; perubahan dikelola melalui penyesuaian profil.",
            "Akses FTP/SFTP Erajaya tersedia sebelum M5.",
            "Personel kunci EBR/SES tersedia untuk UAT dan workshop keputusan Phase 5.",
        ]
    )

    b.h("9. Manajemen Perubahan dan Pelaporan", 1)
    b.p(
        "Perubahan ruang lingkup setelah charter disetujui diajukan melalui change request "
        "kepada Project Sponsor dengan analisis dampak (waktu, biaya, risiko). Status "
        "proyek dilaporkan berkala oleh Project Manager mencakup kemajuan milestone, isu "
        "terbuka (termasuk tiga isu data pada FSD Bab 8.2), dan keputusan yang dibutuhkan. "
        "Persetujuan atas dokumen ini memberikan mandat kepada Project Manager untuk "
        "memobilisasi tim dan memulai eksekusi sesuai tahapan pada Bab 5."
    )
    b.save("Project Charter - Levis Retail Odoo v1.0.docx")


# ============================================================================
# 3. BUSINESS WORKFLOW
# ============================================================================
def build_workflow():
    b = Builder()
    b.cover(
        "BUSINESS WORKFLOW",
        "Business Workflow — Proses Bisnis Retail Levi's pada Odoo",
        SUBTITLE,
        "BW/ERAJAYA/LEVIS-ODOO/2026/004",
    )
    b.approval(APPROVAL_INTRO, APPROVAL_ROWS)

    b.h("1. Pendahuluan", 1)
    b.p(
        "Dokumen ini menjabarkan alur kerja bisnis (business workflow) operasional retail "
        "Levi's lintas empat aktor — Levi's Principal, PT SES (SAP), toko (POS principal), "
        "dan PT EBR (Odoo) — pada tingkat langkah operasional harian, termasuk siklus "
        "harian/bulanan, pembagian peran, dan penanganan pengecualian. Dokumen melengkapi "
        "FSD (kebutuhan fungsional) dan TSD (rancangan teknis); diagram visual tersedia "
        "pada docs/levis/business-process-flow.html."
    )

    b.h("2. Peran dan Tanggung Jawab (RACI Ringkas)", 1)
    b.table(
        ["Aktivitas", "Toko", "SES (SAP)", "EBR Finance (Odoo)", "Erajaya IT"],
        [
            ["PO EBR → SES", "—", "I (terima PO)", "R/A (buat & approve)", "C"],
            ["Pengiriman barang ke toko", "I", "R/A", "I", "—"],
            ["GR di toko (POS principal)", "R/A", "I", "I", "—"],
            ["Ekspor report harian POS", "R", "C", "I", "C"],
            ["Upload/feed file ke Odoo", "—", "—", "R (Phase 1)", "R/A (Phase 2)"],
            ["Validasi import & posting", "—", "—", "R/A", "C"],
            ["Vendor bill & pembayaran SES", "—", "I", "R/A", "—"],
            ["Rekonsiliasi tender vs bank", "—", "—", "R/A", "C"],
            ["Klaim promosi (X31)", "—", "C", "R/A", "—"],
        ],
        widths=[5.2, 2.4, 2.6, 3.4, 2.6],
    )
    b.p("R = Responsible (pelaksana), A = Accountable (penanggung jawab), C = Consulted, I = Informed.", italic=True)

    b.h("3. Workflow Utama", 1)

    b.h("3.1 WF-A — Pengadaan SES kepada Principal (Konteks)", 2)
    b.p(
        "Alur ini sepenuhnya berjalan pada SAP SES dan menjadi konteks hulu: SES "
        "menerbitkan PO kepada principal, menerima barang dengan GR di DC, dan "
        "menyelesaikan pembayaran melalui three-way match. Ketersediaan stok DC SES "
        "menjadi prasyarat pemenuhan PO dari EBR pada WF-B."
    )

    b.h("3.2 WF-B — Pembelian EBR kepada SES sampai GR Toko", 2)
    b.table(
        ["Langkah", "Aktor", "Aktivitas", "Input", "Output"],
        [
            [
                "B1",
                "EBR Purchasing (Odoo)",
                "Membuat PO ke vendor PT SES; approval berjenjang via approval engine sesuai matriks otorisasi.",
                "Rencana kebutuhan toko",
                "PO approved (Odoo)",
            ],
            [
                "B2",
                "SES (SAP)",
                "Menerima PO, membentuk SO dan Delivery, mengirim barang DC → toko.",
                "PO EBR",
                "Surat jalan, barang tiba di toko",
            ],
            [
                "B3",
                "Toko (POS principal)",
                "Memeriksa fisik barang dan melakukan GR pada POS; selisih dicatat sebagai discrepancy ke SES.",
                "Surat jalan + fisik barang",
                "GR POS; stok toko bertambah",
            ],
            [
                "B4",
                "EBR Finance (Odoo)",
                "Menerima mirror GR (file X32P/GR) → memvalidasi receipt atas PO; stok dan akrual hutang tercatat.",
                "File mirror GR",
                "Receipt tervalidasi di Odoo",
            ],
            [
                "B5",
                "EBR Finance (Odoo)",
                "Mencatat vendor bill SES, mencocokkan PO–GR–Bill, menjadwalkan dan membayar.",
                "Invoice SES",
                "Bill posted; pembayaran terlaksana",
            ],
        ],
        widths=[1.4, 3.4, 5.6, 2.9, 2.9],
    )

    b.h("3.3 WF-C — Siklus Harian Toko (EOD) dan Mirroring", 2)
    b.table(
        ["Langkah", "Aktor", "Aktivitas", "Waktu"],
        [
            [
                "C1",
                "Toko",
                "Transaksi sepanjang hari pada POS principal: penjualan, retur, "
                "pembayaran multi-tender; GR/transfer bila ada kiriman.",
                "Sepanjang hari",
            ],
            [
                "C2",
                "Toko / Principal",
                "Tutup kasir; report harian terbentuk di sistem principal (X24DN, X70D/X70T, X32P, X20).",
                "EOD",
            ],
            [
                "C3",
                "Ops (Phase 1) / Sistem (Phase 2)",
                "Phase 1: ekspor dan unggah file via wizard Odoo (dengan preview dry-run). "
                "Phase 2: file di-drop ke SFTP Erajaya dan ditarik cron Odoo otomatis.",
                "EOD / H+1 pagi",
            ],
            [
                "C4",
                "Sistem (Odoo)",
                "Dedup SHA256 → parsing per profil → posting: penjualan + HPP (X24), payment "
                "register (X70D), stock move (X32P); file diarsipkan, log tercatat.",
                "Otomatis setelah C3",
            ],
            [
                "C5",
                "EBR Finance",
                "Memeriksa log import (status, jumlah baris, error); menindaklanjuti exception; "
                "rekonsiliasi settlement tender terhadap mutasi bank per acquirer.",
                "H+1",
            ],
        ],
        widths=[1.4, 3.8, 8.2, 2.8],
    )

    b.h("3.4 WF-D — Siklus Bulanan", 2)
    b.bullets(
        [
            (
                "Rekonsiliasi persediaan. ",
                "Posisi on-hand Odoo dibandingkan dengan X20; selisih ditindaklanjuti sebagai "
                "inventory adjustment dengan persetujuan.",
            ),
            (
                "Rekonsiliasi tender. ",
                "X70T (settlement) dicocokkan terhadap penerimaan bank; selisih fee acquirer "
                "dibukukan sesuai kebijakan.",
            ),
            ("Klaim promosi. ", "X31 (discount journal) direkap sebagai dasar klaim promosi kepada principal/SES."),
            (
                "Tutup buku. ",
                "Jurnal penjualan, HPP, persediaan, AP, dan bank ditutup; laporan keuangan EBR "
                "diterbitkan dari Odoo (custom_accounting_full).",
            ),
        ]
    )

    b.h("4. Penanganan Pengecualian (Exception Workflow)", 1)
    b.table(
        ["Kode", "Pengecualian", "Penanganan", "Peran"],
        [
            [
                "E1",
                "File duplikat diunggah (hash SHA256 sama dengan import sukses/berjalan sebelumnya).",
                "Sistem menolak dengan menyebutkan nomor log sebelumnya. Bila pengulangan memang "
                "disengaja (mis. file dikoreksi di tempat dengan isi identik), Manager dapat "
                "mencentang Force re-import; tindakan tercatat pada log baru.",
                "Retail Import Manager",
            ],
            [
                "E2",
                "File gagal parse (struktur kolom berubah, encoding salah).",
                "Log berstatus failed dengan pesan error; gunakan preview dry-run untuk diagnosis; "
                "perbaiki column map pada profil (data, tanpa deploy kode); unggah ulang — status "
                "failed tidak memblokir dedup.",
                "Erajaya IT + EBR Finance",
            ],
            [
                "E3",
                "Import partial (sebagian baris error).",
                "raw_payload memuat ringkasan error per baris (maks. 200); baris bermasalah "
                "dikoreksi di sumber lalu file revisi diimpor — idempotensi external ID menjadikan "
                "baris yang sudah masuk ter-update, bukan terduplikasi.",
                "EBR Finance",
            ],
            [
                "E4",
                "Feed SFTP error (koneksi/kredensial/direktori).",
                "Status feed menjadi error dengan pesan; uji dengan Test Connection; perbaiki "
                "kredensial pada ir.config_parameter atau private key; Poll Now untuk mengulang.",
                "Erajaya IT",
            ],
            [
                "E5",
                "Selisih GR vs PO (qty kurang/lebih, barang rusak).",
                "Receipt divalidasi sesuai qty aktual; selisih dieskalasi ke SES sebagai "
                "discrepancy untuk debit note/pengiriman susulan sebelum bill dicocokkan.",
                "Toko + EBR Finance + SES",
            ],
            [
                "E6",
                "Selisih rekonsiliasi tender vs bank.",
                "Identifikasi per acquirer per hari; selisih fee dibukukan; selisih abnormal "
                "dieskalasi ke acquirer/principal.",
                "EBR Finance",
            ],
            [
                "E7",
                "Saldo awal X20 dijalankan ulang.",
                "Sistem menolak otomatis (guard one-shot ir.config_parameter); pengulangan sah "
                "memerlukan keputusan formal dan reset penanda oleh IT.",
                "Erajaya IT",
            ],
        ],
        widths=[1.2, 4.4, 7.6, 3.0],
    )

    b.h("5. Kontrol dan Audit", 1)
    b.bullets(
        [
            (
                "Jejak audit penuh. ",
                "Setiap import tercatat: siapa, kapan, file apa (hash SHA256), berapa baris, "
                "berapa record dibuat/di-update/dilewati, dan error apa; file sumber tersimpan "
                "sebagai lampiran.",
            ),
            (
                "Pemisahan tugas. ",
                "Grup Retail Import User menjalankan import; perubahan profil/feed dan force "
                "re-import berada pada grup Manager.",
            ),
            ("Approval berjenjang. ", "PO EBR melewati approval engine sebelum dikirim ke SES."),
            (
                "Pengendalian duplikasi. ",
                "Dedup file (SHA256) dan idempotensi record (external ID) mencegah penggandaan "
                "transaksi dan master data; saldo awal dilindungi guard one-shot.",
            ),
            (
                "Ketertelusuran. ",
                "Nomor baris sumber (_row) disertakan pada pesan error sehingga koreksi di file "
                "sumber dapat dilakukan presisi.",
            ),
        ]
    )

    b.h("6. Indikator Operasional (KPI Workflow)", 1)
    b.table(
        ["KPI", "Target", "Sumber Pengukuran"],
        [
            [
                "Ketepatan waktu mirroring harian",
                "Data EOD termuat ≤ H+1 jam 10.00",
                "retail.import.log (imported_at vs tanggal transaksi)",
            ],
            ["Tingkat keberhasilan import", "≥ 99% log berstatus imported (bukan failed)", "Statistik log per bulan"],
            ["Duplikasi transaksi", "0 kejadian", "Audit dedup + rekonsiliasi"],
            ["Selisih rekonsiliasi tender", "≤ ambang fee acquirer yang disepakati", "Rekonsiliasi bulanan"],
            ["Penyelesaian exception", "E1–E7 selesai ≤ 2 hari kerja", "Tindak lanjut log/feed"],
        ],
        widths=[5.4, 5.4, 5.4],
    )

    b.p(
        "Dengan disetujuinya dokumen ini, alur kerja pada Bab 3–4 menjadi acuan operasional "
        "harian seluruh pihak, dan setiap penyimpangan ditangani melalui jalur pengecualian "
        "yang telah ditetapkan."
    )
    b.save("Business Workflow - Levis Retail Odoo v1.0.docx")


if __name__ == "__main__":
    build_tsd()
    build_charter()
    build_workflow()
