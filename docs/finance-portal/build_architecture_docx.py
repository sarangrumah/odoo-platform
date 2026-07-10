# -*- coding: utf-8 -*-
"""Generate Finance Portal Architecture document as DOCX (v2.0).

Key diagrams are rendered as embedded PNG images via matplotlib.
Usage:  python docs/finance-portal/build_architecture_docx.py
"""

import io
import os

import matplotlib

matplotlib.use("Agg")
import matplotlib.patches as mpatches
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))

# ── Colour palette ──────────────────────────────────────────────────────────
C_INK = RGBColor(0x1A, 0x23, 0x32)
C_ACCENT = RGBColor(0x71, 0x4B, 0x67)
HEAD_BG = "1A2332"
TOTAL_BG = "EDE7EC"
LGREY = "F5F5F5"

# matplotlib palette
M_ODOO = "#1168BD"
M_BRIDGE = "#2D6A4F"
M_KAFKA = "#D4860B"
M_SAP = "#5D6D7E"
M_HRIS = "#717D7E"
M_KC = "#884EA0"
M_PERSON = "#08427B"
M_WARN = "#C0392B"
M_HOLDING = "#8E44AD"
M_PT = "#1A5276"
M_SHARED = "#1E8449"
M_NEW = "#1A8754"
M_WHITE = "#FFFFFF"
M_BG = "#F8F9FA"


# ── matplotlib helpers ───────────────────────────────────────────────────────


def _fig(w=14, h=8, dpi=130):
    fig, ax = plt.subplots(figsize=(w, h))
    ax.set_xlim(0, w)
    ax.set_ylim(0, h)
    ax.axis("off")
    fig.patch.set_facecolor(M_WHITE)
    return fig, ax


def _box(ax, x, y, w, h, label, sublabel="", color=M_ODOO, textcolor=M_WHITE, fontsize=9, radius=0.25, zorder=3):
    box = FancyBboxPatch(
        (x, y),
        w,
        h,
        boxstyle=f"round,pad=0,rounding_size={radius}",
        facecolor=color,
        edgecolor=M_WHITE,
        linewidth=1.5,
        zorder=zorder,
    )
    ax.add_patch(box)
    cy = y + h / 2
    if sublabel:
        ax.text(
            x + w / 2,
            cy + 0.15,
            label,
            ha="center",
            va="center",
            fontsize=fontsize,
            fontweight="bold",
            color=textcolor,
            zorder=zorder + 1,
        )
        ax.text(
            x + w / 2,
            cy - 0.22,
            sublabel,
            ha="center",
            va="center",
            fontsize=fontsize - 1.5,
            color=textcolor,
            zorder=zorder + 1,
            linespacing=1.3,
        )
    else:
        ax.text(
            x + w / 2,
            cy,
            label,
            ha="center",
            va="center",
            fontsize=fontsize,
            fontweight="bold",
            color=textcolor,
            zorder=zorder + 1,
        )
    return box


def _arrow(ax, x1, y1, x2, y2, label="", color="#444", lw=1.5, zorder=2):
    ax.annotate(
        "",
        xy=(x2, y2),
        xytext=(x1, y1),
        arrowprops=dict(arrowstyle="-|>", color=color, lw=lw, mutation_scale=16),
        zorder=zorder,
    )
    if label:
        mx, my = (x1 + x2) / 2, (y1 + y2) / 2
        ax.text(
            mx,
            my + 0.15,
            label,
            ha="center",
            va="bottom",
            fontsize=7.5,
            color=color,
            zorder=zorder + 1,
            bbox=dict(boxstyle="round,pad=0.2", facecolor=M_WHITE, edgecolor="none", alpha=0.8),
        )


def _legend(ax, items, x=0.3, y=0.35, fs=8):
    for i, (color, label) in enumerate(items):
        px = x + i * 2.4
        p = mpatches.Patch(color=color, label=label)
        ax.add_patch(
            FancyBboxPatch(
                (px, y), 0.35, 0.25, boxstyle="round,pad=0,rounding_size=0.05", facecolor=color, edgecolor="none"
            )
        )
        ax.text(px + 0.45, y + 0.12, label, va="center", fontsize=fs, color="#333")


def _png(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=130, bbox_inches="tight", facecolor=M_WHITE)
    plt.close(fig)
    buf.seek(0)
    return buf


# ── Diagram generators ───────────────────────────────────────────────────────


def diagram_system_context():
    fig, ax = _fig(16, 10)
    ax.set_facecolor(M_BG)

    # Title
    ax.text(
        8,
        9.6,
        "System Context — Finance Portal",
        ha="center",
        fontsize=14,
        fontweight="bold",
        color=C_INK.rgb if hasattr(C_INK, "rgb") else "#1A2332",
    )

    # ── Users (top left) ──
    user_data = [
        (0.4, 7.8, "Karyawan\n(Employee)", M_PERSON),
        (0.4, 6.6, "Vendor", M_PERSON),
        (0.4, 5.4, "Tim Finance & Tax", M_PERSON),
        (0.4, 4.2, "Finance Group\n(Holding)", M_HOLDING),
    ]
    for x, y, lbl, col in user_data:
        _box(ax, x, y, 2.0, 0.9, lbl, color=col, fontsize=8)

    # ── Finance Portal DB (center) ──
    _box(ax, 3.2, 3.5, 5.6, 5.2, "", color="#EAF2FB", textcolor="#1A2332", zorder=1)
    ax.text(
        6.0,
        8.5,
        "Finance Portal DB  (BARU, TERPISAH)",
        ha="center",
        fontsize=9,
        fontweight="bold",
        color=M_NEW,
        bbox=dict(boxstyle="round,pad=0.3", facecolor="#D5F5E3", edgecolor=M_NEW),
    )

    module_data = [
        (3.4, 7.5, "custom_finance_portal", "Dokumen + Approval + Vendor Portal"),
        (3.4, 6.5, "custom_finance_budget", "Budget & PR Validation"),
        (3.4, 5.5, "custom_finance_portal_sap", "Adapter + Webhook + Sync Log"),
        (3.4, 4.5, "custom_finance_portal_sso", "Keycloak + Vendor Auth"),
    ]
    for x, y, name, sub in module_data:
        _box(ax, x, y, 5.2, 0.75, name, sub, color=M_ODOO, fontsize=8)

    # ── Existing DB (bottom left of portal) ──
    _box(ax, 3.2, 3.0, 2.4, 0.38, "Existing Odoo DB  (tidak disentuh)", color=M_WARN, fontsize=7.5)

    # ── Bridge ──
    _box(
        ax,
        9.6,
        5.5,
        2.8,
        2.2,
        "SAP Bridge",
        "FastAPI · Kafka Consumer/Producer\nRedis Cache",
        color=M_BRIDGE,
        fontsize=8.5,
    )

    # ── External ──
    ext = [
        (13.2, 8.2, "Keycloak\nOIDC IdP", M_KC),
        (13.2, 6.5, "Kafka Bus", M_KAFKA),
        (13.2, 4.8, "SAP S/4HANA\nSystem of Record", M_SAP),
        (13.2, 3.1, "HRIS\nTravel + Employee", M_HRIS),
    ]
    for x, y, lbl, col in ext:
        _box(ax, x, y, 2.5, 0.9, lbl, color=col, fontsize=8)

    # ── Arrows: Users → Portal ──
    for _, y, _, _ in user_data:
        _arrow(ax, 2.4, y + 0.45, 3.2, 6.0, color=M_PERSON, lw=1.2)

    # ── Portal → Keycloak ──
    _arrow(ax, 8.8, 4.7, 13.2, 8.6, "OIDC", color=M_KC, lw=1.5)

    # ── Portal → Bridge ──
    _arrow(ax, 8.8, 6.5, 9.6, 6.5, "HMAC REST\n(internal)", color=M_BRIDGE, lw=2)

    # ── Bridge → Kafka ──
    _arrow(ax, 12.4, 6.5, 13.2, 6.9, "produce/consume", color=M_KAFKA, lw=1.8)

    # ── Kafka ↔ SAP/HRIS ──
    _arrow(ax, 14.45, 6.5, 14.45, 5.7, "Events", color=M_SAP, lw=1.5)
    _arrow(ax, 14.45, 6.5, 14.45, 4.0, "Events", color=M_HRIS, lw=1.5)

    # Legend
    ax.plot([0.3, 15.7], [1.8, 1.8], color="#ccc", lw=0.8)
    _legend(
        ax,
        [
            (M_PERSON, "Pengguna"),
            (M_ODOO, "Odoo Module"),
            (M_BRIDGE, "SAP Bridge"),
            (M_KAFKA, "Kafka Bus"),
            (M_SAP, "SAP S/4HANA"),
            (M_HRIS, "HRIS"),
            (M_KC, "Keycloak"),
            (M_WARN, "Existing DB"),
        ],
        x=0.3,
        y=1.3,
        fs=8,
    )

    return _png(fig)


def diagram_kafka_integration():
    fig, ax = _fig(16, 9)
    ax.set_facecolor(M_BG)
    ax.text(
        8,
        8.6,
        "Integrasi Kafka-Only — Bridge sebagai Cache Layer",
        ha="center",
        fontsize=13,
        fontweight="bold",
        color="#1A2332",
    )

    # Rule banner
    rect = FancyBboxPatch(
        (2.5, 7.9),
        11.0,
        0.55,
        boxstyle="round,pad=0,rounding_size=0.1",
        facecolor="#FDEBD0",
        edgecolor="#E59866",
        lw=1.5,
    )
    ax.add_patch(rect)
    ax.text(
        8.0,
        8.17,
        "ATURAN: Bridge TIDAK BOLEH call SAP REST langsung — semua via Kafka topic",
        ha="center",
        va="center",
        fontsize=9,
        color="#784212",
        fontweight="bold",
    )

    # Columns: Odoo | Bridge | Kafka | SAP/HRIS
    cols = [1.5, 5.2, 9.0, 12.8]
    headers = ["Odoo 19\nFinance Portal", "SAP Bridge\n(FastAPI + Redis)", "Kafka Bus", "SAP / HRIS"]
    hcols = [M_ODOO, M_BRIDGE, M_KAFKA, M_SAP]
    for x, h, c in zip(cols, headers, hcols):
        _box(ax, x, 6.8, 2.8, 0.8, h, color=c, fontsize=9)
        ax.plot([x + 1.4, x + 1.4], [0.6, 6.8], color=c, lw=0.8, ls="--", alpha=0.4, zorder=1)

    # Rows: flows
    rows = [
        (
            5.8,
            "PUSH outbound\n(push approved doc)",
            [(1.5 + 2.8, 5.95), (5.2, 5.95)],
            [(5.2 + 2.8, 5.95), (9.0, 5.95)],
            [(9.0 + 2.8, 5.95), (12.8, 5.95)],
        ),
        (
            4.5,
            "Posting Confirm\n(SAP → Odoo)",
            [(12.8 + 2.8, 4.65), (9.0 + 2.8, 4.65)],
            [(9.0, 4.65), (5.2 + 2.8, 4.65)],
            [(5.2, 4.65), (1.5 + 2.8, 4.65)],
        ),
        (
            3.3,
            "Master Sync / PR Lookup\n(via Redis cache)",
            [(9.0 + 2.8, 3.45), (9.0, 3.45)],  # SAP→Kafka (dotted, background)
            [(9.0, 3.45), (5.2 + 2.8, 3.45)],  # Kafka→Bridge consume
            [(1.5 + 2.8, 3.45), (5.2, 3.45)],
        ),  # Odoo→Bridge (query cache)
        (
            2.1,
            "Payment Status\n(SAP → Odoo via webhook)",
            [(12.8 + 2.8, 2.25), (9.0 + 2.8, 2.25)],
            [(9.0, 2.25), (5.2 + 2.8, 2.25)],
            [(5.2, 2.25), (1.5 + 2.8, 2.25)],
        ),
    ]

    for y, label, seg1, seg2, seg3 in rows:
        ax.text(0.1, y + 0.12, label, va="center", fontsize=7.5, color="#444", style="italic")
        ax.plot([0.9, 1.5], [y, y], color="#aaa", lw=0.8)
        for (x1, y1), (x2, y2) in [seg1, seg2, seg3]:
            ax.annotate(
                "",
                xy=(x2, y2),
                xytext=(x1, y1),
                arrowprops=dict(arrowstyle="-|>", color="#555", lw=1.6, mutation_scale=14),
            )

    # Redis cache box inside Bridge column
    _box(
        ax,
        5.3,
        0.8,
        2.6,
        1.1,
        "Redis Cache",
        "pr · po_gr · vendor\nbudget · employee · travel",
        color="#C0392B",
        fontsize=8,
    )
    ax.annotate("", xy=(6.6, 1.9), xytext=(6.6, 3.1), arrowprops=dict(arrowstyle="-|>", color="#C0392B", lw=1.5))
    ax.text(6.9, 2.5, "store", fontsize=7.5, color="#C0392B", style="italic")

    return _png(fig)


def diagram_ocr_pipeline():
    fig, ax = _fig(13, 11)
    ax.set_facecolor(M_BG)
    ax.text(
        6.5,
        10.6,
        "OCR Pipeline — Scanned PDF (primary path)",
        ha="center",
        fontsize=12,
        fontweight="bold",
        color="#1A2332",
    )

    def process(x, y, w, h, text, color=M_ODOO):
        _box(ax, x, y, w, h, text, color=color, fontsize=8.5, radius=0.15)

    def diamond(ax, cx, cy, half_w, half_h, text, color=M_KAFKA):
        xs = [cx, cx + half_w, cx, cx - half_w, cx]
        ys = [cy + half_h, cy, cy - half_h, cy, cy + half_h]
        ax.fill(xs, ys, facecolor=color, edgecolor=M_WHITE, lw=1.5, zorder=3)
        ax.text(cx, cy, text, ha="center", va="center", fontsize=8, fontweight="bold", color=M_WHITE, zorder=4)

    def arr(x1, y1, x2, y2, lbl="", col="#555"):
        _arrow(ax, x1, y1, x2, y2, lbl, color=col, lw=1.8)

    # Start
    process(4.5, 9.9, 4.0, 0.55, "Vendor Upload PDF", color=M_PERSON)

    # Quality gate
    arr(6.5, 9.9, 6.5, 9.2)
    diamond(ax, 6.5, 8.75, 2.2, 0.45, "DPI >= 200? Blur OK?")
    arr(6.5, 8.3, 6.5, 7.7)  # yes down
    ax.text(6.65, 8.0, "Ya", fontsize=8, color=M_BRIDGE, fontweight="bold")
    # No → right → reject
    arr(8.7, 8.75, 10.5, 8.75, "", col=M_WARN)
    process(10.5, 8.38, 2.8, 0.75, "Tolak upload\n'Scan ulang, min 300 DPI'", color=M_WARN)
    ax.text(9.5, 8.9, "Tidak", fontsize=8, color=M_WARN, fontweight="bold")

    # Text layer check
    diamond(ax, 6.5, 7.35, 2.2, 0.45, "Ada text layer?\n(pdfplumber)")
    arr(6.5, 7.35, 6.5, 6.8)

    # Two paths
    # Left: pdfplumber (digital)
    arr(4.3, 7.35, 3.5, 7.35, "", col=M_BRIDGE)
    ax.text(3.9, 7.5, "Ya", fontsize=8, color=M_BRIDGE, fontweight="bold")
    process(1.8, 6.9, 2.5, 0.7, "pdfplumber\nextract text", color=M_BRIDGE)
    arr(3.05, 7.25, 3.05, 5.1, "", col=M_BRIDGE)

    # Right: Tesseract (primary)
    ax.text(6.65, 7.05, "Tidak\n(primary)", fontsize=8, color=M_ODOO, fontweight="bold")
    process(4.8, 6.3, 3.4, 0.75, "pdf2image → 300 DPI PNG", color=M_ODOO)
    arr(6.5, 6.3, 6.5, 5.8)
    process(4.8, 5.1, 3.4, 1.0, "OpenCV Pre-processing\nGrayscale · Denoise\nDeskew · Binarize · Upscale", color=M_ODOO)
    arr(6.5, 5.1, 6.5, 4.55)
    process(4.8, 3.85, 3.4, 0.75, "Tesseract 4 LSTM\nlang: ind+eng", color=M_ODOO)
    arr(6.5, 3.85, 6.5, 3.3)

    # Merge → regex parser
    process(4.2, 2.6, 4.6, 0.7, "Regex Parser: invoice no · date · NPWP\ntotal · PPN · PPh · PO/DO ref", color=M_BRIDGE)
    arr(3.05, 5.1, 3.05, 2.95, "", col=M_BRIDGE)
    ax.plot([3.05, 4.2], [2.95, 2.95], color=M_BRIDGE, lw=1.8)

    # Confidence scoring
    arr(6.5, 2.6, 6.5, 2.05)
    diamond(ax, 6.5, 1.65, 2.2, 0.42, "Confidence?")

    # Three outcomes
    # >= 0.88 → green
    arr(8.7, 1.65, 10.0, 1.65, "", col=M_BRIDGE)
    process(10.0, 1.28, 2.7, 0.75, "Pre-fill HIJAU\nVendor konfirmasi", color=M_BRIDGE)
    ax.text(9.2, 1.82, ">=0.88", fontsize=8, color=M_BRIDGE, fontweight="bold")

    # 0.70-0.87 → yellow
    arr(6.5, 1.23, 6.5, 0.7)
    ax.text(6.65, 0.95, "0.70-0.87", fontsize=8, color=M_KAFKA, fontweight="bold")
    process(4.8, 0.2, 3.4, 0.7, "Pre-fill KUNING\nVendor WAJIB periksa", color=M_KAFKA)

    # < 0.70 → red
    arr(4.3, 1.65, 3.0, 1.65, "", col=M_WARN)
    process(1.0, 1.28, 2.3, 0.75, "Field MERAH\nVendor isi manual", color=M_WARN)
    ax.text(3.3, 1.82, "<0.70", fontsize=8, color=M_WARN, fontweight="bold")

    return _png(fig)


def diagram_multicompany():
    fig, ax = _fig(14, 8)
    ax.set_facecolor(M_BG)
    ax.text(
        7.0,
        7.6,
        "Multi-Company Architecture — Single Finance Portal DB",
        ha="center",
        fontsize=12,
        fontweight="bold",
        color="#1A2332",
    )

    # DB container
    db_box = FancyBboxPatch(
        (0.3, 0.5),
        13.4,
        6.8,
        boxstyle="round,pad=0,rounding_size=0.3",
        facecolor="#EAF2FB",
        edgecolor="#1168BD",
        lw=2,
        zorder=1,
    )
    ax.add_patch(db_box)
    ax.text(
        7.0,
        7.1,
        "Finance Portal DB  (PostgreSQL — single database)",
        ha="center",
        fontsize=9,
        color=M_ODOO,
        fontweight="bold",
    )

    # Finance Group Manager (top center)
    _box(
        ax,
        5.0,
        5.9,
        4.0,
        0.85,
        "Finance Group Manager",
        "company_ids = [PT-A, PT-B, PT-C ...]\nDashboard: semua PT",
        color=M_HOLDING,
        fontsize=8,
    )

    # Three PTs
    pt_data = [
        (0.6, "PT ERA Busana\nRetailindo", "SAP company_code: 1000"),
        (5.0, "PT Erajaya\nSwasembada", "SAP company_code: 2000"),
        (9.4, "PT [Brand Lain]", "SAP company_code: 3000"),
    ]
    for x, name, code in pt_data:
        _box(ax, x, 4.5, 3.6, 1.1, name, code, color=M_PT, fontsize=8.5)
        # Arrow from group to PT
        _arrow(ax, 7.0, 5.9, x + 1.8, 5.6, color=M_HOLDING, lw=1.2)
        # Sub-items
        subs = [
            (x + 0.1, 3.45, 1.6, 0.7, "Dokumen\n(company_id)", M_ODOO),
            (x + 1.9, 3.45, 1.6, 0.7, "Budget\n+ Approval Matrix", "#2471A3"),
        ]
        for sx, sy, sw, sh, sl, sc in subs:
            _box(ax, sx, sy, sw, sh, sl, color=sc, fontsize=7.5)

    # Shared resources at bottom
    ax.plot([0.6, 13.1], [2.95, 2.95], color="#aaa", lw=1, ls="--")
    ax.text(7.0, 2.78, "Shared across all companies", ha="center", fontsize=8, color="#555", style="italic")
    shared = [
        (1.0, "res.partner\n(Vendor Master)\nSatu vendor →\nbanyak PT", M_SHARED),
        (4.5, "hr.employee\n(Employee Master)\nDari HRIS sync", M_SHARED),
        (8.0, "res.bank\n(Bank Master)", M_SHARED),
        (11.0, "res.users\n(Satu akun)\nAkses multi-PT", M_SHARED),
    ]
    for x, lbl, col in shared:
        _box(ax, x, 0.8, 2.8, 1.8, lbl, color=col, fontsize=8)

    return _png(fig)


def diagram_state_machine():
    """Document lifecycle state diagram."""
    fig, ax = _fig(15, 7)
    ax.set_facecolor(M_BG)
    ax.text(
        7.5,
        6.6,
        "Document Lifecycle — State Machine (semua tipe dokumen)",
        ha="center",
        fontsize=12,
        fontweight="bold",
        color="#1A2332",
    )

    states = {
        "draft": (1.0, 4.9),
        "submitted": (3.2, 4.9),
        "tax_review": (5.4, 4.9),
        "fin_review": (7.6, 4.9),
        "approved": (9.8, 4.9),
        "pushed": (9.8, 3.3),
        "posted": (9.8, 1.8),
        "paid": (9.8, 0.4),
        "rejected": (5.4, 3.3),
        "cancelled": (3.2, 2.3),
    }
    colors = {
        "draft": "#7F8C8D",
        "submitted": M_ODOO,
        "tax_review": "#D4AC0D",
        "fin_review": "#CA6F1E",
        "approved": M_BRIDGE,
        "pushed": "#1A5276",
        "posted": "#1F618D",
        "paid": "#117A65",
        "rejected": M_WARN,
        "cancelled": "#95A5A6",
    }
    labels = {
        "draft": "Draft",
        "submitted": "Submitted",
        "tax_review": "Tax\nReview",
        "fin_review": "Finance\nReview",
        "approved": "Approved",
        "pushed": "Pushed\n(SAP)",
        "posted": "Posted\n(GL)",
        "paid": "Paid",
        "rejected": "Rejected",
        "cancelled": "Cancelled",
    }
    W, H = 1.9, 0.75
    for name, (x, y) in states.items():
        _box(ax, x, y, W, H, labels[name], color=colors[name], fontsize=8.5)

    def sa(frm, to, lbl="", offset=(0, 0)):
        fx, fy = states[frm]
        tx, ty = states[to]
        x1 = fx + W / 2 + offset[0]
        y1 = fy + H / 2
        x2 = tx + W / 2 + offset[0]
        y2 = ty + H / 2
        _arrow(ax, x1, y1, x2, y2, lbl, color="#555", lw=1.5)

    # Main happy path
    for a, b, l in [
        ("draft", "submitted", "submit()"),
        ("submitted", "tax_review", "approval\ntrigger"),
        ("tax_review", "fin_review", "Tax\nApprove"),
        ("fin_review", "approved", "Finance\nApprove"),
    ]:
        sa(a, b, l)

    # Down path
    for a, b, l in [
        ("approved", "pushed", "push to SAP\n(queue_job)"),
        ("pushed", "posted", "SAP webhook\nposting-confirm"),
        ("posted", "paid", "SAP webhook\npayment-status"),
    ]:
        fx, fy = states[a]
        tx, ty = states[b]
        _arrow(ax, fx + W / 2, fy, tx + W / 2, ty + H, l, color=M_BRIDGE, lw=1.8)

    # Reject paths
    _arrow(
        ax,
        states["tax_review"][0] + W / 2,
        states["tax_review"][1],
        states["rejected"][0] + W / 2,
        states["rejected"][1] + H,
        "Reject",
        color=M_WARN,
        lw=1.5,
    )
    _arrow(
        ax,
        states["fin_review"][0] + W / 2,
        states["fin_review"][1],
        states["rejected"][0] + W / 2,
        states["rejected"][1] + H,
        "Reject",
        color=M_WARN,
        lw=1.5,
    )

    # Rejected → draft
    _arrow(
        ax,
        states["rejected"][0],
        states["rejected"][1] + H / 2,
        states["draft"][0] + W,
        states["draft"][1] + H / 2,
        "Revise",
        color="#E67E22",
        lw=1.3,
    )

    # Cancelled
    ax.text(1.5, 2.95, "cancel()\n(sebelum approved)", fontsize=7.5, color="#7F8C8D", ha="center", style="italic")
    _arrow(
        ax,
        states["submitted"][0] + W / 2,
        states["submitted"][1],
        states["cancelled"][0] + W / 2,
        states["cancelled"][1] + H,
        "",
        color="#95A5A6",
        lw=1.2,
    )

    # Start symbol
    ax.add_patch(plt.Circle((0.6, 5.27), 0.15, color="#333", zorder=5))
    _arrow(ax, 0.75, 5.27, 1.0, 5.27, "", color="#333", lw=2)

    # End symbol
    ax.add_patch(plt.Circle((10.75, 0.1), 0.18, color="#333", zorder=5))
    ax.add_patch(plt.Circle((10.75, 0.1), 0.12, color=M_WHITE, zorder=6))
    _arrow(ax, 10.75, 0.4, 10.75, 0.28, "", color="#333", lw=2)

    return _png(fig)


# ── Builder (docx) ────────────────────────────────────────────────────────────


class Builder:
    def __init__(self):
        self.doc = Document()
        normal = self.doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(10.5)
        normal.paragraph_format.space_after = Pt(5)
        for level, size, color in (
            ("Heading 1", 15, C_ACCENT),
            ("Heading 2", 12, C_INK),
            ("Heading 3", 10.5, C_INK),
        ):
            st = self.doc.styles[level]
            st.font.name = "Calibri"
            st.font.size = Pt(size)
            st.font.color.rgb = color
            st.font.bold = True
        for section in self.doc.sections:
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(2.2)
            section.right_margin = Cm(2.2)

    def p(self, text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=None, color=None, space_after=5):
        par = self.doc.add_paragraph()
        par.alignment = align
        par.paragraph_format.space_after = Pt(space_after)
        run = par.add_run(text)
        run.bold = bold
        run.italic = italic
        if size:
            run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
        return par

    def h(self, text, level):
        self.doc.add_heading(text, level=level)

    def note(self, text):
        par = self.doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        par.paragraph_format.left_indent = Cm(0.5)
        par.paragraph_format.space_after = Pt(4)
        run = par.add_run(text)
        run.italic = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    def code(self, text):
        par = self.doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.LEFT
        par.paragraph_format.left_indent = Cm(0.4)
        par.paragraph_format.space_after = Pt(3)
        run = par.add_run(text)
        run.font.name = "Courier New"
        run.font.size = Pt(8.5)

    def bullets(self, items):
        for it in items:
            par = self.doc.add_paragraph(style="List Bullet")
            par.paragraph_format.space_after = Pt(3)
            if isinstance(it, tuple):
                r = par.add_run(it[0])
                r.bold = True
                r.font.size = Pt(10)
                r2 = par.add_run(it[1])
                r2.font.size = Pt(10)
            else:
                r = par.add_run(str(it))
                r.font.size = Pt(10)

    def image(self, png_buf, width_cm=16.0, caption=""):
        par = self.doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = par.add_run()
        run.add_picture(png_buf, width=Cm(width_cm))
        if caption:
            self.note(caption)

    @staticmethod
    def _shade(cell, hexcolor):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), hexcolor)
        tcPr.append(shd)

    def table(self, headers, rows, widths=None, total_rows=None, highlight_rows=None, head_bg=HEAD_BG, cell_size=9.5):
        total_rows = total_rows or set()
        highlight_rows = highlight_rows or {}
        t = self.doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = t.rows[0].cells
        for i, htxt in enumerate(headers):
            hdr[i].text = ""
            run = hdr[i].paragraphs[0].add_run(htxt)
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(cell_size)
            self._shade(hdr[i], head_bg)
        for ridx, row in enumerate(rows):
            cells = t.add_row().cells
            is_total = ridx in total_rows
            row_bg = highlight_rows.get(ridx)
            for i, val in enumerate(row):
                cells[i].text = ""
                par = cells[i].paragraphs[0]
                par.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = par.add_run(str(val))
                run.font.size = Pt(cell_size)
                run.bold = is_total
                if is_total:
                    self._shade(cells[i], TOTAL_BG)
                elif row_bg:
                    self._shade(cells[i], row_bg)
        if widths:
            for i, w in enumerate(widths):
                for row in t.rows:
                    row.cells[i].width = Cm(w)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)
        return t

    def cover(self, kind, title, subtitle, attrs):
        for _ in range(5):
            self.doc.add_paragraph()
        self.p(kind, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=22, color=C_ACCENT, space_after=2)
        self.doc.add_paragraph()
        self.p(title, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
        self.p(subtitle, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
        for _ in range(3):
            self.doc.add_paragraph()
        self.table(["Atribut", "Keterangan"], attrs, widths=[5, 11])
        self.doc.add_page_break()

    def save(self, path):
        self.doc.save(path)


# ── Build ─────────────────────────────────────────────────────────────────────


def build():
    b = Builder()

    # ── Cover ──────────────────────────────────────────────────────────────
    b.cover(
        "DOKUMEN ARSITEKTUR",
        "Finance Portal on Odoo",
        "v2.0 · System of Engagement over SAP S/4HANA · Kafka-Only Integration",
        [
            ["Dokumen", "Architecture Document — Finance Portal (Odoo)"],
            ["Versi", "2.0"],
            ["Tanggal", "2026-06-25"],
            ["DB", "Finance Portal DB — baru, terpisah dari existing Odoo DB"],
            ["Scope", "Odoo 19 CE + SAP Bridge microservice (FastAPI + Redis)"],
            ["Lihat juga", "architecture.md (Mermaid) · Project-Charter-Odoo-Finance-Portal.md"],
        ],
    )

    # ── 1. ADR ─────────────────────────────────────────────────────────────
    b.h("1. Architecture Decision Record (ADR)", 1)
    b.table(
        ["#", "Keputusan", "Pilihan", "Alasan"],
        [
            ["ADR-01", "Peran Odoo", "System of Engagement (tanpa GL posting)", "SAP tetap System of Record"],
            [
                "ADR-02",
                "Database",
                "DB BARU terpisah dari existing Odoo DB",
                "Isolasi lifecycle; tidak ganggu operasional existing",
            ],
            [
                "ADR-03",
                "Multi-PT Erajaya",
                "Single Finance Portal DB + res.company per PT",
                "Consolidated dashboard; shared vendor master; satu SSO",
            ],
            [
                "ADR-04",
                "Integrasi SAP",
                "Kafka-only — Bridge TIDAK boleh call SAP REST langsung",
                "Constraint historis: direct SAP connection tidak disarankan",
            ],
            [
                "ADR-05",
                "Bridge pattern",
                "Kafka consumer/producer + Redis cache + HMAC REST relay",
                "Odoo tidak kenal Kafka; Bridge satu-satunya integration point",
            ],
            ["ADR-06", "SSO karyawan", "Keycloak OIDC via auth_oauth", "HRIS maintain employee identity"],
            [
                "ADR-07",
                "Login vendor",
                "Odoo native portal: self-registration + invite",
                "HRIS tidak maintain vendor; vendor bukan employee",
            ],
            [
                "ADR-08",
                "OCR engine",
                "Tesseract 4 LSTM (primary) + pdfplumber (fallback)",
                "CPU-only (VM tanpa GPU); scanned PDF sebagai format utama",
            ],
            ["ADR-09", "Travel dinas", "Read-only mirror dari HRIS", "HRIS owns travel engine"],
            ["ADR-10", "Async push", "queue_job with_delay()", "Hindari timeout; decouple dari SAP latency"],
        ],
        widths=[1.5, 3.5, 5, 7],
        cell_size=8.5,
    )

    # ── 2. System Context ──────────────────────────────────────────────────
    b.doc.add_page_break()
    b.h("2. System Context", 1)
    b.p(
        "Finance Portal berjalan di Odoo tenant DB baru (finance_portal) pada instance yang sama "
        "dengan existing Odoo DB — tidak saling mengganggu. Seluruh komunikasi ke SAP/HRIS "
        "hanya melalui Kafka; Bridge tidak pernah call SAP REST langsung."
    )
    b.image(
        diagram_system_context(),
        width_cm=16.5,
        caption="Gambar 1 — System Context: Finance Portal DB terpisah, integrasi Kafka-only.",
    )

    # ── 3. Module Architecture ─────────────────────────────────────────────
    b.h("3. Arsitektur Modul Odoo", 1)
    b.table(
        ["Modul", "Tanggung Jawab", "Depends Utama"],
        [
            [
                "custom_finance_portal",
                "Domain engagement inti: semua model dokumen (CA, Reimbursement, Vendor Invoice, Travel), "
                "finance.document.mixin, master data, approval Tax->Finance, vendor portal, dashboard, reporting",
                "custom_core, custom_approval_engine, custom_pdp_audit, mail, portal, hr, account",
            ],
            [
                "custom_finance_budget",
                "Cost budget per divisi/company/tahun; _check_document_budget() soft/hard; "
                "_finance_pr_threshold() resolving limitation",
                "custom_finance_portal",
            ],
            [
                "custom_finance_portal_sap",
                "Adapter SAP & HRIS (SapBridgeAdapter, HrisBridgeAdapter); async push via queue_job; "
                "cron master sync; inbound webhook secure_endpoint; finance.sync.log + Sync menu",
                "custom_finance_portal, custom_finance_budget, custom_adapter_framework, queue_job",
            ],
            [
                "custom_finance_portal_sso",
                "Override _auth_oauth_signin(); vendor self-registration controller (/finance/vendor/register); "
                "role → group mapping; auto-provision vendor ke res.partner",
                "custom_finance_portal, auth_oauth, custom_core, auth_jwt",
            ],
        ],
        widths=[4.5, 9, 3.5],
        cell_size=9,
    )

    # ── 4. Integrasi Kafka-Only ─────────────────────────────────────────────
    b.doc.add_page_break()
    b.h("4. Integrasi — Kafka Only (REVISED)", 1)
    b.p(
        "Constraint utama: Bridge TIDAK PERNAH membuka koneksi HTTP/JDBC ke SAP secara langsung. "
        "Semua data dari SAP/HRIS masuk via Kafka topics yang di-produce oleh tim SAP/HRIS. "
        "Bridge mengkonsumsi semua topics dan menyimpan data di Redis cache. "
        "Odoo query Bridge via HMAC REST; Bridge serve dari Redis — SAP tidak perlu expose API apapun."
    )
    b.image(
        diagram_kafka_integration(),
        width_cm=16.5,
        caption="Gambar 2 — Kafka-Only Integration: Bridge sebagai satu-satunya integration hub dengan Redis cache.",
    )

    b.h("4.1 Kafka Topic Map", 2)
    b.table(
        ["#", "Kafka Topic", "Producer", "Consumer", "Frekuensi"],
        [
            ["1", "sap.to.portal.pr-master", "SAP", "Bridge", "Event-driven (create/update/close)"],
            ["2", "sap.to.portal.po-gr-master", "SAP", "Bridge", "Event-driven"],
            ["3", "sap.to.portal.vendor-master", "SAP", "Bridge", "Daily + on-change"],
            ["4", "sap.to.portal.budget-master", "SAP", "Bridge", "Daily + on-revision"],
            ["5", "sap.to.portal.item-category", "SAP", "Bridge", "Weekly"],
            ["6", "sap.to.portal.bank-master", "SAP", "Bridge", "Weekly"],
            ["7", "sap.to.portal.payment-status", "SAP", "Bridge", "Event-driven"],
            ["8", "sap.to.portal.posting-confirm", "SAP", "Bridge", "Event-driven"],
            ["9", "hris.to.portal.employee-master", "HRIS", "Bridge", "Daily + on-change"],
            ["10", "hris.to.portal.travel-request", "HRIS", "Bridge", "Event-driven"],
            ["11", "portal.to.sap.cash-advance", "Bridge", "SAP", "Event-driven"],
            ["12", "portal.to.sap.reimbursement", "Bridge", "SAP", "Event-driven"],
            ["13", "portal.to.sap.vendor-invoice", "Bridge", "SAP", "Event-driven"],
            ["14", "portal.to.sap.travel-settlement", "Bridge", "SAP", "Event-driven"],
        ],
        widths=[0.7, 5.8, 2, 2, 6.5],
        cell_size=9,
        highlight_rows={10: "E8F8F5", 11: "E8F8F5", 12: "E8F8F5", 13: "E8F8F5"},
    )
    b.note(
        "Baris hijau = outbound (Odoo -> SAP). Nama topik, retention, dan partition count disepakati tim Kafka klien."
    )

    b.h("4.2 PR & PO Lookup via Cache (bukan direct SAP)", 2)
    b.p(
        "Lookup PR/PO tidak lagi memanggil SAP REST. SAP produce PR events ke Kafka secara "
        "event-driven; Bridge consume dan simpan di Redis dengan TTL 1 jam. "
        "Odoo query Bridge, Bridge serve dari cache."
    )
    b.bullets(
        [
            ("Cache HIT: ", "Response instan dari Redis — latensi <50ms"),
            (
                "Cache MISS (PR baru): ",
                "Bridge return {found: false, retry_after: 30}. "
                "UI Odoo tampilkan: 'PR sedang disinkronisasi dari SAP, coba lagi dalam 30 detik'",
            ),
            (
                "Implikasi ke SAP team: ",
                "SAP harus produce PR events secara event-driven (tidak batch) agar cache selalu fresh. PO/GR sama.",
            ),
        ]
    )

    # ── 5. Document Lifecycle ───────────────────────────────────────────────
    b.doc.add_page_break()
    b.h("5. Document Lifecycle", 1)
    b.image(
        diagram_state_machine(),
        width_cm=16.5,
        caption="Gambar 3 — State machine berlaku untuk semua tipe dokumen: CA, Reimbursement, Vendor Invoice, Travel Settlement.",
    )

    b.table(
        ["State", "Siapa yang menggerakkan", "Action di Odoo"],
        [
            ["draft → submitted", "Requester (submit())", "Validasi PR threshold + budget check"],
            ["submitted → tax_review", "approval.mixin (tier 1)", "Notifikasi Tax Reviewer"],
            ["tax_review → fin_review", "Tax Approver", "Notifikasi Finance Reviewer"],
            ["fin_review → approved", "Finance Approver", "Dokumen siap di-push"],
            ["approved → pushed", "queue_job worker", "_finance_push_to_sap() → Bridge → Kafka → SAP"],
            ["pushed → posted", "SAP webhook (Bridge)", "sap_document_no terisi; state = posted"],
            ["posted → paid", "SAP webhook (Bridge)", "payment_plan_date + amount terisi; state = paid"],
        ],
        widths=[4, 4, 9],
        cell_size=9,
    )

    # ── 6. HMAC Security ────────────────────────────────────────────────────
    b.h("6. HMAC Security Model", 1)
    b.p(
        "Canonical message: timestamp.encode() + raw_body_bytes — byte-identical di ketiga titik "
        "(Odoo outbound, Bridge inbound, Bridge outbound, Odoo inbound). "
        "Drift toleransi +/-300 detik; nonce di-cache untuk cegah replay."
    )
    b.table(
        ["Titik", "File", "Fungsi"],
        [
            [
                "Odoo Outbound",
                "custom_adapter_framework/models/adapter_base.py",
                "BaseAdapter._build_headers(): X-Timestamp + X-Signature",
            ],
            [
                "Bridge Inbound",
                "services/finance-sap-bridge/app/hmac_util.py",
                "verify(): canonical ts+body, drift +-300s",
            ],
            ["Bridge Outbound", "services/finance-sap-bridge/app/hmac_util.py", "sign(): X-Timestamp + X-Signature"],
            [
                "Odoo Inbound",
                "custom_core/controllers/secure_endpoint.py",
                "_verify_hmac(): canonical ts.encode()+raw_body, nonce cache",
            ],
        ],
        widths=[3.5, 6.5, 7],
        cell_size=9,
    )

    # ── 7. SSO & Vendor Auth ────────────────────────────────────────────────
    b.doc.add_page_break()
    b.h("7. Autentikasi — SSO & Vendor", 1)

    b.h("7.1 Employee SSO (Keycloak OIDC)", 2)
    b.bullets(
        [
            "User akses /web/login → Odoo redirect ke Keycloak realm.",
            "Keycloak: autentikasi (username/password + MFA jika diaktifkan).",
            "Keycloak kirim id_token + access_token (JWT) ke callback Odoo.",
            "Odoo _auth_oauth_signin(): verify token via Keycloak JWKS.",
            "Odoo _finance_sso_apply_roles(): parse realm_access.roles + resource_access.*.roles.",
            "Additive mapping: group_finance_* += user (tidak pernah revoke).",
            "Company switcher tersedia jika user punya akses ke beberapa PT.",
        ]
    )

    b.h("7.2 Vendor Login — Self-Registration", 2)
    b.p(
        "HRIS tidak maintain data vendor, sehingga vendor menggunakan Odoo native portal "
        "authentication (bukan Keycloak SSO). Dua jalur: self-registration dan invite."
    )
    b.table(
        ["Jalur", "Flow", "Syarat"],
        [
            [
                "Self-Registration",
                "Vendor buka /finance/vendor/register → isi nama PT + NPWP + email PIC → "
                "sistem match NPWP dengan res.partner dari SAP master sync → "
                "Finance admin approve → email aktivasi → vendor set password → login",
                "NPWP harus sudah ada di vendor master SAP (sync Kafka)",
            ],
            [
                "Invite (Finance → Vendor)",
                "Finance buka res.partner vendor → klik 'Grant Portal Access' → "
                "sistem buat res.users (portal) → kirim email invite → "
                "vendor klik link → set password → login",
                "res.partner vendor sudah ada (dari SAP sync)",
            ],
        ],
        widths=[3, 9, 5],
        cell_size=9,
    )
    b.note(
        "Jika NPWP tidak ditemukan saat self-registration: tampilkan pesan "
        "'NPWP belum terdaftar sebagai vendor aktif. Hubungi tim Procurement.'"
    )

    b.h("7.3 Role Mapping", 2)
    b.table(
        ["Keycloak Role", "Odoo Group", "Akses"],
        [
            ["finance_user", "group_finance_user", "Submit dokumen; lihat milik sendiri"],
            ["finance_tax", "group_finance_tax", "Review & approve tier 1 (Tax)"],
            ["finance_officer", "group_finance_officer", "Review tier 2; akses semua dokumen"],
            ["finance_manager", "group_finance_manager", "Full access + konfigurasi"],
            ["(Odoo portal)", "group_finance_vendor", "Vendor Portal: submit invoice + tracking"],
        ],
        widths=[4, 4, 9],
        cell_size=9.5,
    )

    # ── 8. OCR Pipeline ─────────────────────────────────────────────────────
    b.doc.add_page_break()
    b.h("8. OCR Pipeline", 1)
    b.p(
        "Format utama yang diterima adalah scanned PDF. Tesseract 4 LSTM berjalan di CPU "
        "(VM tanpa GPU). Pipeline memiliki quality gate di awal dan confidence scoring di akhir — "
        "semua hasil OCR wajib dikonfirmasi vendor sebelum submit."
    )
    b.image(
        diagram_ocr_pipeline(),
        width_cm=14.0,
        caption="Gambar 4 — OCR Pipeline: Tesseract sebagai primary path; pdfplumber shortcut untuk digital PDF.",
    )

    b.h("8.1 Akurasi Ekspektasi (CPU Tesseract, tanpa GPU)", 2)
    b.table(
        ["Kualitas Scan", "DPI", "Akurasi field numerik", "Akurasi tabel/line item"],
        [
            ["Bersih + datar", ">= 300", "87-90%", "75-82%"],
            ["Normal, sedikit miring", "200-300", "78-84%", "65-72%"],
            ["Buruk / terlipat / stamp", "< 200", "55-68%", "40-55%"],
        ],
        widths=[5, 2, 4, 6],
        cell_size=9.5,
    )

    b.h("8.2 Paket Python (CPU-only, tanpa GPU)", 2)
    b.code(
        "pdfplumber>=0.10          # digital PDF text extraction\n"
        "pytesseract>=0.3.10       # Tesseract Python binding\n"
        "tesseract-ocr             # system package (apt)\n"
        "tesseract-ocr-ind         # bahasa Indonesia\n"
        "opencv-python-headless>=4.8  # pre-processing\n"
        "Pillow>=10\n"
        "pdf2image>=1.16           # convert image PDF ke PNG\n"
        "poppler-utils             # dependency pdf2image (apt)"
    )

    b.h("8.3 Panduan Upload untuk Vendor", 2)
    b.bullets(
        [
            "Format: PDF hasil scan (bukan foto dari kamera HP)",
            "DPI minimal: 300 DPI",
            "Warna: hitam-putih / grayscale (lebih bersih dari berwarna)",
            "Dokumen: rata, tidak terlipat, tidak ada bayangan atau stamp menutupi angka",
            "Ukuran file: maksimal 10 MB per halaman",
        ]
    )

    # ── 9. Multi-Company ─────────────────────────────────────────────────────
    b.doc.add_page_break()
    b.h("9. Multi-Company Architecture", 1)
    b.p(
        "Satu Finance Portal DB berisi res.company untuk setiap PT Erajaya. "
        "Isolasi data per PT via company_id pada semua model dokumen, budget, dan approval matrix. "
        "Finance Group Manager memiliki akses ke semua PT dalam satu dashboard terpadu."
    )
    b.image(
        diagram_multicompany(),
        width_cm=15.5,
        caption="Gambar 5 — Multi-Company: satu DB, company_id per PT, vendor/employee/bank shared.",
    )

    b.h("9.1 Perubahan Teknis di Modul", 2)
    b.code(
        "# Semua model dokumen + budget + approval matrix\n"
        "company_id = fields.Many2one(\n"
        "    'res.company', required=True,\n"
        "    default=lambda self: self.env.company\n"
        ")\n\n"
        "# Security rule: user hanya lihat dokumen perusahaan yang dia akses\n"
        "[('company_id', 'in', self.env.user.company_ids.ids)]\n\n"
        "# res.company: tambah field SAP company code\n"
        "x_sap_company_code = fields.Char(string='SAP Company Code')\n\n"
        "# Outbound payload ke SAP: sertakan company_code\n"
        "payload['company_code'] = self.company_id.x_sap_company_code"
    )

    b.h("9.2 User Experience Multi-PT", 2)
    b.table(
        ["Role", "Akses Company", "Behaviour"],
        [
            [
                "Karyawan PT A",
                "company_ids = [PT-A]",
                "Submit → company_id = PT-A; company switcher tidak tampil (single company)",
            ],
            ["Finance Officer PT B", "company_ids = [PT-B]", "Hanya lihat dokumen PT-B"],
            [
                "Finance Group Manager",
                "company_ids = [PT-A, PT-B, PT-C ...]",
                "Dashboard 'All Companies' aggregate semua PT; bisa drill-down per PT",
            ],
            [
                "Vendor 'Supplier X'",
                "res.partner (shared, tidak company-specific)",
                "Lihat invoice yang dia submit ke semua PT; pilih 'Tujuan PT' saat submit invoice baru",
            ],
        ],
        widths=[4, 4, 9],
        cell_size=9,
    )

    # ── 10. Budget & PR Validation ─────────────────────────────────────────
    b.h("10. Budget & PR Validation", 1)
    b.table(
        ["Langkah", "Logic", "Konfigurasi"],
        [
            [
                "1. Resolve threshold",
                "_finance_pr_threshold(): (1) finance.limitation._resolve_for(doc) best-match "
                "submission_type+division; (2) config param custom_finance_portal.pr_required_threshold; "
                "(3) default Rp 1.000.000",
                "finance.limitation master atau config param",
            ],
            [
                "2. PR validation (jika amount > threshold)",
                "SapBridgeAdapter.lookup_pr(pr_number) → Bridge Redis → {is_valid, status, remaining_value}. "
                "Tolak jika not valid atau status = closed/cancelled",
                "PR data dari Kafka topic sap.to.portal.pr-master (event-driven)",
            ],
            [
                "3. Budget check",
                "finance.budget._check_document_budget(doc): hitung consumed dari dokumen "
                "in-state submitted→paid; bandingkan dengan total_budget",
                "custom_finance_budget.enforce = False (soft) atau True (hard)",
            ],
        ],
        widths=[3.5, 9, 4.5],
        cell_size=9,
    )

    # ── 11. Data Requirements ───────────────────────────────────────────────
    b.doc.add_page_break()
    b.h("11. Data yang Dikonsumsi Odoo", 1)
    b.p(
        "SEMUA data dari SAP/HRIS masuk ke Odoo melalui: Kafka -> Bridge (Redis cache) -> HMAC REST -> Odoo. "
        "Bridge tidak pernah call SAP REST. Odoo tidak pernah call Kafka langsung."
    )

    b.h("11.1 Master Data SAP (via Kafka -> Redis)", 2)
    b.table(
        ["Data Feed", "Kafka Topic", "Field Wajib", "Keterangan"],
        [
            [
                "Supplier/Vendor",
                "sap.to.portal.vendor-master",
                "external_id, name, npwp [PDP], bank_account_no [PDP], currency_code, sap_company_codes",
                "Upsert ke res.partner by x_sap_external_id",
            ],
            [
                "Division/Cost Center",
                "sap.to.portal.vendor-master",
                "external_id, name, code, company_code",
                "Upsert ke finance.vertical; mapping ke res.company",
            ],
            [
                "Cost Budget",
                "sap.to.portal.budget-master",
                "external_id, division_external_id, fiscal_year, total_budget, company_code",
                "Consumed dihitung mandiri oleh Odoo",
            ],
            [
                "Item Category",
                "sap.to.portal.item-category",
                "external_id, name, code, is_active",
                "Dropdown baris CA/Reimb/Invoice",
            ],
            [
                "Approval Matrix",
                "sap.to.portal.vendor-master",
                "document_type, tier (1=Tax/2=Finance), approver_group",
                "Workaround: konfigurasi manual jika feed tidak ada",
            ],
            [
                "Bank Master",
                "sap.to.portal.bank-master",
                "external_id, name, bic, country_code",
                "Ref rekening karyawan; workaround: manual input",
            ],
        ],
        widths=[3.5, 5, 5, 3.5],
        cell_size=8.5,
    )

    b.h("11.2 Lookup via Bridge Cache", 2)
    b.table(
        ["Lookup", "Bridge Endpoint", "Field Response Wajib", "Behaviour jika Cache MISS"],
        [
            [
                "PR Lookup",
                "GET /from-odoo/finance/pr/lookup",
                "is_valid, status (open/partial/closed), total_value, remaining_value, cost_center_external_id",
                "Return {found:false, retry_after:30}; UI tampilkan pesan retry",
            ],
            [
                "PO/GR Lookup",
                "GET /from-odoo/finance/po-gr/lookup",
                "po_status, vendor_external_id, total_po_value, items[].unbilled_quantity",
                "Return {found:false}; Finance input manual atau tunggu sync",
            ],
        ],
        widths=[2.5, 4.5, 6, 4],
        cell_size=8.5,
    )

    b.h("11.3 Inbound Webhook dari SAP", 2)
    b.code(
        "// Posting Confirmation (state: pushed -> posted)\n"
        '{"odoo_ref":"CA/2025/00042","sap_document_no":"5100000123",\n'
        ' "sap_posting_date":"2025-01-08","event_type":"posted"}\n\n'
        "// Payment Status (state: posted -> paid)\n"
        '{"odoo_ref":"CA/2025/00042","sap_document_no":"5100000123",\n'
        ' "event_type":"payment_scheduled","payment_plan_date":"2025-01-15",\n'
        ' "payment_amount":3500000.00,"payment_currency":"IDR"}'
    )

    b.h("11.4 Data HRIS", 2)
    b.table(
        ["Data Feed", "Kafka Topic", "Field Wajib [PDP]", "Dipakai untuk"],
        [
            [
                "Employee Master",
                "hris.to.portal.employee-master",
                "employee_id, email, division_external_id, cost_center_external_id, "
                "bank_account_no [PDP], nik [PDP], employment_status",
                "Auto-fill divisi; budget scope; rekening payment",
            ],
            [
                "Travel Request",
                "hris.to.portal.travel-request",
                "hris_travel_id, employee_id, purpose, departure_date, return_date, "
                "approved_advance, status (approved/travelling), allowance_items[]",
                "finance.travel.settlement; filter status = approved/travelling",
            ],
        ],
        widths=[3.5, 5, 5.5, 3],
        cell_size=8.5,
    )

    b.h("11.5 Readiness Tracker", 2)
    b.table(
        ["#", "Data Feed", "Kafka Topic", "Frekuensi", "Status", "Blocker jika Tidak Ada"],
        [
            [
                "1",
                "Supplier/Vendor",
                "sap.to.portal.vendor-master",
                "Daily+change",
                "TBD",
                "Vendor Invoice tidak bisa divalidasi",
            ],
            ["2", "Division/Cost Center", "sap.to.portal.vendor-master", "Daily", "TBD", "Budget check tidak jalan"],
            ["3", "Cost Budget", "sap.to.portal.budget-master", "Daily+rev.", "TBD", "W: soft bypass"],
            ["4", "Item Category", "sap.to.portal.item-category", "Weekly", "TBD", "Dropdown kosong"],
            ["5", "Approval Matrix", "sap.to.portal.vendor-master", "Daily", "TBD", "W: manual config di Odoo"],
            ["6", "Bank Master", "sap.to.portal.bank-master", "Weekly", "TBD", "W: manual input"],
            ["7", "PR Master (event!)", "sap.to.portal.pr-master", "Event-driven", "TBD", "PR validation tidak aktif"],
            [
                "8",
                "PO/GR Master (event!)",
                "sap.to.portal.po-gr-master",
                "Event-driven",
                "TBD",
                "Vendor Invoice PO tidak bisa validasi",
            ],
            ["9", "Posting Confirm", "sap.to.portal.posting-confirm", "Event", "TBD", "State tidak bergerak ke posted"],
            ["10", "Payment Status", "sap.to.portal.payment-status", "Event", "TBD", "Payment tracking tidak update"],
            [
                "11",
                "Employee Master",
                "hris.to.portal.employee-master",
                "Daily+change",
                "TBD",
                "User provisioning manual",
            ],
            [
                "12",
                "Travel Request",
                "hris.to.portal.travel-request",
                "Event",
                "TBD",
                "Travel Settlement tidak bisa dibuat",
            ],
        ],
        widths=[0.6, 3.5, 5, 2.5, 1.5, 4],
        cell_size=8.5,
    )
    b.note(
        "Kolom Status diisi tim SAP/HRIS klien saat workshop kontrak integrasi (Milestone M1, W3). "
        "W = workaround tersedia. Feed #7 dan #8 (PR/PO) WAJIB event-driven — batch harian tidak cukup."
    )

    # ── 12. Outbound Payloads ───────────────────────────────────────────────
    b.h("12. Outbound Payloads (Odoo → Bridge → Kafka → SAP)", 1)

    b.h("12.1 Cash Advance GL Posting", 2)
    b.code(
        '{"odoo_ref":"CA/2025/00042","doc_type":"cash_advance","company_code":"1000",\n'
        ' "posting_date":"2025-01-08","employee_id":"EMP-00234",\n'
        ' "cost_center_external_id":"CC-FINANCE-01","currency_code":"IDR",\n'
        ' "total_amount":3500000.00,"pr_number":"PR-2025-00010",\n'
        ' "bank_account_no":"1234567890","bank_name":"BCA",\n'
        ' "lines":[{"item_external_id":"ITEM-001","description":"Tiket","amount":1000000}],\n'
        ' "attachments":[{"filename":"surat_tugas.pdf","content_base64":"..."}]}'
    )

    b.h("12.2 Vendor Invoice MIRO", 2)
    b.code(
        '{"odoo_ref":"INV/2025/00015","doc_type":"vendor_invoice",\n'
        ' "invoice_subtype":"po_non_trade","company_code":"1000",\n'
        ' "vendor_external_id":"V-10001","invoice_date":"2025-01-05",\n'
        ' "vendor_invoice_no":"INV-2024-456","po_number":"PO-2024-00456",\n'
        ' "currency_code":"IDR","amount_net":25000000,"tax_code":"PPh23_2pct",\n'
        ' "tax_amount":500000,"total_amount":24500000,\n'
        ' "cost_center_external_id":"CC-FINANCE-01",\n'
        ' "lines":[{"po_line":1,"gr_number":"GR-2024-00789","quantity":1,"unit_price":25000000}],\n'
        ' "attachments":[{"filename":"invoice.pdf","content_base64":"..."}]}'
    )

    # ── 13. Deployment ──────────────────────────────────────────────────────
    b.doc.add_page_break()
    b.h("13. Deployment", 1)

    b.h("13.1 Stack Docker Compose", 2)
    b.table(
        ["Container", "Port", "Keterangan"],
        [
            ["nginx", "443 (ext)", "TLS termination; routing ke finance_portal DB via dbfilter"],
            ["odoo-app", "8069 (int)", "Multi-worker; Finance Portal modules + queue_job worker"],
            ["odoo-mgmt", "8072 (int)", "1 worker; admin only; share filestore dengan odoo-app"],
            ["postgres", "5432 (int)", "existing_db + finance_portal DB (dua DB, satu instance)"],
            ["redis", "6379 (int)", "queue_job broker"],
            ["finance-sap-bridge", "8080 (int)", "Kafka consumer/producer + Redis cache + HMAC REST"],
            ["redis (bridge)", "6380 (int)", "Redis untuk Bridge cache (terpisah dari queue_job)"],
        ],
        widths=[4, 2.5, 10.5],
        cell_size=9,
    )
    b.note("PENTING: odoo-app dan odoo-mgmt WAJIB mount volume ./data/odoo-filestore yang sama.")

    b.h("13.2 Urutan Deploy Finance Portal DB", 2)
    b.code(
        "# 1. Buat DB baru\n"
        "make create-db DB=finance_portal\n\n"
        "# 2. Install modul (urutan sesuai depends)\n"
        "make install MODULE=custom_finance_portal      DB=finance_portal\n"
        "make install MODULE=custom_finance_budget      DB=finance_portal\n"
        "make install MODULE=custom_finance_portal_sap  DB=finance_portal\n"
        "make install MODULE=custom_finance_portal_sso  DB=finance_portal\n\n"
        "# 3. WAJIB restart (Python tidak re-import dari make update)\n"
        "docker compose restart odoo-app\n\n"
        "# 4. Deploy bridge\n"
        "docker compose up -d finance-sap-bridge\n\n"
        "# 5. Konfigurasi di Odoo Settings:\n"
        "#    - Bridge URL + HMAC secret (dua arah)\n"
        "#    - Keycloak OAuth provider (realm + client_id)\n"
        "#    - res.company per PT + x_sap_company_code\n"
        "#    - Approval matrix per company (tier 1=Tax, tier 2=Finance)\n"
        "#    - Aktifkan cron sync master SAP"
    )

    # ── 14. PDP ─────────────────────────────────────────────────────────────
    b.h("14. PDP Compliance", 1)
    b.table(
        ["Field", "Model", "Klasifikasi UU PDP", "Perlakuan"],
        [
            ["nik", "hr.employee", "Spesifik (Pasal 4)", "pdp_class='specific'; masking view non-privileged"],
            [
                "bank_account_no",
                "res.partner, hr.employee",
                "Spesifik",
                "Masking 4 digit terakhir; full hanya group_finance_officer+",
            ],
            ["npwp", "res.partner", "Spesifik", "Masking 6 digit tengah; full hanya group_finance_officer+"],
            [
                "Transisi state",
                "finance.document.mixin",
                "—",
                "pdp.audited.mixin → INSERT INTO pdp.audit_log (raw SQL)",
            ],
            [
                "Log adapter",
                "custom.adapter.call.log",
                "—",
                "BaseAdapter._sanitize_log() strip field PDP sebelum insert log",
            ],
        ],
        widths=[4, 4, 3.5, 5.5],
        cell_size=9,
    )

    # ── 15. Component Summary ───────────────────────────────────────────────
    b.h("15. Ringkasan Komponen", 1)
    b.table(
        ["Komponen", "Stack", "Lokasi", "In-Scope"],
        [
            ["custom_finance_portal", "Odoo 19 Python/XML", "addons/ee_gap/custom_finance_portal/", "Ya"],
            ["custom_finance_budget", "Odoo 19 Python/XML", "addons/ee_gap/custom_finance_budget/", "Ya"],
            ["custom_finance_portal_sap", "Odoo 19 Python/XML", "addons/ee_gap/custom_finance_portal_sap/", "Ya"],
            ["custom_finance_portal_sso", "Odoo 19 Python/XML", "addons/ee_gap/custom_finance_portal_sso/", "Ya"],
            ["finance-sap-bridge", "FastAPI Py3.12 + Redis", "services/finance-sap-bridge/", "Ya (scaffold)"],
            ["OCR service", "pdfplumber + Tesseract + OpenCV", "services/ocr-service/ atau inline bridge", "Ya"],
            ["Keycloak", "Keycloak 24+", "Infra klien", "Tidak"],
            ["Kafka Cluster + konektor", "Confluent / MSK", "Infra klien", "Tidak"],
            ["SAP S/4HANA + Kafka producer", "ABAP / CPI", "Tim SAP klien", "Tidak"],
            ["HRIS + Kafka producer", "Vendor HRIS", "Tim HC klien", "Tidak"],
        ],
        widths=[5, 4, 5.5, 2.5],
        highlight_rows={6: LGREY, 7: LGREY, 8: LGREY, 9: LGREY},
        cell_size=9,
    )

    # ── Save ────────────────────────────────────────────────────────────────
    out = os.path.join(HERE, "Architecture - Finance Portal Odoo v2.0.docx")
    b.save(out)
    print(f"Saved : {out}")
    print(f"Size  : {os.path.getsize(out):,} bytes")


if __name__ == "__main__":
    build()
