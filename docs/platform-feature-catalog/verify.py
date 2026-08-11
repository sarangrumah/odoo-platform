#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Check the built artefacts against the repository they claim to describe.

The catalog's whole selling point is that its numbers came from the repo. These
checks are what make that claim testable rather than asserted.

Usage:
    python3 verify.py               # everything
    python3 verify.py --skip-pdf    # skip the PDF/DOCX checks
    python3 verify.py --published   # additionally check the share copy
"""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import subprocess
import sys

HERE = os.path.dirname(os.path.abspath(__file__))
REPO = os.path.abspath(os.path.join(HERE, "..", ".."))
DIST = os.path.join(HERE, "dist")
PUBLISH_DIR = "/srv/sftp-share/files/katalog-fitur-platform"

PDF = os.path.join(DIST, "Katalog_Fitur_Platform_Odoo_Erajaya.pdf")
DOCX = os.path.join(DIST, "Katalog_Fitur_Platform_Odoo_Erajaya.docx")
XLSX = os.path.join(DIST, "Matriks_Fitur_Platform_Odoo_Erajaya.xlsx")

PDF_PAGES_MIN, PDF_PAGES_MAX = 90, 260

results: list[tuple[str, str, str]] = []


def check(name):
    def wrap(fn):
        def run(*a, **k):
            try:
                detail = fn(*a, **k)
                results.append(("PASS", name, detail or ""))
            except AssertionError as exc:
                results.append(("FAIL", name, str(exc)))
            except Exception as exc:  # noqa: BLE001 - report, never abort the run
                results.append(("ERROR", name, f"{type(exc).__name__}: {exc}"))
        return run
    return wrap


def load_catalog():
    with open(os.path.join(HERE, "catalog.json"), "r", encoding="utf-8") as fh:
        return json.load(fh)


def scan_disk() -> dict[str, int]:
    """Recount modules straight off the filesystem, independently of the
    generator, so a bug in the generator cannot validate itself."""
    counts: dict[str, int] = {}
    addons = os.path.join(REPO, "addons")
    for root, dirs, files in os.walk(addons):
        dirs[:] = [d for d in dirs if d not in ("__pycache__", ".git")]
        if "__manifest__.py" in files:
            group = os.path.relpath(root, addons).split(os.sep)[0]
            counts[group] = counts.get(group, 0) + 1
    return counts


# --- checks -----------------------------------------------------------------

@check("1. module total matches a fresh filesystem scan")
def c_total(cat):
    disk = sum(scan_disk().values())
    claimed = cat["meta"]["counts"]["modules_total"]
    assert disk == claimed, f"disk {disk} != catalog {claimed}"
    return f"{disk} modules"


@check("2. per-group counts match a fresh scan")
def c_groups(cat):
    disk, claimed = scan_disk(), cat["meta"]["counts"]["by_group"]
    diff = {g: (disk.get(g), claimed.get(g)) for g in set(disk) | set(claimed)
            if disk.get(g) != claimed.get(g)}
    assert not diff, f"mismatch {diff}"
    return f"{len(claimed)} groups"


@check("3. every module on disk is classified in taxonomy.py")
def c_taxonomy(cat):
    sys.path.insert(0, HERE)
    import taxonomy
    on_disk = set()
    addons = os.path.join(REPO, "addons")
    for root, dirs, files in os.walk(addons):
        dirs[:] = [d for d in dirs if d not in ("__pycache__", ".git")]
        if "__manifest__.py" in files:
            on_disk.add(os.path.basename(root))
    missing = on_disk - set(taxonomy.DOMAIN_BY_MODULE)
    stale = set(taxonomy.DOMAIN_BY_MODULE) - on_disk
    assert not missing, f"unclassified: {sorted(missing)}"
    assert not stale, f"classified but gone: {sorted(stale)}"
    return f"{len(on_disk)} classified"


@check("4. domain counts reconcile to the total")
def c_domains(cat):
    total = sum(d["module_count"] for d in cat["domains"])
    assert total == cat["meta"]["counts"]["modules_total"], \
        f"domains sum {total} != {cat['meta']['counts']['modules_total']}"
    return f"{len(cat['domains'])} domains"


@check("5. XLSX matrix has one row per module")
def c_xlsx_rows(cat):
    from openpyxl import load_workbook
    wb = load_workbook(XLSX, read_only=True)
    expected = cat["meta"]["counts"]["modules_total"] + 1
    got = wb["Matriks Fitur"].max_row
    assert got == expected, f"{got} rows, expected {expected}"
    return f"{got - 1} modules + header"


@check("6. XLSX pivot grand total equals the module total")
def c_xlsx_pivot(cat):
    from openpyxl import load_workbook
    wb = load_workbook(XLSX, read_only=True)
    ws = wb["Domain x Grup"]
    grand = ws.cell(row=ws.max_row, column=ws.max_column).value
    assert grand == cat["meta"]["counts"]["modules_total"], \
        f"pivot total {grand}"
    return f"{grand}"


@check("7. PDF page count is within a sane band")
def c_pdf_pages(cat):
    out = subprocess.run(["pdfinfo", PDF], capture_output=True, text=True, check=True).stdout
    pages = int(next(l for l in out.splitlines() if l.startswith("Pages:")).split()[1])
    assert PDF_PAGES_MIN <= pages <= PDF_PAGES_MAX, \
        f"{pages} pages outside {PDF_PAGES_MIN}-{PDF_PAGES_MAX}"
    return f"{pages} pages"


def pdf_text() -> str:
    return subprocess.run(["pdftotext", "-layout", PDF, "-"],
                          capture_output=True, text=True, check=True).stdout


@check("8. every module name reached the PDF")
def c_pdf_modules(cat):
    text = pdf_text()
    missing = [m["name"] for m in cat["modules"] if m["name"] not in text]
    assert not missing, f"absent from PDF: {missing[:8]}{' …' if len(missing) > 8 else ''}"
    return f"{len(cat['modules'])} names found"


@check("9. no mojibake in the PDF text")
def c_pdf_mojibake(cat):
    text = pdf_text()
    bad = [tok for tok in ("â€", "Ã¢", "Ã©", "Â·", "ï¿½") if tok in text]
    assert not bad, f"found {bad}"
    return "clean"


@check("10. DOCX opens and carries a table of contents")
def c_docx(cat):
    from docx import Document
    doc = Document(DOCX)
    paras = len(doc.paragraphs)
    assert paras > 500, f"only {paras} paragraphs — conversion probably truncated"
    # pandoc writes the TOC as a Word field, not as literal text: Word fills it
    # in when the document is opened. Looking for the words "Daftar Isi" finds
    # nothing even on a correct file.
    xml = doc.element.xml
    assert "fldChar" in xml and "TOC" in xml, "no TOC field in the document body"
    return f"{paras} paragraphs, {len(doc.tables)} tables, TOC field present"


@check("11. no secret values leaked into the PDF")
def c_secrets(cat):
    """Architecture and parameter *names* are published on purpose. Actual
    secret *values* are not, so scan for anything that looks like one."""
    text = pdf_text()
    patterns = [
        r"\b[0-9a-f]{64}\b",                       # hex key / raw digest
        r"-----BEGIN [A-Z ]*PRIVATE KEY-----",
    ]
    hits = [m if isinstance(m, str) else " ".join(m)
            for pat in patterns for m in re.findall(pat, text)]

    # `<name> = <value>` where the name reads like a credential. Documentation
    # legitimately shows how a credential is *produced*
    # ("token = secrets.token_urlsafe(24)") or where one goes
    # ("Token: <api_token>"); a literal value is the thing to catch.
    for assign in re.findall(
        r"(?i)(?:secret|password|api[_-]?key|token)\s*[:=]\s*(\S{8,})", text
    ):
        looks_like_code = "(" in assign or assign.startswith(("<", "{", "$", "%"))
        if not looks_like_code:
            hits.append(assign[:60])

    # A long base64 run needs its own rule. The naive character-class version
    # fires on every slash-separated enum the document quotes
    # ("draft/confirmed/exported/…", "BCA/Mandiri/BNI/…"), which buries any real
    # finding. Require mixed case, a digit, and at most one slash.
    for tok in re.findall(r"\b[A-Za-z0-9+/]{40,}={0,2}\b", text):
        if (tok.count("/") <= 1
                and any(c.isdigit() for c in tok)
                and any(c.islower() for c in tok)
                and any(c.isupper() for c in tok)):
            hits.append(tok[:60])

    assert not hits, f"possible secret-shaped strings: {hits[:5]}"
    return "no secret-shaped strings"


@check("12. docs/architecture.md tier table agrees with the repo")
def c_arch(cat):
    path = os.path.join(REPO, "docs", "architecture.md")
    with open(path, "r", encoding="utf-8") as fh:
        doc = fh.read()
    claimed = cat["meta"]["counts"]["by_group"]
    bad = []
    for group, count in claimed.items():
        m = re.search(rf"^\|\s*`{re.escape(group)}/`\s*\|\s*(\d+)\s*\|", doc, re.M)
        if not m:
            bad.append(f"{group}: no row")
        elif int(m.group(1)) != count:
            bad.append(f"{group}: doc says {m.group(1)}, repo has {count}")
    assert not bad, "; ".join(bad)
    return f"{len(claimed)} rows agree"


@check("13. published copies match dist/ byte for byte")
def c_published(cat):
    bad = []
    for f in (PDF, DOCX, XLSX):
        target = os.path.join(PUBLISH_DIR, os.path.basename(f))
        if not os.path.isfile(target):
            bad.append(f"{os.path.basename(f)}: not published")
            continue
        if sha256(f) != sha256(target):
            bad.append(f"{os.path.basename(f)}: checksum differs")
    assert not bad, "; ".join(bad)
    return f"{PUBLISH_DIR}"


@check("14. published files are world-readable, not world-writable")
def c_perms(cat):
    bad = []
    for f in (PDF, DOCX, XLSX):
        target = os.path.join(PUBLISH_DIR, os.path.basename(f))
        if not os.path.isfile(target):
            continue
        mode = os.stat(target).st_mode & 0o777
        if mode & 0o002 or not mode & 0o004:
            bad.append(f"{os.path.basename(f)}: {oct(mode)}")
    assert not bad, "; ".join(bad)
    return "0644"


def sha256(path: str) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as fh:
        for chunk in iter(lambda: fh.read(1 << 20), b""):
            h.update(chunk)
    return h.hexdigest()


# --- main -------------------------------------------------------------------


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--skip-pdf", action="store_true")
    ap.add_argument("--published", action="store_true")
    args = ap.parse_args()

    cat = load_catalog()

    c_total(cat)
    c_groups(cat)
    c_taxonomy(cat)
    c_domains(cat)
    c_xlsx_rows(cat)
    c_xlsx_pivot(cat)
    if not args.skip_pdf:
        c_pdf_pages(cat)
        c_pdf_modules(cat)
        c_pdf_mojibake(cat)
        c_docx(cat)
        c_secrets(cat)
    c_arch(cat)
    if args.published:
        c_published(cat)
        c_perms(cat)

    width = max(len(n) for _s, n, _d in results)
    for status, name, detail in results:
        mark = {"PASS": "ok  ", "FAIL": "FAIL", "ERROR": "ERR "}[status]
        print(f"  {mark}  {name.ljust(width)}   {detail}")

    failed = [r for r in results if r[0] != "PASS"]
    print(f"\n{len(results) - len(failed)}/{len(results)} checks passed")
    return 1 if failed else 0


if __name__ == "__main__":
    raise SystemExit(main())
