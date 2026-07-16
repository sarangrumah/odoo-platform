# -*- coding: utf-8 -*-
"""Generate the Finance Portal (Odoo) Project Charter as DOCX.

Mandays-only (no commercial/rate). Scope: Odoo only — SAP & HC/HRIS effort excluded.
House style mirrors docs/projects/levis/build_tsd_charter_workflow_docx.py.

Usage:  python docs/projects/finance-portal/build_charter_docx.py
"""
import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
INK = RGBColor(0x1A, 0x23, 0x32)
ACCENT = RGBColor(0x71, 0x4B, 0x67)
HEAD_BG = "1A2332"


class Builder:
    def __init__(self):
        self.doc = Document()
        normal = self.doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)
        normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        normal.paragraph_format.space_after = Pt(6)
        for level, size in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)):
            st = self.doc.styles[level]
            st.font.name = "Calibri"
            st.font.size = Pt(size)
            st.font.color.rgb = ACCENT if level == "Heading 1" else INK
            st.font.bold = True
        for section in self.doc.sections:
            section.top_margin = Cm(2.2)
            section.bottom_margin = Cm(2.2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

    def p(self, text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
          size=None, color=None, space_after=6):
        par = self.doc.add_paragraph()
        par.alignment = align
        par.paragraph_format.space_after = Pt(space_after)
        run = par.add_run(text)
        run.bold = bold
        run.italic = italic
        if size:
            run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
        return par

    def h(self, text, level):
        self.doc.add_heading(text, level=level)

    def bullets(self, items):
        for it in items:
            par = self.doc.add_paragraph(style="List Bullet")
            par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if isinstance(it, tuple):
                r = par.add_run(it[0]); r.bold = True
                par.add_run(it[1])
            else:
                par.add_run(it)

    @staticmethod
    def _shade(cell, hexcolor):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), hexcolor)
        tcPr.append(shd)

    def table(self, headers, rows, widths=None, total_rows=None):
        total_rows = total_rows or set()
        t = self.doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = t.rows[0].cells
        for i, htxt in enumerate(headers):
            hdr[i].text = ""
            run = hdr[i].paragraphs[0].add_run(htxt)
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9.5)
            self._shade(hdr[i], HEAD_BG)
        for ridx, row in enumerate(rows):
            cells = t.add_row().cells
            emphasize = ridx in total_rows
            for i, val in enumerate(row):
                cells[i].text = ""
                par = cells[i].paragraphs[0]
                par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                run = par.add_run(str(val))
                run.font.size = Pt(9.5)
                run.bold = emphasize
                if emphasize:
                    self._shade(cells[i], "EDE7EC")
        if widths:
            for i, w in enumerate(widths):
                for row in t.rows:
                    row.cells[i].width = Cm(w)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)
        return t

    def cover(self, kind, title, subtitle, attrs):
        for _ in range(5):
            self.doc.add_paragraph()
        self.p(kind, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=22,
               color=ACCENT, space_after=2)
        self.doc.add_paragraph()
        self.p(title, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
        self.p(subtitle, align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
        for _ in range(3):
            self.doc.add_paragraph()
        self.table(["Atribut", "Keterangan"], attrs, widths=[5, 11])
        self.doc.add_page_break()

    def save(self, path):
        self.doc.save(path)


def build():
    b = Builder()

    # ---------------- Cover ----------------
    b.cover(
        "PROJECT CHARTER",
        "Finance Portal on Odoo",
        "System of Engagement over SAP — Scope Odoo Only (Mandays)",
        [
            ["Judul Dokumen", "Project Charter — Finance Portal on Odoo"],
            ["Versi", "1.0"],
            ["Tanggal", "2026-06-23"],
            ["Disusun oleh", "Delivery Team (Custom Platform)"],
            ["Status", "Draft for Approval"],
            ["Scope Effort", "Odoo only — effort SAP & HC/HRIS dikeluarkan"],
        ],
    )

    # ---------------- Revision History ----------------
    b.h("Document Control", 1)
    b.table(
        ["Versi", "Tanggal", "Penulis", "Catatan"],
        [["1.0", "2026-06-23", "Delivery Team", "Initial charter"]],
        widths=[2, 3, 4, 7],
    )

    # ---------------- 1. Latar Belakang ----------------
    b.h("1. Latar Belakang & Justifikasi", 1)
    b.p(
        "Proses keuangan (Cash Advance, Reimbursement & Expenses, Vendor Invoice Non-Trade, "
        "Perjalanan Dinas) belum memiliki portal pengajuan terpadu di sisi pengguna, sementara "
        "SAP S/4HANA tetap menjadi system of record (posting GL, MIRO, pembayaran, master data). "
        "Dibutuhkan Finance Portal sebagai system of engagement yang menyediakan form pengajuan, "
        "approval Tax -> Finance, validasi budget/PR, lalu mendorong dokumen yang disetujui ke SAP "
        "dan menampilkan kembali status pembayaran."
    )
    b.p(
        "Solusi dibangun di atas Odoo 19 dengan memanfaatkan komponen platform yang sudah ada "
        "(approval engine, adapter framework, queue/async, audit PDP, reporting) sehingga lebih "
        "cepat dan hemat dibanding aplikasi bespoke. Integrasi ke SAP/HRIS dilakukan melalui "
        "Kafka + REST dengan pendekatan contract-first."
    )

    # ---------------- 2. Tujuan ----------------
    b.h("2. Tujuan Proyek (Objectives)", 1)
    b.bullets([
        "Menyediakan portal terpadu: Cash Advance (+Realization), Reimbursement & Expenses, "
        "Vendor Invoice (PO/Non-PO Non-Trade), dan settlement Perjalanan Dinas.",
        "Menjalankan approval dua tahap Tax Review -> Finance Review dengan matrix konfigurable.",
        "Validasi budget per divisi dan aturan PR wajib untuk pengajuan > Rp 1.000.000.",
        "Integrasi dengan SAP (push GL/journal/MIRO, terima payment plan & status) dan HRIS "
        "(read travel) melalui kontrak yang disepakati.",
        "Menyediakan SSO (Keycloak) untuk login karyawan & vendor.",
        "Odoo tidak melakukan posting GL — pembukuan tetap di SAP.",
    ])

    # ---------------- 3. Success Criteria ----------------
    b.h("3. Kriteria Sukses (Success Criteria / KPI)", 1)
    b.table(
        ["KPI", "Target"],
        [
            ["Modul Odoo terinstal & lulus uji (unit + SIT)", "100% modul in-scope"],
            ["End-to-end happy path (submit -> Tax -> Finance -> push SAP -> status mirror)", "Lulus di SIT & UAT"],
            ["UAT sign-off oleh user Finance/Tax", "Tercapai sebelum Go-Live"],
            ["Login SSO (karyawan & vendor) + role mapping benar", "100% skenario"],
            ["Defect Sev-1/Sev-2 saat Go-Live", "0 open"],
            ["Go-Live sesuai timeline", "±W21"],
        ],
        widths=[11, 5],
    )

    # ---------------- 4. Scope ----------------
    b.h("4. Ruang Lingkup (Scope)", 1)
    b.h("4.1 In-Scope (Odoo)", 2)
    b.bullets([
        "Modul: custom_finance_portal, custom_finance_budget, custom_finance_portal_sap "
        "(adapter sisi Odoo), custom_finance_portal_sso.",
        "Dokumen + workflow approval Tax->Finance, budget control, aturan PR > Rp 1jt.",
        "Adapter Odoo ke kontrak REST integrasi, upsert master idempoten, webhook status (HMAC), "
        "push job async, sync log + menu.",
        "Vendor Portal (login SSO, submit invoice, tracking), SSO Keycloak (config sisi Odoo).",
        "Reporting (login/transaction/sync log), dashboard, konfigurasi, security/PDP audit, hardening.",
    ])
    b.h("4.2 Out-of-Scope (DIKELUARKAN)", 2)
    b.table(
        ["Area", "Penjelasan", "Pemilik"],
        [
            ["SAP development", "ABAP/CPI/OData, posting GL/journal/MIRO, expose PR/PO/GR non-trade + "
             "nilai/status, payment list, attachment SAP Basis, approval status PO", "Tim SAP klien"],
            ["HC / HRIS development", "Konektor HRIS->Kafka, perubahan modul travel, ekstraksi "
             "employee master", "Tim HC/HRIS klien"],
            ["Kafka & konektor", "Cluster Kafka, topik, konektor SAP<->Kafka & Kafka<->Portal", "Tim integrasi/Kafka klien"],
            ["Infrastruktur non-Odoo", "Provisioning Keycloak server, jaringan, sertifikat", "Tim infra klien"],
        ],
        widths=[3.5, 9, 3.5],
    )
    b.p("Sisi Odoo build terhadap kontrak (JSON + topik) sehingga dapat dikerjakan paralel.", italic=True)

    # ---------------- 5. Deliverables ----------------
    b.h("5. Deliverables Utama", 1)
    b.bullets([
        "4 modul Odoo terinstal & teruji.",
        "Konfigurasi SSO Keycloak (sisi Odoo) + role mapping.",
        "Dokumentasi: FSD, TSD, kontrak integrasi (JSON + topik), runbook deploy, user guide.",
        "Laporan UAT & sign-off, rencana cutover, laporan hypercare.",
    ])

    # ---------------- 6. Milestone ----------------
    b.h("6. Milestone & Timeline (±24 minggu / ±6 bulan)", 1)
    b.table(
        ["#", "Milestone", "Target Minggu"],
        [
            ["M1", "Requirement & kontrak integrasi final", "W3"],
            ["M2", "Design (FSD/TSD) sign-off", "W5"],
            ["M3", "Feature complete (Build)", "W15"],
            ["M4", "SIT pass (gated: konektor SAP/Kafka/HRIS ready)", "W17"],
            ["M5", "UAT sign-off", "W20"],
            ["M6", "Go-Live", "W21"],
            ["M7", "Hypercare selesai & handover", "W24"],
        ],
        widths=[1.5, 11, 3.5],
    )

    # ---------------- 7. Effort ----------------
    b.h("7. Estimasi Effort & Resource (Mandays — Odoo Only)", 1)
    b.table(
        ["Fase", "PMO", "BA", "Dev", "QA", "Total"],
        [
            ["1. Requirement Gathering & Analysis", "8", "22", "3", "4", "37"],
            ["2. Design (FSD/TSD)", "3", "8", "10", "6", "27"],
            ["3. Build / Development", "18", "35", "123", "44", "220"],
            ["4. SIT", "5", "6", "12", "16", "39"],
            ["5. UAT", "5", "12", "12", "8", "37"],
            ["6. Cutover & Go-Live", "4", "5", "8", "3", "20"],
            ["7. Post Go-Live / Hypercare", "4", "5", "9", "3", "21"],
            ["Subtotal", "47", "93", "177", "84", "401"],
            ["Kontingensi 15%", "7", "14", "27", "13", "61"],
            ["TOTAL", "54", "107", "204", "97", "≈462"],
        ],
        widths=[7, 1.8, 1.8, 1.8, 1.8, 2],
        total_rows={7, 9},
    )
    b.p("Komposisi tim: 1 PMO (part-time), 1–2 BA, 2–3 Developer, 1–2 QA.")

    b.h("7.1 Detail Build per Modul (Dev / BA / QA)", 2)
    b.table(
        ["Workstream (Odoo)", "Dev", "BA", "QA"],
        [
            ["Foundation (SSO Odoo-side, base module, security, CI/CD)", "12", "2", "3"],
            ["Master data + sync upsert (Odoo-side) + cron", "14", "6", "4"],
            ["Cash Advance + Realization (+ approval Tax->Finance)", "14", "4", "5"],
            ["Reimbursement & Expenses", "9", "3", "4"],
            ["Vendor Invoice (PO/Non-PO Non-Trade) + Vendor Portal SSO", "18", "4", "6"],
            ["Perjalanan Dinas (read HRIS via kontrak + settlement)", "7", "3", "3"],
            ["Budget control + aturan PR > Rp 1jt", "8", "3", "3"],
            ["Integration Odoo-side (adapter + push job + webhook + sync log)", "16", "4", "6"],
            ["Reporting (login/transaction/sync log) + Dashboards", "12", "3", "4"],
            ["Configuration + Limitation master", "4", "2", "2"],
            ["Non-functional (PDP audit, security, perf, resilience)", "9", "1", "4"],
            ["Subtotal Build", "123", "35", "44"],
        ],
        widths=[10, 2, 2, 2],
        total_rows={11},
    )

    # ---------------- 8. Governance ----------------
    b.h("8. Organisasi Proyek & Governance", 1)
    b.h("8.1 Stakeholder & Peran", 2)
    b.table(
        ["Peran", "Tanggung Jawab"],
        [
            ["Project Sponsor (klien)", "Mandat, keputusan strategis, sign-off"],
            ["Steering Committee", "Arahan, eskalasi, keputusan scope mayor"],
            ["PMO", "Perencanaan, monitoring, reporting, change control, koordinasi lintas tim"],
            ["Business Analyst (BA)", "Requirement, FSD, mapping master/approval, fasilitasi UAT"],
            ["Developer", "Pengembangan modul Odoo + adapter integrasi sisi Odoo"],
            ["QA", "Test design, SIT, regression, otomasi"],
            ["Product Owner (klien)", "Prioritas backlog, validasi fungsional, UAT"],
            ["Tim SAP/Kafka/HC (klien)", "Konektor & perubahan SAP/Kafka/HRIS (di luar scope)"],
        ],
        widths=[5, 11],
    )
    b.h("8.2 RACI (ringkas)", 2)
    b.table(
        ["Aktivitas", "PMO", "BA", "Dev", "QA", "Sponsor/PO"],
        [
            ["Requirement & FSD", "A", "R", "C", "C", "C/A"],
            ["Design / TSD", "C", "C", "R", "C", "A"],
            ["Build modul Odoo", "A", "C", "R", "C", "I"],
            ["Kontrak integrasi", "A", "R", "C", "I", "C"],
            ["SIT", "A", "C", "C", "R", "I"],
            ["UAT", "C", "A", "C", "C", "R"],
            ["Go-Live", "A", "C", "R", "C", "A"],
        ],
        widths=[5.5, 2, 2, 2, 2, 2.5],
    )
    b.p("R = Responsible, A = Accountable, C = Consulted, I = Informed.", italic=True, size=9)
    b.h("8.3 Cadence & Change Control", 2)
    b.bullets([
        "Sprint 2-mingguan; demo akhir sprint; stand-up harian tim delivery.",
        "Steering Committee bulanan; status report mingguan oleh PMO.",
        "Change control: setiap perubahan scope/kontrak via formulir CR, disetujui PMO + Sponsor.",
        "Eskalasi: blocker > 2 hari kerja dieskalasi ke Steering.",
    ])

    # ---------------- 9-11 ----------------
    b.h("9. Asumsi", 1)
    b.bullets([
        "1 manday = 1 orang-hari (±20 hari/bulan); mandays ≠ durasi kalender.",
        "Reuse komponen Odoo (approval engine, adapter framework, queue_job, PDP audit, reporting).",
        "Kontrak integrasi (JSON + topik) final di fase Requirement; tim SAP/Kafka/HC menyediakannya.",
        "Konektor SAP/Kafka/HRIS siap sebelum SIT (±W13).",
        "Lingkungan dev/staging/prod Odoo + Keycloak disediakan klien tepat waktu.",
        "Master data SAP berkualitas memadai.",
    ])
    b.h("10. Batasan (Constraints)", 1)
    b.bullets([
        "Odoo tidak memposting GL — pembukuan tetap di SAP.",
        "Integrasi hanya via kontrak REST/Kafka yang disepakati (tidak akses langsung DB SAP).",
        "Kepatuhan UU PDP (audit trail, masking data sensitif: NIK, rekening).",
        "Timeline bergantung pada kesiapan dependency eksternal.",
    ])
    b.h("11. Dependensi", 1)
    b.bullets([
        "Konektor & endpoint SAP/Kafka/HRIS (eksternal).",
        "Keycloak realm/client + role mapper (infra klien).",
        "Akses & approval data master dari SAP/HRIS.",
    ])

    # ---------------- 12. Risk ----------------
    b.h("12. Risiko Tingkat Tinggi (RAID)", 1)
    b.table(
        ["#", "Risiko", "Dampak", "Mitigasi"],
        [
            ["1", "Konektor SAP/Kafka/HRIS belum siap saat SIT", "Geser SIT->Go-Live", "Contract-first + stub/mock; SIT gated milestone"],
            ["2", "Kualitas master data SAP", "Rework mapping", "Cleansing dini di Build"],
            ["3", "Attachment SAP Basis & approval status PO belum ada", "Potensi CR", "Tandai sebagai opsi/CR di awal"],
            ["4", "Ketersediaan user UAT", "UAT molor", "Jadwalkan slot UAT sejak Design"],
            ["5", "Perubahan kontrak integrasi", "Effort bertambah", "Change control via PMO"],
        ],
        widths=[1, 5.5, 4, 5.5],
    )

    # ---------------- 13. Acceptance ----------------
    b.h("13. Kriteria Penerimaan (Acceptance)", 1)
    b.bullets([
        "Seluruh deliverable in-scope diterima & sign-off.",
        "KPI Bab 3 tercapai; 0 defect Sev-1/Sev-2 terbuka saat Go-Live.",
        "UAT sign-off oleh Product Owner & user Finance/Tax.",
        "Handover dokumentasi & runbook ke tim support.",
    ])

    # ---------------- 14. Sign-off ----------------
    b.h("14. Persetujuan (Sign-off)", 1)
    b.table(
        ["Peran", "Nama", "Tanda Tangan", "Tanggal"],
        [
            ["Project Sponsor", "", "", ""],
            ["Steering Committee", "", "", ""],
            ["Project Manager / PMO", "", "", ""],
            ["Product Owner (klien)", "", "", ""],
            ["Delivery Lead", "", "", ""],
        ],
        widths=[5, 5, 4, 2.5],
    )

    out = os.path.join(HERE, "Project Charter - Finance Portal Odoo v1.0.docx")
    b.save(out)
    print("WROTE", out)


if __name__ == "__main__":
    build()
