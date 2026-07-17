# -*- coding: utf-8 -*-
"""Generate two Levi's Retail documents for PT Era Busana Retailindo (EBR) as DOCX:

  1. FSD       — Functional Specification Document (EBR-branded)
  2. UAT Detail — User Acceptance Testing detail dengan ~45 test case ter-traceable

Usage:  python docs/projects/levis/build_fsd_uat_ebr_docx.py
Output: docs/projects/levis/FSD - ERA Busana Retailindo (Levi's) v1.0.docx
        docs/projects/levis/UAT Detail - ERA Busana Retailindo (Levi's) v1.0.docx

Companion of build_fsd_docx.py and build_tsd_charter_workflow_docx.py.
Semua paragraf body di-justify; heading Odoo purple; header tabel navy.
Test case di-grounding pada modul addons/ee_gap/custom_retail_import (label tombol,
menu, state log, dan gating Phase-5 sesuai implementasi nyata).
"""

import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
INK = RGBColor(0x1A, 0x23, 0x32)
ACCENT = RGBColor(0x71, 0x4B, 0x67)  # Odoo purple
HEAD_BG = "1A2332"
CASE_BG = "714B67"
LABEL_BG = "EDE7EC"
TODAY = "22 Juni 2026"


class Builder:
    def __init__(self):
        self.doc = Document()
        normal = self.doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)
        normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        normal.paragraph_format.space_after = Pt(6)
        for level, size in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)):
            st = self.doc.styles[level]
            st.font.name = "Calibri"
            st.font.size = Pt(size)
            st.font.color.rgb = ACCENT if level == "Heading 1" else INK
            st.font.bold = True
        for section in self.doc.sections:
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(2.3)
            section.right_margin = Cm(2.3)

    # -- primitives ----------------------------------------------------------
    def p(self, text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=None, color=None, space_after=6):
        par = self.doc.add_paragraph()
        par.alignment = align
        par.paragraph_format.space_after = Pt(space_after)
        run = par.add_run(text)
        run.bold = bold
        run.italic = italic
        if size:
            run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
        return par

    def h(self, text, level):
        self.doc.add_heading(text, level=level)

    def bullets(self, items):
        for it in items:
            par = self.doc.add_paragraph(style="List Bullet")
            par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if isinstance(it, tuple):
                r = par.add_run(it[0])
                r.bold = True
                par.add_run(it[1])
            else:
                par.add_run(it)

    @staticmethod
    def _shade(cell, hexcolor):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), hexcolor)
        tcPr.append(shd)

    def table(self, headers, rows, widths=None, fs=9.5):
        t = self.doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = t.rows[0].cells
        for i, htxt in enumerate(headers):
            hdr[i].text = ""
            run = hdr[i].paragraphs[0].add_run(htxt)
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(fs)
            self._shade(hdr[i], HEAD_BG)
        for row in rows:
            cells = t.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = ""
                par = cells[i].paragraphs[0]
                par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                run = par.add_run(str(val))
                run.font.size = Pt(fs)
        if widths:
            for i, w in enumerate(widths):
                for row in t.rows:
                    row.cells[i].width = Cm(w)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)
        return t

    # -- UAT test-case block -------------------------------------------------
    @staticmethod
    def _cell_lines(cell, value, fs=9.5, bold=False):
        cell.text = ""
        if isinstance(value, (list, tuple)):
            lines = list(value)
        else:
            lines = [value]
        for idx, ln in enumerate(lines):
            par = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
            par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            par.paragraph_format.space_after = Pt(1)
            run = par.add_run(str(ln))
            run.font.size = Pt(fs)
            run.bold = bold

    def testcase(self, tc):
        t = self.doc.add_table(rows=1, cols=2)
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        # header (merged)
        hc = t.rows[0].cells
        head = hc[0].merge(hc[1])
        head.text = ""
        run = head.paragraphs[0].add_run(f"{tc['id']}  —  {tc['title']}")
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        self._shade(head, CASE_BG)

        def addrow(label, value, bold_val=False):
            cells = t.add_row().cells
            cells[0].text = ""
            lp = cells[0].paragraphs[0]
            lr = lp.add_run(label)
            lr.bold = True
            lr.font.size = Pt(9)
            self._shade(cells[0], LABEL_BG)
            self._cell_lines(cells[1], value, bold=bold_val)
            cells[0].width = Cm(3.7)
            cells[1].width = Cm(12.8)

        addrow("Referensi FR", tc.get("fr", "-"))
        addrow("Prioritas", tc.get("prio", "Sedang"))
        addrow("Prakondisi", tc.get("pre", "-"))
        addrow("Langkah Pengujian", tc.get("steps", "-"))
        addrow("Data Uji", tc.get("data", "-"))
        addrow("Hasil yang Diharapkan", tc.get("expected", "-"))
        addrow("Hasil Aktual", tc.get("actual", ""))
        addrow("Status (Lulus / Gagal)", tc.get("status", ""))
        addrow("Penguji / Tanggal", "")
        addrow("Catatan / No. Defect", "")
        self.doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # -- shared cover/approval ----------------------------------------------
    def cover(self, kind, abbr, title, subtitle, number, status):
        for _ in range(5):
            self.doc.add_paragraph()
        self.p(kind, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=22, color=ACCENT, space_after=2)
        self.p(abbr, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, color=ACCENT)
        self.doc.add_paragraph()
        self.p(title, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
        self.p(subtitle, align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
        for _ in range(3):
            self.doc.add_paragraph()
        self.table(
            ["Atribut", "Keterangan"],
            [
                ["Judul Dokumen", title],
                ["Entitas", "PT Era Busana Retailindo (EBR) — operator retail Levi's"],
                ["Nomor Dokumen", number],
                ["Versi", "1.0"],
                ["Tanggal", TODAY],
                ["Klasifikasi", "Internal — Erajaya Group"],
                ["Status", status],
                ["Dokumen Terkait", "FSD/TSD/Charter/Business Workflow Levi's Retail (SES x EBR)"],
            ],
            widths=[5, 11],
        )
        self.doc.add_page_break()

    def approval(self, intro, rows):
        self.h("Lembar Persetujuan (Approval Sheet)", 1)
        self.p(intro)
        self.table(
            ["Peran", "Nama", "Jabatan / Fungsi", "Tanda Tangan", "Tanggal"], rows, widths=[4.4, 3.2, 4.6, 2.5, 2.0]
        )
        self.h("Riwayat Revisi", 2)
        self.table(
            ["Versi", "Tanggal", "Penyusun", "Deskripsi Perubahan"],
            [
                ["1.0", TODAY, "Tim Implementasi Erajaya", "Penerbitan awal untuk pengajuan persetujuan."],
            ],
            widths=[2, 3, 4, 7.2],
        )
        self.doc.add_page_break()

    def save(self, filename):
        path = os.path.join(HERE, filename)
        self.doc.save(path)
        print("Saved:", path)


SUBTITLE = "PT Sinar Eka Selaras (SES — SAP)  ×  PT Era Busana Retailindo (EBR — Odoo)  ×  POS Principal Levi's"

APPROVAL_INTRO = (
    "Dokumen ini diajukan untuk mendapatkan persetujuan. Dengan menandatangani lembar ini, "
    "para pihak menyatakan telah membaca, memahami, dan menyetujui isi dokumen. Perubahan "
    "setelah persetujuan dikelola melalui mekanisme change request."
)


# ============================================================================
# 1. FUNCTIONAL SPECIFICATION DOCUMENT (EBR-branded)
# ============================================================================
def build_fsd():
    b = Builder()
    b.cover(
        "FUNCTIONAL SPECIFICATION DOCUMENT",
        "(FSD)",
        "FSD — Implementasi Odoo Bisnis Retail Levi's PT Era Busana Retailindo",
        SUBTITLE,
        "FSD/ERAJAYA/EBR-LEVIS/2026/001",
        "Draft — Diajukan untuk Persetujuan (Approval Request)",
    )
    b.approval(
        APPROVAL_INTRO,
        [
            ["Disusun oleh (Prepared by)", "", "Functional Consultant — Erajaya", "", ""],
            ["Diperiksa oleh (Reviewed by)", "", "IT / ERP Lead — Erajaya", "", ""],
            ["Diperiksa oleh (Reviewed by)", "", "Finance & Accounting — PT EBR", "", ""],
            ["Disetujui oleh (Approved by)", "", "Business Owner — Levi's Retail", "", ""],
            ["Disetujui oleh (Approved by)", "", "Manajemen PT SES", "", ""],
        ],
    )

    # ---- 1. Pendahuluan ----
    b.h("1. Pendahuluan", 1)
    b.h("1.1 Latar Belakang", 2)
    b.p(
        "Erajaya Group melalui PT Sinar Eka Selaras (SES) bertindak sebagai distributor resmi "
        "produk Levi's di Indonesia, dengan seluruh proses distribusi dan pengadaan dijalankan "
        "pada sistem SAP. Operasional penjualan retail dijalankan oleh PT Era Busana "
        "Retailindo (EBR), sementara seluruh toko beroperasi menggunakan sistem POS bawaan "
        "principal yang menjadi standar global Levi's. PT EBR membutuhkan sistem pembukuan dan "
        "pelaporan yang andal namun tidak menggantikan sistem transaksional toko; untuk "
        "kebutuhan tersebut dipilih Odoo sebagai sistem ERP pencatatan (system of record "
        "akuntansi) PT EBR. Data operasional toko — goods receipt, penjualan POS beserta tender "
        "pembayarannya, dan mutasi persediaan — dimirror ke Odoo melalui file Excel/CSV pada "
        "fase awal dan melalui FTP/SFTP Erajaya pada fase berikutnya."
    )
    b.h("1.2 Tujuan Dokumen", 2)
    b.p(
        "Dokumen ini mendefinisikan kebutuhan fungsional implementasi Odoo untuk PT Era Busana "
        "Retailindo, mencakup proses bisnis end-to-end antara Levi's Principal, PT SES, toko "
        "retail, dan PT EBR; modul Odoo yang diinstal beserta kustomisasinya; spesifikasi file "
        "data yang dipertukarkan; serta konfigurasi data yang harus disiapkan sebelum go-live. "
        "Dokumen ini menjadi dasar persetujuan (approval) dan acuan penyusunan skenario User "
        "Acceptance Testing (UAT) sebelum tahap konfigurasi, pengembangan, dan migrasi data "
        "dilaksanakan."
    )
    b.h("1.3 Ruang Lingkup", 2)
    b.p("Ruang lingkup pekerjaan yang tercakup (in-scope):", bold=True)
    b.bullets(
        [
            "Pencatatan Purchase Order PT EBR kepada PT SES di Odoo, termasuk approval berjenjang, "
            "pencatatan vendor bill, dan pembayaran kepada SES.",
            "Mirroring data operasional toko dari POS principal ke Odoo: goods receipt toko, "
            "penjualan POS dan tender pembayaran, mutasi persediaan, posisi stok on-hand, serta "
            "data diskon/promosi.",
            "Pembangunan adapter import Excel/CSV (modul custom_retail_import) dengan profil per "
            "jenis file, pratinjau dry-run, deduplikasi, log audit, dan eksekusi asinkron.",
            "Penyediaan feed otomatis melalui FTP/SFTP Erajaya pada Phase 2.",
            "Penyiapan master data dan konfigurasi awal: company, chart of accounts, rekening bank "
            "dan jurnal, pajak, vendor, master produk, warehouse per toko, dan saldo awal stok.",
            "Pelaporan keuangan dan rekonsiliasi (penjualan, HPP, persediaan, settlement tender "
            "terhadap penerimaan bank).",
        ]
    )
    b.p("Hal-hal di luar ruang lingkup (out-of-scope):", bold=True)
    b.bullets(
        [
            "Proses transaksional di toko. Toko tetap menggunakan POS bawaan principal; Odoo tidak "
            "menggantikan fungsi POS toko.",
            "Proses internal SAP PT SES, termasuk pengadaan SES kepada Levi's Principal dan goods "
            "receipt di DC SES; proses tersebut dijelaskan hanya sebagai konteks alur bisnis.",
            "Integrasi API real-time dua arah antara SAP, POS principal, dan Odoo (di luar mekanisme "
            "file/SFTP yang didefinisikan).",
            "Perubahan proses bisnis pada sisi Levi's Principal.",
        ]
    )
    b.h("1.4 Definisi dan Singkatan", 2)
    b.table(
        ["Istilah", "Definisi"],
        [
            ["SES", "PT Sinar Eka Selaras — distributor produk Levi's (Erajaya Group), bersistem SAP."],
            ["EBR", "PT Era Busana Retailindo — entitas retail Levi's yang pembukuannya di Odoo."],
            ["Principal", "Levi's selaku pemilik merek dan pemasok produk."],
            ["POS Principal", "Sistem point-of-sale bawaan principal yang digunakan seluruh toko Levi's."],
            ["GR", "Goods Receipt — penerimaan barang."],
            ["Tender", "Metode pembayaran pada transaksi POS (tunai, EDC, e-wallet) beserta settlement-nya."],
            ["Mirroring", "Replikasi data transaksi dari sistem sumber (POS principal) ke Odoo melalui file."],
            ["UAT", "User Acceptance Testing — pengujian penerimaan oleh pengguna bisnis."],
            ["FSD / TSD", "Functional / Technical Specification Document."],
        ],
        widths=[3.2, 13.0],
    )
    b.h("1.5 Referensi", 2)
    b.bullets(
        [
            "Diagram proses bisnis: docs/projects/levis/business-process-flow.html.",
            "Dokumen pendukung: TSD, Project Charter, Business Workflow Levi's Retail (SES x EBR).",
            "Paket data principal/SES: X101 Material Master, X20 On-hand, X24DN Retail Sales, X31 "
            "Discount Journal, X32P Stock Movement, X70D Tender Detail, X70T Tender Settlement, "
            "Store Master, CoA EBR.",
            "Source code adapter: addons/ee_gap/custom_retail_import (Odoo 19).",
            "Runbook onboarding data: scripts/tenants/levis/README.md.",
        ]
    )
    b.doc.add_page_break()

    # ---- 2. Proses bisnis ----
    b.h("2. Gambaran Umum Proses Bisnis", 1)
    b.p(
        "Rantai proses melibatkan empat aktor dengan tiga sistem berbeda. Levi's Principal "
        "memasok produk dan menyediakan sistem POS toko. PT SES menjalankan distribusi dengan "
        "SAP sebagai sistem pencatatan. Toko retail beroperasi sepenuhnya pada POS bawaan "
        "principal. PT EBR menjalankan pembukuan dan pelaporan keuangan pada Odoo yang menerima "
        "mirroring data operasional toko. Terdapat tiga alur utama berikut."
    )
    b.h("2.1 Flow A — Pengadaan SES kepada Levi's Principal (SAP, Konteks)", 2)
    b.table(
        ["Langkah", "Aktivitas", "Sistem"],
        [
            ["A1", "SES membuat PO produk Levi's dan mengirimkannya kepada principal.", "SAP — SES"],
            ["A2", "Principal mengonfirmasi pesanan dan mengirim barang beserta dokumen.", "Principal"],
            ["A3", "Goods Receipt atas kiriman principal; stok bertambah di DC SES.", "SAP — SES"],
            ["A4", "Verifikasi invoice (three-way match) dan pembayaran kepada principal.", "SAP — SES"],
        ],
        widths=[2, 10.7, 3.5],
    )
    b.h("2.2 Flow B — Pembelian EBR kepada SES dan GR di Toko", 2)
    b.table(
        ["Langkah", "Aktivitas", "Sistem"],
        [
            ["B1", "EBR membuat PO pembelian produk Levi's kepada SES (intercompany B2B).", "Odoo — EBR"],
            ["B2", "SES menerima PO, membuat SO dan Delivery, mengirim barang ke toko.", "SAP — SES"],
            ["B3", "Toko menerima barang; Goods Receipt dilakukan pada POS principal.", "POS Principal"],
            ["B4", "Data GR toko dimirror ke Odoo untuk validasi receipt atas PO EBR.", "Mirroring"],
            ["B5", "Vendor bill SES dicatat, dicocokkan dengan PO dan GR, lalu dibayar.", "Odoo — EBR"],
        ],
        widths=[2, 10.7, 3.5],
    )
    b.h("2.3 Flow C — Operasional Toko dan Mirroring ke Odoo", 2)
    b.table(
        ["Langkah", "Aktivitas", "Sistem"],
        [
            ["C1", "Transaksi di POS: GR toko, penjualan & tender, mutasi stok, on-hand.", "POS Principal"],
            ["C2", "Ekspor report standar principal per periode (Excel/CSV).", "Export"],
            ["C3", "Transfer file: Phase 1 unggah manual; Phase 2 FTP/SFTP Erajaya.", "File / SFTP"],
            ["C4", "Adapter import memetakan report ke dokumen Odoo.", "Odoo — EBR"],
            ["C5", "Posting jurnal otomatis dan rekonsiliasi tender vs bank.", "Odoo — EBR"],
        ],
        widths=[2, 10.7, 3.5],
    )
    b.doc.add_page_break()

    # ---- 3. Kebutuhan fungsional ----
    b.h("3. Kebutuhan Fungsional", 1)
    b.p(
        "Kebutuhan fungsional berikut menjadi dasar konfigurasi, pengembangan, dan penyusunan "
        "skenario UAT. Prioritas dinyatakan sebagai M (Mandatory — wajib pada go-live) atau "
        "P2 (Phase 2 / menunggu keputusan). Kolom ID dipakai sebagai kunci traceability pada "
        "dokumen UAT."
    )
    b.table(
        ["ID", "Kebutuhan Fungsional", "Deskripsi", "Prioritas"],
        [
            [
                "FR-01",
                "Purchase Order EBR → SES",
                "Pengguna EBR membuat, mengubah, dan mengonfirmasi PO kepada SES dengan approval "
                "berjenjang sesuai matriks otorisasi (custom_approval_engine).",
                "M",
            ],
            [
                "FR-02",
                "Mirror GR toko ke Odoo",
                "GR di POS principal direplikasi ke Odoo sebagai receipt atas PO EBR, menambah stok "
                "warehouse toko dan menjadi dasar pencocokan vendor bill.",
                "M",
            ],
            [
                "FR-03",
                "Vendor bill & pembayaran SES",
                "Vendor bill SES dicatat dan dicocokkan terhadap PO dan GR (three-way match) sebelum "
                "pembayaran diproses.",
                "M",
            ],
            [
                "FR-04",
                "Mirror penjualan POS (X24DN)",
                "Detail penjualan harian per toko diimpor membentuk jurnal penjualan dan HPP. "
                "Representasi pos.order, kedalaman histori, dan tax mapping menunggu keputusan "
                "Phase 5; executor menahan posting hingga diaktifkan.",
                "P2",
            ],
            [
                "FR-05",
                "Mirror tender & settlement (X70D/X70T)",
                "Rincian metode pembayaran dan settlement diimpor sebagai payment register untuk "
                "rekonsiliasi terhadap penerimaan bank per acquirer.",
                "P2",
            ],
            [
                "FR-06",
                "Mirror mutasi & on-hand (X32P/X20)",
                "Mutasi persediaan bernilai direplikasi sebagai stock move; X20 untuk saldo awal "
                "(one-shot) dan validasi stok periodik.",
                "M",
            ],
            [
                "FR-07",
                "Import master data",
                "Master produk X101 (artikel, varian size/inseam, barcode, kategori, harga), CoA, "
                "company detail, dan store master diimpor melalui profil import.",
                "M",
            ],
            [
                "FR-08",
                "Wizard import dengan dry-run",
                "Setiap unggah menampilkan pratinjau hasil parsing sebelum commit; file besar diproses "
                "asinkron via queue_job.",
                "M",
            ],
            [
                "FR-09",
                "Deduplikasi & idempotensi",
                "File identik (SHA256) ditolak saat diimpor ulang; re-import record menghasilkan update "
                "via external ID, bukan duplikasi.",
                "M",
            ],
            [
                "FR-10",
                "Log audit import",
                "Setiap import tercatat (pengguna, waktu, jumlah baris, error) dan file sumber "
                "diarsipkan sebagai attachment untuk audit.",
                "M",
            ],
            [
                "FR-11",
                "SFTP feed otomatis",
                "Odoo polling terjadwal ke FTP/SFTP Erajaya per jenis file, memproses file baru melalui "
                "executor yang sama dengan unggah manual, dan memberi notifikasi bila parsing gagal.",
                "P2",
            ],
            [
                "FR-12",
                "Klaim promosi (X31)",
                "Data diskon/promo diimpor sebagai dasar klaim promosi kepada principal/SES.",
                "P2",
            ],
            [
                "FR-13",
                "Pelaporan keuangan",
                "Laporan keuangan lengkap (Excel/PDF) via custom_accounting_full; dokumen transaksi "
                "(Invoice, Quotation, PO) ber-branding via custom_report_templates.",
                "M",
            ],
            [
                "FR-14",
                "Pajak Indonesia",
                "PPN keluaran/masukan dikonfigurasi sesuai tarif; integrasi Coretax/e-Faktur dan bukti "
                "potong via custom_coretax dan custom_coretax_bupot.",
                "M",
            ],
        ],
        widths=[1.6, 4.0, 8.6, 2.0],
    )
    b.doc.add_page_break()

    # ---- 4. Modul ----
    b.h("4. Modul Odoo yang Diinstal dan Kustomisasi", 1)
    b.p(
        'Tenant EBR diprovision menggunakan Industry Pack "Retail / POS" pada platform Odoo '
        "Erajaya (deployment satu aksi). Pack memasang modul Odoo Community beserta modul custom "
        "platform; di atasnya dipasang modul khusus kebutuhan Levi's."
    )
    b.h("4.1 Modul Standar Odoo (Community)", 2)
    b.table(
        ["Modul", "Fungsi pada Implementasi Ini"],
        [
            ["purchase", "Purchase Order EBR → SES; pencocokan vendor bill (PO–GR–Invoice)."],
            ["stock (Inventory)", "Warehouse per toko, receipt (mirror GR), stock move, valuasi persediaan."],
            ["account (Accounting)", "CoA, jurnal, AP/AR, rekonsiliasi bank."],
            ["point_of_sale", "Representasi transaksi POS hasil mirroring (bukan POS transaksional toko)."],
            ["sale_management", "Sales order dan pricelist."],
        ],
        widths=[4.5, 11.7],
    )
    b.h("4.2 Modul Custom Platform (Bawaan Pack Retail)", 2)
    b.table(
        ["Modul", "Fungsi"],
        [
            ["custom_accounting_full", "Pelengkap akuntansi CE→EE: laporan keuangan lengkap (Excel/PDF)."],
            ["custom_coretax / custom_coretax_bupot", "Integrasi pajak Indonesia: e-Faktur/Coretax; bukti potong."],
            ["custom_pos_id", "Lokalisasi POS Indonesia (format struk, pembayaran lokal)."],
            ["custom_approval_engine", "Approval berjenjang untuk PO dan dokumen lainnya."],
            ["custom_report_templates", "Template PDF ber-branding per company: Invoice, Quotation, PO."],
        ],
        widths=[5.2, 11],
    )
    b.h("4.3 Modul Custom Khusus Levi's — custom_retail_import", 2)
    b.p(
        "Adapter mirroring Excel/CSV → Odoo dan SFTP feed; komponen inti integrasi Flow C, "
        "dirancang reusable untuk tenant retail lain. Kemampuan utamanya:"
    )
    b.bullets(
        [
            (
                "Import profile per jenis file. ",
                "Setiap report (X101, X20, X24, X70D, CoA, Company, Store Master) memiliki profil "
                "tersendiri: format, baris awal data, mapping kolom (JSON), aturan parsing "
                "tanggal/desimal, serta perbaikan karakter rusak (U+FFFD → ®).",
            ),
            (
                "Upload wizard dengan dry-run preview. ",
                "Pengguna mengunggah file lalu melihat pratinjau parsing sebelum commit.",
            ),
            (
                "Deduplikasi dan idempotensi. ",
                "Hash SHA256 per file menolak file identik; external ID per record menjadikan re-import "
                "bersifat update, bukan duplikasi.",
            ),
            (
                "Log audit dan arsip file. ",
                "Setiap import tercatat lengkap dan file sumber disimpan sebagai attachment.",
            ),
            (
                "SFTP feed poller dan eksekusi asinkron. ",
                "Cron terjadwal menarik file dari FTP/SFTP Erajaya; file besar (mis. X101 ±159 ribu SKU) "
                "diproses di background via queue_job.",
            ),
            (
                "Phase-5 gating. ",
                "Executor X24 (penjualan) dan X70D (tender) mampu memparsing dan mengelompokkan data "
                "namun menahan posting hingga keputusan bisnis Phase 5 ditetapkan.",
            ),
        ]
    )
    b.doc.add_page_break()

    # ---- 5. Spesifikasi file ----
    b.h("5. Spesifikasi File Feed Mirroring", 1)
    b.table(
        ["Report / File", "Isi Data", "Tujuan di Odoo", "Frekuensi"],
        [
            [
                "X101_Material_Master",
                "Master produk Levi's (artikel, SKU, barcode, harga).",
                "Product master (create/update produk).",
                "Saat ada perubahan",
            ],
            [
                "X20_Current_Onhand_Inventory",
                "Posisi stok on-hand per toko.",
                "Saldo awal (one-shot) dan validasi stok.",
                "Harian / periodik",
            ],
            [
                "X24DN_Retail_Sales_Detail",
                "Detail transaksi penjualan POS per toko.",
                "Penjualan (jurnal penjualan + HPP).",
                "Harian (EOD)",
            ],
            [
                "X70D / X70T (Tender)",
                "Metode pembayaran dan settlement.",
                "Payment register & rekonsiliasi bank.",
                "Harian (EOD)",
            ],
            [
                "X32P_Stock_Movement",
                "Mutasi persediaan bernilai (GR, transfer, retur, adjustment).",
                "Stock move dan receipt.",
                "Harian / periodik",
            ],
            ["X31_Discount_Journal", "Data diskon/promo.", "Dasar klaim promosi.", "Bulanan / per program"],
            [
                "Store Master",
                "Daftar 24 lokasi (1 DC + 23 toko).",
                "Warehouse per toko + crosswalk kode toko.",
                "Saat ada perubahan",
            ],
            ["CoA EBR", "Bagan akun (code, name, account_type).", "Chart of Accounts PT EBR.", "Satu kali"],
        ],
        widths=[4.0, 5.0, 4.7, 2.5],
    )

    # ---- 6. Konfigurasi data ----
    b.h("6. Konfigurasi Data dan Urutan Setup", 1)
    b.p(
        "Penyiapan data tenant EBR mengikuti urutan: konfigurasi dasar (company, CoA, bank, "
        "pajak) → master data → saldo awal → transaksi harian via mirroring. Setiap langkah "
        "merupakan prasyarat bagi langkah setelahnya, dan menjadi prakondisi suite UAT terkait."
    )
    b.table(
        ["No", "Data", "Sumber", "Cara Load"],
        [
            [
                "1",
                "Company Detail (EBR)",
                "Sheet legal entity (key/value).",
                "Profil levis_company atau Settings ▸ Companies.",
            ],
            ["2", "Chart of Accounts", "CoA EBR.xlsx.", "Profil levis_coa atau Accounting ▸ CoA ▸ Import."],
            ["3", "Bank Account & Journals", "Data Finance EBR.", "Manual (Accounting ▸ Configuration)."],
            ["4", "Pajak (PPN)", "Konfigurasi tax + Coretax.", "custom_coretax + konfigurasi tax."],
            ["5", "Vendor PT SES", "Data Procurement.", "Manual / import partner."],
            ["6", "Master Produk", "X101 (±14.885 artikel / ±159.658 SKU).", "Profil levis_x101 (asinkron ±30 menit)."],
            ["7", "Store Master → Warehouse", "Store Master (24 lokasi).", "Loader Track A / profil store_master."],
            ["8", "Saldo Awal Stok", "X20 on-hand cut-over.", "Profil levis_x20 (one-shot; prasyarat 6 & 7)."],
            ["9", "Transaksi Harian", "X24/X70D/X32P/X31.", "Wizard (Phase 1) → SFTP feed (Phase 2)."],
        ],
        widths=[1.0, 3.4, 5.0, 6.8],
    )

    # ---- 7. Roadmap ----
    b.h("7. Roadmap Integrasi", 1)
    b.h("7.1 Phase 1 — Mirroring Manual Excel/CSV", 2)
    b.p(
        "Tim ops mengekspor report principal lalu mengunggahnya ke Odoo melalui wizard import "
        "dengan mapping kolom pada profil. Validasi (jumlah baris, total nilai, exception "
        "report) dilakukan sebelum posting. Sesuai untuk masa awal pembukaan toko dan "
        "stabilisasi master data."
    )
    b.h("7.2 Phase 2 — FTP/SFTP Erajaya", 2)
    b.p(
        "File report di-drop otomatis ke server FTP/SFTP Erajaya sesuai jadwal. Odoo polling "
        "terjadwal — mengambil, parsing, import, lalu mengarsipkan ke folder "
        "incoming/processed/error — dan mengirim notifikasi bila file gagal diparse atau ada "
        "selisih rekonsiliasi. Menghilangkan proses manual sehingga data toko mutakhir H+0–H+1."
    )

    # ---- 8. Asumsi & open items ----
    b.h("8. Asumsi, Dependensi, dan Isu Terbuka", 1)
    b.h("8.1 Asumsi", 2)
    b.bullets(
        [
            "Odoo PT EBR berperan sebagai sistem pembukuan, bukan sistem transaksional toko; sumber "
            "transaksi tetap POS principal dan SAP SES.",
            "Struktur report principal stabil; perubahan ditangani melalui penyesuaian profil import.",
            "Kode toko SAP menjadi kunci penghubung seluruh file X terhadap warehouse Odoo.",
            "Akses server FTP/SFTP Erajaya disediakan sebelum Phase 2 dimulai.",
        ]
    )
    b.h("8.2 Isu Terbuka (Sebelum Go-Live Penuh)", 2)
    b.table(
        ["No", "Isu", "Dampak", "Tindak Lanjut"],
        [
            [
                "1",
                "File X20/X24/X70D masih sampel satu toko (store 14694 — Tunjungan Plaza 3).",
                "Saldo awal & penjualan hanya untuk satu toko.",
                "Minta ekstrak lengkap seluruh toko atau aktifkan feed SFTP.",
            ],
            [
                "2",
                "Store Master tanpa kode toko untuk 23 dari 24 lokasi.",
                "Crosswalk kode toko → warehouse tidak lengkap.",
                "Minta store master lengkap dengan kode SAP; loader idempoten.",
            ],
            [
                "3",
                "Keputusan Phase 5 (representasi POS, histori, tax mapping) belum ditetapkan.",
                "Posting X24/X70D tertahan gating executor.",
                "Workshop keputusan Finance EBR + tim pajak; gating dibuka setelahnya.",
            ],
        ],
        widths=[1.0, 5.4, 4.6, 5.2],
    )
    b.p(
        "Dengan disetujuinya dokumen ini, tim implementasi melanjutkan ke konfigurasi, migrasi "
        "data, UAT, dan persiapan go-live sesuai ruang lingkup dan roadmap. Hal di luar ruang "
        "lingkup dikelola melalui change request."
    )
    b.save("FSD - ERA Busana Retailindo (Levi's) v1.0.docx")


# ============================================================================
# 2. UAT DETAIL DOCUMENT
# ============================================================================
def build_uat():
    b = Builder()
    b.cover(
        "USER ACCEPTANCE TESTING",
        "(UAT Detail)",
        "UAT Detail — Implementasi Odoo Bisnis Retail Levi's PT Era Busana Retailindo",
        SUBTITLE,
        "UAT/ERAJAYA/EBR-LEVIS/2026/005",
        "Draft — Siap Eksekusi UAT (Test Execution Ready)",
    )
    b.approval(
        "Dokumen ini berisi rencana dan skenario rinci User Acceptance Testing (UAT) atas "
        "implementasi Odoo PT Era Busana Retailindo. Dengan menandatangani lembar persetujuan "
        "hasil UAT (Bab 9), para pihak menyatakan bahwa sistem telah diuji sesuai skenario dan "
        "memenuhi kriteria penerimaan, atau mencatat defect/keberatan yang harus diselesaikan "
        "sebelum go-live.",
        [
            ["Disusun oleh (Prepared by)", "", "Functional Consultant — Erajaya", "", ""],
            ["Pelaksana UAT (UAT Lead)", "", "Business Analyst — Erajaya", "", ""],
            ["Penguji (Tester)", "", "Finance & Accounting — PT EBR", "", ""],
            ["Penguji (Tester)", "", "Operasional / Master Data — PT EBR", "", ""],
            ["Disetujui oleh (Sign-off)", "", "Business Owner — Levi's Retail", "", ""],
        ],
    )

    # ---- 1. Pendahuluan ----
    b.h("1. Pendahuluan", 1)
    b.h("1.1 Tujuan", 2)
    b.p(
        "Dokumen ini mendefinisikan pendekatan, kriteria, dan skenario rinci User Acceptance "
        "Testing (UAT) untuk memastikan implementasi Odoo PT Era Busana Retailindo memenuhi "
        "kebutuhan fungsional yang ditetapkan pada FSD (FSD/ERAJAYA/EBR-LEVIS/2026/001) "
        "sebelum go-live. UAT dijalankan oleh pengguna bisnis EBR dengan pendampingan tim "
        "implementasi Erajaya, dan hasilnya menjadi dasar keputusan penerimaan (acceptance) "
        "sistem."
    )
    b.h("1.2 Ruang Lingkup Pengujian", 2)
    b.p("In-scope UAT:", bold=True)
    b.bullets(
        [
            "Import master data (company, CoA, produk X101, store master).",
            "Saldo awal stok (X20) dan kontrol one-shot.",
            "Kontrol adapter import: dry-run preview, deduplikasi, force re-import, audit log, eksekusi asinkron.",
            "Procurement EBR → SES: PO, approval berjenjang, GR mirror, vendor bill three-way match, dan pembayaran.",
            "Mirroring operasional toko (Flow C): X24 sales, X70D tender, X32P movement, X31 "
            "discount — termasuk verifikasi gating Phase 5.",
            "SFTP feed (Phase 2): test connection, poll, dedup skip, cron.",
            "Pelaporan keuangan, template PDF ber-branding, dan rekonsiliasi tender.",
            "Pajak (PPN, Coretax/e-Faktur) dan keamanan/hak akses.",
        ]
    )
    b.p("Out-of-scope UAT:", bold=True)
    b.bullets(
        [
            "Pengujian POS transaksional toko (tetap pada POS principal).",
            "Pengujian proses internal SAP SES dan proses Levi's Principal.",
            "Performance/stress test di luar verifikasi durasi import X101.",
            "Posting penjualan/tender produksi penuh sebelum keputusan Phase 5 (diuji sebagai "
            "verifikasi gating, bukan posting).",
        ]
    )
    b.h("1.3 Referensi", 2)
    b.bullets(
        [
            "FSD — FSD/ERAJAYA/EBR-LEVIS/2026/001 v1.0 (sumber kebutuhan fungsional FR-01..FR-14).",
            "TSD — spesifikasi teknis custom_retail_import.",
            "Business Workflow — alur operasional dan exception (E1–E7).",
            "Modul: addons/ee_gap/custom_retail_import (label menu, tombol, state log).",
        ]
    )
    b.doc.add_page_break()

    # ---- 2. Pendekatan UAT ----
    b.h("2. Pendekatan dan Metodologi UAT", 1)
    b.h("2.1 Lingkungan Pengujian", 2)
    b.p(
        "UAT dijalankan pada database tenant Odoo terpisah (mis. levis_uat) yang merupakan "
        "salinan konfigurasi produksi tanpa data live, di-provision dengan Industry Pack "
        "Retail / POS dan modul custom_retail_import terpasang beserta 6 profil Levi's ter-seed. "
        "Akses melalui browser ke URL tenant UAT; tiap penguji memakai akun dengan grup hak "
        "akses yang sesuai perannya."
    )
    b.h("2.2 Data Uji", 2)
    b.table(
        ["Data Uji", "Sumber", "Catatan"],
        [
            ["Master produk", "X101_Material_Master.xlsx (full).", "±14.885 artikel / ±159.658 SKU."],
            ["Chart of Accounts", "CoA EBR.xlsx.", "Lengkap & clean."],
            ["Company legal entity", "Sheet key/value EBR.", "Acuan format: PT SES.xlsx."],
            ["On-hand awal", "X20 (sampel store 14694).", "Single-store sample (isu terbuka #1)."],
            ["Penjualan / tender", "X24DN, X70D/X70T (sampel store 14694).", "Untuk uji parsing & gating."],
            ["Mutasi / diskon", "X32P, X31.", "Untuk uji staged parse/count."],
            ["Store master", "LEVI'S - STORE MASTER DATA.xlsx.", "Kode toko 23/24 belum ada (isu #2)."],
        ],
        widths=[3.6, 5.6, 7.0],
    )
    b.h("2.3 Peran dan Tanggung Jawab", 2)
    b.table(
        ["Peran", "Pihak", "Tanggung Jawab pada UAT"],
        [
            ["UAT Lead", "Business Analyst — Erajaya", "Koordinasi jadwal, pencatatan hasil, triase defect."],
            ["Functional Consultant", "Erajaya", "Menyiapkan lingkungan & data uji, dukungan teknis."],
            ["Tester Finance", "Finance & Accounting — PT EBR", "Uji procurement, akuntansi, pelaporan, pajak."],
            ["Tester Ops/Master Data", "Operasional — PT EBR", "Uji import master, saldo awal, mirroring."],
            ["IT / Developer", "Erajaya", "Perbaikan defect, dukungan SFTP/queue_job."],
            ["Approver / Sign-off", "Business Owner Levi's Retail", "Keputusan penerimaan akhir."],
        ],
        widths=[3.6, 4.6, 8.0],
    )
    b.h("2.4 Kriteria Masuk (Entry Criteria)", 2)
    b.bullets(
        [
            "FSD & TSD telah disetujui; lingkungan UAT siap dan dapat diakses.",
            "Modul custom_retail_import terpasang; 6 profil Levi's ter-seed dan terverifikasi.",
            "Data uji tersedia sesuai Bab 2.2; akun penguji & hak akses telah dibuat.",
            "Smoke test (UAT-ENV) lolos sebelum eksekusi suite fungsional.",
        ]
    )
    b.h("2.5 Kriteria Keluar (Exit Criteria)", 2)
    b.bullets(
        [
            "100% test case prioritas Tinggi berstatus Lulus (tidak ada yang Gagal terbuka).",
            "Tidak ada defect Critical/High yang masih terbuka; defect Medium/Low memiliki rencana "
            "penyelesaian yang disepakati.",
            "Seluruh test case telah dieksekusi dan hasilnya tercatat (Lulus/Gagal/Blocked).",
            "Lembar sign-off UAT (Bab 9) ditandatangani oleh approver.",
        ]
    )
    b.h("2.6 Klasifikasi Severity dan Prioritas Defect", 2)
    b.table(
        ["Severity", "Definisi", "Target Penyelesaian"],
        [
            [
                "Critical",
                "Fungsi inti gagal/blok go-live; tidak ada workaround (mis. import master gagal total, posting jurnal salah).",
                "Sebelum lanjut UAT",
            ],
            [
                "High",
                "Fungsi penting salah hasil; workaround sulit (mis. dedup tidak jalan, three-way match keliru).",
                "≤ 2 hari kerja",
            ],
            ["Medium", "Fungsi salah namun ada workaround (mis. label/preview tidak akurat).", "Sebelum go-live"],
            ["Low", "Kosmetik/minor; tidak memengaruhi hasil (mis. teks UI).", "Backlog / pasca go-live"],
        ],
        widths=[2.2, 10.6, 3.4],
    )
    b.h("2.7 Siklus dan Alur Defect", 2)
    b.p(
        "Setiap test case dieksekusi dan diberi status Lulus, Gagal, atau Blocked. Test case "
        "Gagal menghasilkan defect dengan alur: Baru → Triase (severity/prioritas) → Perbaikan "
        "(IT/Developer) → Retest (penguji) → Closed/Reopen. Defect dicatat pada Defect Log "
        "(Bab 8) dan dirujuk pada kolom Catatan / No. Defect di tiap test case."
    )
    b.h("2.8 Konvensi Penomoran Test Case", 2)
    b.table(
        ["Prefix Suite", "Area", "FR Terkait"],
        [
            ["UAT-ENV", "Lingkungan & Smoke Test", "—"],
            ["UAT-MD", "Master Data Import", "FR-07, FR-08, FR-09, FR-10"],
            ["UAT-IMP", "Kontrol Adapter Import", "FR-08, FR-09, FR-10"],
            ["UAT-OS", "Saldo Awal Stok (X20)", "FR-06"],
            ["UAT-PR", "Procurement EBR → SES", "FR-01, FR-02, FR-03"],
            ["UAT-MR", "Mirroring Operasional (Flow C)", "FR-04, FR-05, FR-06, FR-12"],
            ["UAT-FT", "SFTP Feed (Phase 2)", "FR-11"],
            ["UAT-RP", "Pelaporan & Rekonsiliasi", "FR-13, FR-05"],
            ["UAT-TX", "Pajak (PPN/Coretax)", "FR-14"],
            ["UAT-SEC", "Keamanan & Hak Akses", "FR-08, FR-10"],
        ],
        widths=[2.8, 6.8, 6.6],
    )
    b.doc.add_page_break()

    # ---- 3. Traceability ----
    b.h("3. Matriks Keterlacakan (Requirement Traceability Matrix)", 1)
    b.p("Setiap kebutuhan fungsional FSD dipetakan ke satu atau lebih test case UAT untuk memastikan cakupan penuh.")
    b.table(
        ["FR", "Kebutuhan Fungsional", "Test Case UAT"],
        [
            ["FR-01", "Purchase Order EBR → SES", "UAT-PR-01, UAT-PR-02"],
            ["FR-02", "Mirror GR toko ke Odoo", "UAT-PR-03"],
            ["FR-03", "Vendor bill & pembayaran SES", "UAT-PR-04, UAT-PR-05"],
            ["FR-04", "Mirror penjualan POS (X24DN)", "UAT-MR-01, UAT-MR-05"],
            ["FR-05", "Mirror tender & settlement", "UAT-MR-02, UAT-MR-06, UAT-RP-03"],
            ["FR-06", "Mirror mutasi & on-hand", "UAT-OS-01, UAT-OS-02, UAT-OS-03, UAT-MR-03"],
            ["FR-07", "Import master data", "UAT-MD-01..UAT-MD-07"],
            ["FR-08", "Wizard import dengan dry-run", "UAT-MD-03, UAT-IMP-06, UAT-SEC-01"],
            ["FR-09", "Deduplikasi & idempotensi", "UAT-MD-06, UAT-IMP-01, UAT-IMP-02, UAT-IMP-03"],
            ["FR-10", "Log audit import", "UAT-IMP-04, UAT-IMP-05"],
            ["FR-11", "SFTP feed otomatis", "UAT-FT-01..UAT-FT-05"],
            ["FR-12", "Klaim promosi (X31)", "UAT-MR-04"],
            ["FR-13", "Pelaporan keuangan", "UAT-RP-01, UAT-RP-02, UAT-RP-04"],
            ["FR-14", "Pajak Indonesia", "UAT-TX-01, UAT-TX-02, UAT-TX-03"],
        ],
        widths=[1.4, 7.0, 7.8],
    )
    b.doc.add_page_break()

    # ---- 4..7 test cases ----
    b.h("4. Skenario Pengujian Rinci (Detailed Test Cases)", 1)
    b.p(
        "Setiap test case berikut disiapkan untuk dieksekusi dan dicatat hasilnya. Kolom Hasil "
        "Aktual, Status, Penguji/Tanggal, dan Catatan diisi saat eksekusi UAT.",
        italic=True,
    )

    for suite_title, cases in SUITES:
        b.h(suite_title, 2)
        for tc in cases:
            b.testcase(tc)
        b.doc.add_paragraph()

    b.doc.add_page_break()

    # ---- 8. Defect log ----
    b.h("8. Catatan Defect (Defect Log)", 1)
    b.p(
        "Seluruh test case berstatus Gagal/Blocked dicatat pada tabel berikut dan ditindaklanjuti "
        "sesuai alur defect (Bab 2.7)."
    )
    b.table(
        ["No", "Test Case", "Deskripsi Defect", "Severity", "Status", "PIC", "Tgl Selesai"],
        [
            ["", "", "", "", "", "", ""],
            ["", "", "", "", "", "", ""],
            ["", "", "", "", "", "", ""],
            ["", "", "", "", "", "", ""],
            ["", "", "", "", "", "", ""],
        ],
        widths=[1.0, 2.4, 5.0, 2.0, 1.8, 2.0, 2.0],
    )

    # ---- 9. Ringkasan & sign-off ----
    b.h("9. Ringkasan Hasil dan Sign-off UAT", 1)
    b.h("9.1 Ringkasan Eksekusi", 2)
    b.table(
        ["Suite", "Jml Test Case", "Lulus", "Gagal", "Blocked", "% Lulus"],
        [
            ["UAT-ENV — Lingkungan & Smoke", "2", "", "", "", ""],
            ["UAT-MD — Master Data Import", "7", "", "", "", ""],
            ["UAT-IMP — Kontrol Adapter Import", "6", "", "", "", ""],
            ["UAT-OS — Saldo Awal Stok", "3", "", "", "", ""],
            ["UAT-PR — Procurement EBR → SES", "5", "", "", "", ""],
            ["UAT-MR — Mirroring Operasional", "6", "", "", "", ""],
            ["UAT-FT — SFTP Feed", "5", "", "", "", ""],
            ["UAT-RP — Pelaporan & Rekonsiliasi", "4", "", "", "", ""],
            ["UAT-TX — Pajak", "3", "", "", "", ""],
            ["UAT-SEC — Keamanan & Hak Akses", "3", "", "", "", ""],
            ["TOTAL", "44", "", "", "", ""],
        ],
        widths=[6.2, 2.6, 1.8, 1.8, 1.8, 2.0],
    )
    b.h("9.2 Keputusan Penerimaan", 2)
    b.p("Tandai salah satu keputusan berikut berdasarkan hasil eksekusi dan kriteria keluar (Bab 2.5):")
    b.bullets(
        [
            "Diterima (Accepted) — seluruh kriteria keluar terpenuhi; sistem disetujui untuk go-live.",
            "Diterima Bersyarat (Accepted with conditions) — go-live dilanjutkan dengan daftar "
            "defect Medium/Low beserta rencana penyelesaian terlampir.",
            "Ditolak (Rejected) — terdapat defect Critical/High terbuka; UAT diulang setelah perbaikan.",
        ]
    )
    b.h("9.3 Lembar Tanda Tangan Hasil UAT", 2)
    b.table(
        ["Peran", "Nama", "Keputusan", "Tanda Tangan", "Tanggal"],
        [
            ["UAT Lead — Erajaya", "", "", "", ""],
            ["Tester Finance — PT EBR", "", "", "", ""],
            ["Tester Ops/Master Data — PT EBR", "", "", "", ""],
            ["IT / ERP Lead — Erajaya", "", "", "", ""],
            ["Business Owner — Levi's Retail (Sign-off)", "", "", "", ""],
        ],
        widths=[5.0, 3.2, 3.0, 2.5, 2.0],
    )
    b.save("UAT Detail - ERA Busana Retailindo (Levi's) v1.0.docx")


# ============================================================================
# TEST-CASE DATA
# ============================================================================
SUITES = [
    (
        "4.1 Suite UAT-ENV — Lingkungan dan Smoke Test",
        [
            {
                "id": "UAT-ENV-01",
                "title": "Verifikasi modul & menu Retail Import terpasang",
                "fr": "Prasyarat seluruh FR",
                "prio": "Tinggi",
                "pre": "Tenant UAT ter-provision dengan Industry Pack Retail / POS.",
                "steps": [
                    "1. Login sebagai administrator tenant UAT.",
                    "2. Buka Apps ▸ filter Installed; pastikan terpasang: purchase, stock, account, "
                    "point_of_sale, custom_retail_import, custom_approval_engine, "
                    "custom_accounting_full, custom_coretax, custom_report_templates.",
                    '3. Pastikan menu utama "Retail Import" tampil pada menu bar.',
                ],
                "data": "Tenant UAT (mis. levis_uat).",
                "expected": [
                    "Seluruh modul tercantum sebagai Installed.",
                    'Menu "Retail Import" tampil dengan submenu Import Data, Import Logs, dan '
                    "Configuration (Import Profiles, SFTP Feeds).",
                ],
            },
            {
                "id": "UAT-ENV-02",
                "title": "Verifikasi 6 profil import Levi's ter-seed",
                "fr": "FR-07, FR-08",
                "prio": "Tinggi",
                "pre": "UAT-ENV-01 lolos; login sebagai Retail Import Manager.",
                "steps": [
                    "1. Buka Retail Import ▸ Configuration ▸ Import Profiles.",
                    "2. Verifikasi keberadaan profil: levis_x101, levis_coa, levis_company, "
                    "levis_x20, levis_x24, levis_x70d.",
                    "3. Buka levis_x101; periksa file_format=xlsx, data_start_row=3, namespace=levis, "
                    "dan column_map_json terisi.",
                ],
                "data": "Data seed modul custom_retail_import.",
                "expected": [
                    "Keenam profil ada dengan code, file_type, dan column_map_json sesuai jenis file.",
                    "data_start_row sesuai: X101=3, CoA=3, Company=4, X20=2, X24=2, X70D=2.",
                ],
            },
        ],
    ),
    (
        "4.2 Suite UAT-MD — Import Master Data",
        [
            {
                "id": "UAT-MD-01",
                "title": "Import Company / Legal Entity (levis_company)",
                "fr": "FR-07",
                "prio": "Tinggi",
                "pre": "Profil levis_company tersedia; file legal entity EBR disiapkan.",
                "steps": [
                    "1. Retail Import ▸ Import Data.",
                    "2. Pilih Profile = levis_company, unggah file legal entity, klik Import.",
                    "3. Buka Settings ▸ Companies ▸ PT Era Busana Retailindo.",
                ],
                "data": "Sheet key/value EBR (Nama, NPWP, Alamat, Telepon, Email).",
                "expected": [
                    "res.company terisi: name, NPWP→vat, alamat→street, phone, email.",
                    "Log import berstatus imported.",
                ],
            },
            {
                "id": "UAT-MD-02",
                "title": "Import Chart of Accounts (levis_coa)",
                "fr": "FR-07",
                "prio": "Tinggi",
                "pre": "Profil levis_coa tersedia; CoA EBR.xlsx clean.",
                "steps": [
                    "1. Retail Import ▸ Import Data; Profile = levis_coa; unggah CoA EBR.xlsx.",
                    "2. Klik Preview (dry-run) lalu Import.",
                    "3. Buka Accounting ▸ Configuration ▸ Chart of Accounts.",
                ],
                "data": "CoA EBR.xlsx (code, name, account_type).",
                "expected": [
                    "Jumlah account.account = jumlah baris CoA; account_type valid sesuai enum Odoo.",
                    "Re-import memperbarui nama akun yang sudah ada, bukan menggandakan.",
                ],
            },
            {
                "id": "UAT-MD-03",
                "title": "Dry-run preview sebelum commit (X101)",
                "fr": "FR-08",
                "prio": "Tinggi",
                "pre": "Profil levis_x101 tersedia.",
                "steps": [
                    "1. Retail Import ▸ Import Data; Profile = levis_x101; unggah X101.",
                    '2. Klik tombol "Preview (dry-run)".',
                    "3. Amati panel Preview; jangan klik Import.",
                    "4. Cek Import Logs dan product.template count.",
                ],
                "data": "X101_Material_Master.xlsx.",
                "expected": [
                    "Panel Preview menampilkan profil, jumlah sample row terparse, jumlah blank "
                    "skipped, dan sampel field logis per baris (mis. product_code, sku, size).",
                    "TIDAK ada record dibuat dan tidak ada log import baru (dry-run, no commit).",
                ],
            },
            {
                "id": "UAT-MD-04",
                "title": "Import master produk X101 (asinkron via queue_job)",
                "fr": "FR-07, FR-08",
                "prio": "Tinggi",
                "pre": "queue_job runner aktif; profil levis_x101 tersedia.",
                "steps": [
                    "1. Retail Import ▸ Import Data; Profile = levis_x101; unggah X101; klik Import.",
                    "2. Amati log import berpindah ke state queued lalu running (job_uuid terisi).",
                    "3. Tunggu job selesai (±30 menit); buka log saat state imported.",
                    "4. Verifikasi Inventory ▸ Products: count template & varian, default_code, "
                    "barcode, atribut, kategori.",
                ],
                "data": "X101 full (±14.885 artikel / ±159.658 SKU; Size 89, Inseam 24, kategori 170).",
                "expected": [
                    "Import dijalankan asinkron (job_uuid terisi) tanpa ERR_EMPTY_RESPONSE.",
                    "product.template ±14.885; product.product ±159.658 dengan default_code=SKU dan barcode=GTIN.",
                    "Kategori 3 tingkat (CATEGORY>CLASS>SUBCLASS) dan atribut Size+Inseam terbentuk.",
                    "Log: state imported, records_created/matched terisi, error_count = 0 (atau "
                    "partial dengan error per baris bila ada).",
                ],
            },
            {
                "id": "UAT-MD-05",
                "title": "Perbaikan encoding U+FFFD → ® pada deskripsi produk",
                "fr": "FR-07",
                "prio": "Sedang",
                "pre": 'X101 berisi karakter rusak (mis. "LEVI\'S�"); fix_encoding aktif pada profil.',
                "steps": [
                    "1. Import X101 (atau gunakan hasil UAT-MD-04).",
                    "2. Buka beberapa produk yang deskripsinya mengandung simbol merek dagang.",
                ],
                "data": "Baris X101 dengan karakter U+FFFD.",
                "expected": [
                    "Karakter rusak U+FFFD tergantikan '®' pada deskripsi/nama produk; tidak ada "
                    "karakter rusak tersisa pada field string.",
                ],
            },
            {
                "id": "UAT-MD-06",
                "title": "Idempotensi re-import X101 (external ID, tanpa duplikasi)",
                "fr": "FR-09",
                "prio": "Tinggi",
                "pre": "UAT-MD-04 selesai (X101 sudah terimpor).",
                "steps": [
                    "1. Catat jumlah product.template dan product.product saat ini.",
                    '2. Re-import file X101 yang sama; karena dedup file aktif, centang "Force '
                    're-import" agar diproses ulang.',
                    "3. Tunggu selesai; bandingkan kembali jumlah record.",
                ],
                "data": "File X101 identik.",
                "expected": [
                    "Jumlah template & varian TIDAK bertambah (record di-update via external ID "
                    "namespace 'levis', bukan diduplikasi).",
                    "Log mencatat records_matched > 0 dan records_created ≈ 0.",
                ],
            },
            {
                "id": "UAT-MD-07",
                "title": "Pembentukan warehouse dari Store Master",
                "fr": "FR-07",
                "prio": "Tinggi",
                "pre": "Store Master tersedia; loader Track A (04_load_stores.py) disiapkan.",
                "steps": [
                    "1. Jalankan loader store (Track A) atau profil store_master sesuai prosedur.",
                    "2. Buka Inventory ▸ Configuration ▸ Warehouses.",
                    "3. Verifikasi external ID wh_<storecode> terbentuk untuk tiap toko.",
                ],
                "data": "LEVI'S - STORE MASTER DATA.xlsx (24 lokasi: 1 DC + 23 toko).",
                "expected": [
                    "Warehouse terbentuk per toko (target 24); external ID wh_<storecode> menjadi kunci join file X.",
                    "Catatan: bila kode toko 23/24 belum tersedia (isu terbuka #2), hanya toko "
                    "berkode yang terbentuk — dicatat sebagai Blocked, bukan Gagal.",
                ],
            },
        ],
    ),
    (
        "4.3 Suite UAT-IMP — Kontrol Adapter Import",
        [
            {
                "id": "UAT-IMP-01",
                "title": "Deduplikasi file identik (SHA256) ditolak",
                "fr": "FR-09",
                "prio": "Tinggi",
                "pre": "Sebuah file telah berhasil diimpor (mis. CoA dari UAT-MD-02).",
                "steps": [
                    "1. Retail Import ▸ Import Data; pilih profil yang sama; unggah file yang sama persis.",
                    '2. Pastikan "Force re-import" TIDAK dicentang; klik Import.',
                ],
                "data": "File identik dengan import sebelumnya.",
                "expected": [
                    'Import ditolak dengan pesan: "This exact file was already imported/queued '
                    "(log #N). Tick 'Force re-import' to override.\"",
                    "Tidak ada log baru berstatus imported dibuat.",
                ],
            },
            {
                "id": "UAT-IMP-02",
                "title": "Force re-import menimpa penolakan dedup",
                "fr": "FR-09",
                "prio": "Sedang",
                "pre": "UAT-IMP-01 menampilkan penolakan dedup.",
                "steps": [
                    '1. Pada wizard, centang "Force re-import".',
                    "2. Klik Import; buka Import Logs.",
                ],
                "data": "File identik + opsi Force re-import.",
                "expected": [
                    "Import diproses dan log baru terbentuk; tindakan force tercatat (file_hash sama, "
                    "log baru terpisah).",
                    "Untuk master data, idempotensi external ID tetap mencegah duplikasi record.",
                ],
            },
            {
                "id": "UAT-IMP-03",
                "title": "File gagal tidak memblokir unggah ulang",
                "fr": "FR-09",
                "prio": "Sedang",
                "pre": "Tersedia file rusak yang menyebabkan import failed.",
                "steps": [
                    "1. Unggah file rusak; amati log berstatus failed.",
                    "2. Perbaiki file; unggah ulang TANPA Force re-import.",
                ],
                "data": "File rusak lalu versi perbaikannya.",
                "expected": [
                    "Status failed dikecualikan dari guard dedup, sehingga unggah ulang tidak "
                    "tertolak sebagai duplikat dan dapat diproses.",
                ],
            },
            {
                "id": "UAT-IMP-04",
                "title": "Kelengkapan log audit & arsip file sumber",
                "fr": "FR-10",
                "prio": "Tinggi",
                "pre": "Minimal satu import berhasil.",
                "steps": [
                    "1. Retail Import ▸ Import Logs; buka satu record imported.",
                    "2. Periksa field: profile, file_hash, records_created/matched/skipped, "
                    "error_count, imported_by, imported_at, dan attachment file sumber.",
                ],
                "data": "Log import hasil suite UAT-MD.",
                "expected": [
                    "Log mencatat siapa (imported_by), kapan (imported_at), file apa (file_hash), "
                    "dan penghitung record per tahap.",
                    "File sumber terarsip sebagai ir.attachment yang dapat diunduh kembali.",
                ],
            },
            {
                "id": "UAT-IMP-05",
                "title": "Import partial — pelaporan error per baris",
                "fr": "FR-10",
                "prio": "Sedang",
                "pre": "Tersedia file dengan sebagian baris bermasalah (mis. account_type invalid).",
                "steps": [
                    "1. Unggah file yang sebagian barisnya error; klik Import.",
                    "2. Buka log; periksa state dan ringkasan error (raw_payload).",
                ],
                "data": "File dengan ±beberapa baris invalid.",
                "expected": [
                    "State akhir = partial; baris valid tetap terimpor, baris error tercatat.",
                    "raw_payload memuat pasangan (nomor baris sumber, pesan) maksimal 200 baris "
                    "pertama untuk koreksi presisi.",
                ],
            },
            {
                "id": "UAT-IMP-06",
                "title": "Dispatch asinkron vs sinkron sesuai jenis file",
                "fr": "FR-08",
                "prio": "Sedang",
                "pre": "Profil tersedia untuk file besar (X101) dan kecil (CoA).",
                "steps": [
                    "1. Import X101 → amati log queued/running dengan job_uuid (asinkron).",
                    "2. Import CoA → amati log langsung imported tanpa job_uuid (sinkron).",
                ],
                "data": "X101 (besar) dan CoA (kecil).",
                "expected": [
                    "Jenis file besar (x101, x20, x24, x70d, x32p) diproses asinkron via queue_job (job_uuid terisi).",
                    "Jenis file kecil diproses sinkron dan langsung imported.",
                ],
            },
        ],
    ),
    (
        "4.4 Suite UAT-OS — Saldo Awal Stok (X20)",
        [
            {
                "id": "UAT-OS-01",
                "title": "Import saldo awal on-hand (X20) ke stock.quant",
                "fr": "FR-06",
                "prio": "Tinggi",
                "pre": "Master produk (UAT-MD-04) dan warehouse (UAT-MD-07) selesai.",
                "steps": [
                    "1. Retail Import ▸ Import Data; Profile = levis_x20; unggah X20; klik Import.",
                    "2. Tunggu selesai; buka Inventory ▸ Reporting ▸ Stock / On Hand.",
                    "3. Bandingkan total on-hand terhadap total baris X20.",
                ],
                "data": "X20_Current_Onhand_Inventory (sampel store 14694).",
                "expected": [
                    "stock.quant saldo awal terbentuk per warehouse toko per SKU.",
                    "Total quant = total on-hand pada file X20; log imported.",
                ],
            },
            {
                "id": "UAT-OS-02",
                "title": "Guard one-shot saldo awal menolak re-run",
                "fr": "FR-06",
                "prio": "Tinggi",
                "pre": "UAT-OS-01 sukses (saldo awal sudah diterapkan).",
                "steps": [
                    "1. Coba import X20 lagi untuk profil yang sama.",
                    "2. Amati respons sistem.",
                ],
                "data": "File X20 (ulang).",
                "expected": [
                    "Sistem menolak eksekusi ulang (guard one-shot via penanda; mencegah penggandaan on-hand).",
                    "On-hand tidak berubah/berlipat.",
                ],
            },
            {
                "id": "UAT-OS-03",
                "title": "Resolusi varian & lokasi pada X20",
                "fr": "FR-06",
                "prio": "Sedang",
                "pre": "Master produk & warehouse tersedia; X20 berisi EAN/ITEM ID & kode toko.",
                "steps": [
                    "1. Pada saat/setelah import X20, periksa beberapa baris.",
                    "2. Verifikasi varian teridentifikasi via barcode(EAN) atau default_code(ITEM ID), "
                    "dan lokasi via wh_<storecode>.",
                ],
                "data": "Baris X20 dengan EAN, ITEM ID, kode toko.",
                "expected": [
                    "Varian benar termatch; quant masuk ke lokasi warehouse toko yang tepat.",
                    "Baris yang varian/lokasinya tidak ditemukan dilaporkan sebagai error baris "
                    "(bukan diam-diam dilewati).",
                ],
            },
        ],
    ),
    (
        "4.5 Suite UAT-PR — Procurement EBR → SES",
        [
            {
                "id": "UAT-PR-01",
                "title": "Buat Purchase Order EBR ke vendor PT SES",
                "fr": "FR-01",
                "prio": "Tinggi",
                "pre": "Vendor PT SES dan master produk tersedia.",
                "steps": [
                    "1. Purchase ▸ Orders ▸ New; Vendor = PT Sinar Eka Selaras.",
                    "2. Tambah baris produk Levi's, qty, harga; simpan.",
                ],
                "data": "Beberapa SKU Levi's, vendor SES.",
                "expected": [
                    "PO tersimpan sebagai RFQ/Draft dengan vendor SES dan baris produk benar; subtotal "
                    "& pajak terhitung.",
                ],
            },
            {
                "id": "UAT-PR-02",
                "title": "Approval berjenjang PO via approval engine",
                "fr": "FR-01",
                "prio": "Tinggi",
                "pre": "Matriks otorisasi custom_approval_engine terkonfigurasi (mis. threshold nilai).",
                "steps": [
                    "1. Buat PO bernilai DI BAWAH threshold; konfirmasi.",
                    "2. Buat PO bernilai DI ATAS threshold; coba konfirmasi tanpa approval.",
                    "3. Approver menyetujui PO bernilai besar; konfirmasi ulang.",
                ],
                "data": "Dua PO: nilai kecil & nilai besar.",
                "expected": [
                    "PO kecil dapat dikonfirmasi sesuai aturan; PO besar tidak dapat dikonfirmasi "
                    "sebelum disetujui approver sesuai matriks.",
                    "Setelah disetujui, PO besar dapat dikonfirmasi dan dikirim ke SES.",
                ],
            },
            {
                "id": "UAT-PR-03",
                "title": "Mirror GR toko memvalidasi receipt atas PO",
                "fr": "FR-02",
                "prio": "Tinggi",
                "pre": "PO EBR→SES terkonfirmasi; file GR/X32P (atau receipt) disiapkan.",
                "steps": [
                    "1. Proses mirror GR toko ke Odoo (sesuai mekanisme yang ditetapkan untuk receipt).",
                    "2. Buka receipt terkait PO; validasi qty.",
                    "3. Periksa stok warehouse toko.",
                ],
                "data": "Data GR toko untuk PO uji.",
                "expected": [
                    "Receipt atas PO EBR tervalidasi; stok warehouse toko bertambah sesuai qty diterima.",
                    "Selisih qty (kurang/lebih) terlihat untuk ditindaklanjuti sebagai discrepancy.",
                ],
            },
            {
                "id": "UAT-PR-04",
                "title": "Vendor bill three-way match (PO–GR–Bill)",
                "fr": "FR-03",
                "prio": "Tinggi",
                "pre": "PO terkonfirmasi & receipt tervalidasi (UAT-PR-03).",
                "steps": [
                    "1. Dari PO, buat Bill; cocokkan terhadap PO dan GR.",
                    "2. Uji skenario mismatch (qty/harga berbeda) dan amati indikasi selisih.",
                    "3. Posting bill saat cocok.",
                ],
                "data": "Invoice SES sesuai & tidak sesuai.",
                "expected": [
                    "Bill tercocokkan dengan PO & GR; mismatch qty/harga ter-flag sebelum posting.",
                    "Bill cocok dapat di-post; AP & jurnal terbentuk benar.",
                ],
            },
            {
                "id": "UAT-PR-05",
                "title": "Pembayaran kepada SES dan pelunasan AP",
                "fr": "FR-03",
                "prio": "Tinggi",
                "pre": "Vendor bill ter-post (UAT-PR-04).",
                "steps": [
                    "1. Register Payment atas bill SES; pilih jurnal bank EBR.",
                    "2. Verifikasi status pembayaran & saldo AP.",
                ],
                "data": "Bill SES ter-post; jurnal bank EBR.",
                "expected": [
                    "Pembayaran tercatat; bill berstatus Paid; saldo AP terhadap SES berkurang sesuai "
                    "nilai pembayaran.",
                ],
            },
        ],
    ),
    (
        "4.6 Suite UAT-MR — Mirroring Operasional (Flow C)",
        [
            {
                "id": "UAT-MR-01",
                "title": "Gating Phase 5 — import penjualan X24 menahan posting",
                "fr": "FR-04",
                "prio": "Tinggi",
                "pre": "Profil levis_x24 tersedia; keputusan Phase 5 BELUM ditetapkan.",
                "steps": [
                    "1. Retail Import ▸ Import Data; Profile = levis_x24; unggah X24DN; klik Import.",
                    "2. Amati respons sistem dan state log.",
                ],
                "data": "X24DN_Retail_Sales_Detail (sampel store 14694).",
                "expected": [
                    "Executor memparsing & mengelompokkan transaksi, namun menahan posting (UserError "
                    "berisi panduan: butuh keputusan representasi POS, histori, payment method, tax "
                    "mapping serta prasyarat pos.config/session).",
                    "Tidak ada pos.order/jurnal penjualan ter-post (perilaku gating sesuai harapan).",
                ],
            },
            {
                "id": "UAT-MR-02",
                "title": "Gating Phase 5 — import tender X70D menahan posting",
                "fr": "FR-05",
                "prio": "Tinggi",
                "pre": "Profil levis_x70d tersedia; keputusan Phase 5 BELUM ditetapkan.",
                "steps": [
                    "1. Import X70D via Profile = levis_x70d.",
                    "2. Amati state log dan pesan.",
                ],
                "data": "X70D_Tender_Detail (sampel store 14694).",
                "expected": [
                    "Tender terparse & terkelompok namun posting tertahan (gating Phase 5); tidak ada "
                    "pos.payment ter-post.",
                ],
            },
            {
                "id": "UAT-MR-03",
                "title": "Staged — mutasi persediaan X32P di-parse & dihitung",
                "fr": "FR-06",
                "prio": "Sedang",
                "pre": "Profil X32P tersedia.",
                "steps": [
                    "1. Import X32P.",
                    "2. Buka log; periksa line_count dan attachment.",
                ],
                "data": "X32P_Stock_Movement_With_Price.",
                "expected": [
                    "File terparse, jumlah baris terhitung, dan attachment diarsipkan; belum menulis "
                    "model stok (staged, reference/audit) — sesuai keputusan Phase 3/5.",
                ],
            },
            {
                "id": "UAT-MR-04",
                "title": "Staged — klaim promosi X31 di-parse & dihitung",
                "fr": "FR-12",
                "prio": "Rendah",
                "pre": "Profil X31 tersedia.",
                "steps": [
                    "1. Import X31_Discount_Journal.",
                    "2. Buka log; periksa jumlah baris & attachment.",
                ],
                "data": "X31_Discount_Journal.",
                "expected": [
                    "File terparse & dihitung, attachment tersimpan sebagai dasar klaim promosi; "
                    "belum menulis model (staged).",
                ],
            },
            {
                "id": "UAT-MR-05",
                "title": "(Pasca Phase 5) Posting penjualan X24 → pos.order + jurnal",
                "fr": "FR-04",
                "prio": "Tinggi",
                "pre": "Keputusan Phase 5 ditetapkan; pos.config/session/payment method disiapkan. "
                "Dieksekusi pada siklus UAT lanjutan bila Phase 5 sudah aktif.",
                "steps": [
                    "1. Aktifkan gating sesuai keputusan; import X24.",
                    "2. Verifikasi pos.order/pos.order.line terbentuk dan jurnal penjualan + HPP "
                    "terbentuk TANPA menggerakkan stok (on-hand sudah dari X20).",
                ],
                "data": "X24DN pasca keputusan Phase 5.",
                "expected": [
                    "pos.order ter-post; jurnal penjualan & HPP benar; stok tidak bergerak ganda.",
                    "Bila Phase 5 belum aktif saat UAT, test case ini dicatat Blocked.",
                ],
            },
            {
                "id": "UAT-MR-06",
                "title": "(Pasca Phase 5) Tender X70D → pos.payment ter-join ke X24",
                "fr": "FR-05",
                "prio": "Tinggi",
                "pre": "UAT-MR-05 aktif; X70D & X24 untuk periode yang sama tersedia.",
                "steps": [
                    "1. Import X70D pasca Phase 5.",
                    "2. Verifikasi pos.payment ter-join ke order X24 via kunci (store, date, register, transnum).",
                ],
                "data": "X70D + X24 periode sama.",
                "expected": [
                    "pos.payment terbentuk dan tertaut ke order yang benar; total tender = total "
                    "penjualan per transaksi.",
                    "Bila Phase 5 belum aktif, dicatat Blocked.",
                ],
            },
        ],
    ),
    (
        "4.7 Suite UAT-FT — SFTP Feed (Phase 2)",
        [
            {
                "id": "UAT-FT-01",
                "title": "Test Connection feed SFTP",
                "fr": "FR-11",
                "prio": "Tinggi",
                "pre": "paramiko terpasang di image; akses SFTP Erajaya (host/kredensial/folder) "
                "tersedia; login sebagai Retail Import Manager.",
                "steps": [
                    "1. Retail Import ▸ Configuration ▸ SFTP Feeds ▸ New.",
                    "2. Isi name, profile, host, port, username, auth_type, password_param/"
                    "private_key_path, remote_dir, file_glob (mis. X20_*.csv).",
                    '3. Klik tombol "Test Connection".',
                ],
                "data": "Konfigurasi SFTP Erajaya + glob per jenis file.",
                "expected": [
                    'Muncul notifikasi "Connection OK" dengan pesan jumlah file yang cocok dengan '
                    'glob pada remote_dir (mis. "3 file(s) match X20_*.csv in /incoming").',
                ],
            },
            {
                "id": "UAT-FT-02",
                "title": "Poll Now menarik & memproses file baru",
                "fr": "FR-11",
                "prio": "Tinggi",
                "pre": "UAT-FT-01 sukses; ada file baru di remote_dir sesuai glob.",
                "steps": [
                    '1. Pada form feed, klik "Poll Now".',
                    "2. Periksa field Last run/Last status/Files imported.",
                    "3. Buka Import Logs; verifikasi log per file ter-create dan diproses executor.",
                ],
                "data": "File baru di SFTP.",
                "expected": [
                    "File baru terunduh, log dibuat (executor sama dengan wizard), last_status = ok, "
                    "files_imported bertambah, last_message mencatat jumlah file baru.",
                ],
            },
            {
                "id": "UAT-FT-03",
                "title": "Dedup skip pada polling berulang",
                "fr": "FR-11, FR-09",
                "prio": "Sedang",
                "pre": "UAT-FT-02 selesai; file di remote tidak berubah.",
                "steps": [
                    '1. Klik "Poll Now" kembali tanpa menambah file baru.',
                    "2. Periksa Import Logs.",
                ],
                "data": "Direktori remote tanpa file baru.",
                "expected": [
                    "File yang hash-nya sudah dikenal dilewati diam-diam (skip); tidak ada log/import "
                    "ganda. File 0 byte juga dilewati.",
                ],
            },
            {
                "id": "UAT-FT-04",
                "title": "Cron poller terjadwal",
                "fr": "FR-11",
                "prio": "Sedang",
                "pre": 'Feed aktif; cron "Retail Import: poll SFTP feeds" diaktifkan.',
                "steps": [
                    "1. Settings ▸ Technical ▸ Scheduled Actions; aktifkan cron poll feeds (default nonaktif).",
                    "2. Picu manual (Run) atau tunggu jadwal; periksa feed Last run.",
                ],
                "data": "Cron poll feeds + file baru di SFTP.",
                "expected": [
                    "Cron memproses seluruh feed aktif; file baru terimpor otomatis tanpa intervensi "
                    "manual; default cron nonaktif terverifikasi.",
                ],
            },
            {
                "id": "UAT-FT-05",
                "title": "Penanganan error koneksi feed",
                "fr": "FR-11",
                "prio": "Sedang",
                "pre": "Feed dengan kredensial/host sengaja salah.",
                "steps": [
                    '1. Set host/kredensial salah; klik "Poll Now".',
                    "2. Periksa Last status & Last message.",
                ],
                "data": "Konfigurasi feed invalid.",
                "expected": [
                    "last_status = error dan last_message memuat pesan kegagalan; sistem tidak crash "
                    "dan feed lain tidak terdampak.",
                ],
            },
        ],
    ),
    (
        "4.8 Suite UAT-RP — Pelaporan dan Rekonsiliasi",
        [
            {
                "id": "UAT-RP-01",
                "title": "Laporan keuangan lengkap (Excel/PDF)",
                "fr": "FR-13",
                "prio": "Tinggi",
                "pre": "Modul custom_accounting_full terpasang; ada transaksi (PO/Bill/Payment).",
                "steps": [
                    "1. Accounting ▸ Reporting; jalankan General Ledger, Trial Balance, P&L, Balance Sheet.",
                    "2. Ekspor ke Excel dan PDF.",
                ],
                "data": "Transaksi hasil suite UAT-PR.",
                "expected": [
                    "Laporan keuangan tersedia & akurat; ekspor Excel/PDF berhasil dengan format rapi.",
                ],
            },
            {
                "id": "UAT-RP-02",
                "title": "Template PDF transaksi ber-branding per company",
                "fr": "FR-13",
                "prio": "Sedang",
                "pre": "Modul custom_report_templates terpasang; branding EBR dikonfigurasi.",
                "steps": [
                    "1. Cetak PDF Purchase Order, Quotation, dan Invoice.",
                    "2. Periksa header/branding, format mata uang, dan tata letak.",
                ],
                "data": "PO/Quotation/Invoice EBR.",
                "expected": [
                    "PDF tampil dengan branding PT EBR; format uang benar (tanpa artefak NBSP/Â); "
                    "layout sesuai template.",
                ],
            },
            {
                "id": "UAT-RP-03",
                "title": "Rekonsiliasi settlement tender (X70T) vs bank",
                "fr": "FR-05",
                "prio": "Sedang",
                "pre": "Data X70T & mutasi bank tersedia (Phase 5/manual).",
                "steps": [
                    "1. Bandingkan settlement tender per acquirer terhadap penerimaan bank per hari.",
                    "2. Identifikasi selisih fee acquirer.",
                ],
                "data": "X70T_Tender_Settlement + rekening koran.",
                "expected": [
                    "Settlement terekonsiliasi per acquirer; selisih fee teridentifikasi & dapat "
                    "dibukukan; selisih abnormal ter-flag untuk eskalasi.",
                ],
            },
            {
                "id": "UAT-RP-04",
                "title": "Valuasi persediaan mencerminkan saldo awal & mutasi",
                "fr": "FR-13, FR-06",
                "prio": "Sedang",
                "pre": "Saldo awal X20 (UAT-OS-01) & receipt (UAT-PR-03) selesai.",
                "steps": [
                    "1. Inventory ▸ Reporting ▸ Inventory Valuation.",
                    "2. Cocokkan nilai persediaan terhadap saldo awal + penerimaan.",
                ],
                "data": "Hasil X20 + receipt PO.",
                "expected": [
                    "Nilai persediaan akurat mencerminkan saldo awal X20 dan mutasi penerimaan; "
                    "akun persediaan di GL konsisten.",
                ],
            },
        ],
    ),
    (
        "4.9 Suite UAT-TX — Pajak (PPN / Coretax)",
        [
            {
                "id": "UAT-TX-01",
                "title": "Konfigurasi & penerapan PPN pada PO/Bill",
                "fr": "FR-14",
                "prio": "Tinggi",
                "pre": "Pajak PPN keluaran/masukan terkonfigurasi sesuai tarif berlaku.",
                "steps": [
                    "1. Buat PO/Bill dengan produk kena PPN.",
                    "2. Verifikasi perhitungan PPN dan akun pajak.",
                ],
                "data": "Transaksi kena PPN.",
                "expected": [
                    "PPN dihitung sesuai tarif; jurnal pajak masuk akun yang benar.",
                ],
            },
            {
                "id": "UAT-TX-02",
                "title": "Ekspor e-Faktur / Coretax",
                "fr": "FR-14",
                "prio": "Sedang",
                "pre": "Modul custom_coretax terpasang; NPWP & data Coretax terisi.",
                "steps": [
                    "1. Jalankan fungsi ekspor e-Faktur/Coretax pada invoice/bill terkait.",
                    "2. Periksa output sesuai format Coretax.",
                ],
                "data": "Invoice/bill ber-PPN.",
                "expected": [
                    "Output e-Faktur/Coretax dihasilkan sesuai format; data NPWP & nilai pajak benar.",
                ],
            },
            {
                "id": "UAT-TX-03",
                "title": "Bukti potong (custom_coretax_bupot)",
                "fr": "FR-14",
                "prio": "Rendah",
                "pre": "Modul custom_coretax_bupot terpasang; ada transaksi relevan.",
                "steps": [
                    "1. Buat/cetak bukti potong pada transaksi relevan.",
                    "2. Periksa nilai & format.",
                ],
                "data": "Transaksi dengan pemotongan pajak.",
                "expected": [
                    "Bukti potong dihasilkan dengan nilai & format sesuai ketentuan.",
                ],
            },
        ],
    ),
    (
        "4.10 Suite UAT-SEC — Keamanan dan Hak Akses",
        [
            {
                "id": "UAT-SEC-01",
                "title": "Hak akses Retail Import User",
                "fr": "FR-08, FR-10",
                "prio": "Tinggi",
                "pre": 'Akun dengan grup "Retail Import / User" (bukan Manager).',
                "steps": [
                    "1. Login sebagai user grup User.",
                    "2. Buka Retail Import; coba akses Import Data, Import Logs, dan Configuration.",
                ],
                "data": "Akun grup Retail Import User.",
                "expected": [
                    "User dapat menjalankan wizard import dan melihat log; menu Configuration "
                    "(Profiles, SFTP Feeds) TIDAK tampil/teramankan.",
                ],
            },
            {
                "id": "UAT-SEC-02",
                "title": "Hak akses Retail Import Manager",
                "fr": "FR-08, FR-10",
                "prio": "Sedang",
                "pre": 'Akun dengan grup "Retail Import / Manager".',
                "steps": [
                    "1. Login sebagai Manager.",
                    "2. Akses Configuration ▸ Import Profiles & SFTP Feeds; ubah profil; lakukan Force re-import.",
                ],
                "data": "Akun grup Retail Import Manager.",
                "expected": [
                    "Manager dapat mengelola profil & feed dan melakukan force re-import (mewarisi hak "
                    "User); pemisahan tugas (SoD) terjaga.",
                ],
            },
            {
                "id": "UAT-SEC-03",
                "title": "Penyimpanan kredensial SFTP terbatas",
                "fr": "FR-11",
                "prio": "Sedang",
                "pre": "Feed SFTP dengan auth password menggunakan ir.config_parameter.",
                "steps": [
                    "1. Verifikasi password disimpan pada ir.config_parameter (kunci pada "
                    "password_param), bukan plaintext pada form.",
                    "2. Verifikasi akses baca parameter dibatasi; uji opsi auth key-based.",
                ],
                "data": "Konfigurasi feed password & key.",
                "expected": [
                    "Secret SFTP dirujuk via ir.config_parameter; akses baca dibatasi; auth key-based "
                    "berfungsi sebagai alternatif yang lebih aman (catatan: belum ada enkripsi "
                    "at-rest — risiko terdokumentasi).",
                ],
            },
        ],
    ),
]


if __name__ == "__main__":
    build_fsd()
    build_uat()
